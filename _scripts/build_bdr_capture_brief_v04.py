#!/usr/bin/env python3
"""
build_bdr_capture_brief_v04.py — Build BDR-FLEET-READINESS capture brief v0.4.

v0.4 fixes per operator feedback 2026-05-26 (fourth round):
  1. Back off confirmatory tone. Brief presents a hypothesis backed by adjacent evidence, not a
     verified finding.
  2. Expand every acronym on first appearance. Format: "Spelled-out form (ACRONYM)."
  3. Drop "operator" and "analyst" role labels. Single author voice. OCI section rewritten so it
     stands without invoking either label.
"""

from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

VAULT_ROOT = Path(__file__).parent.parent
OUT = VAULT_ROOT / "opportunities/BDR-FLEET-READINESS/04_artifacts/capture-brief-v0.4-draft.docx"


def _set_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(18)


def _h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(14)


def _h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(12)


def _body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def _bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(f"• {text}")


def _italic(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.italic = True


def _pb(doc):
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_styles(doc)

    # Cover
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(36)
    r = cover.add_run("CAPTURE BRIEF — v0.4 DRAFT\n")
    r.font.bold = True
    r.font.size = Pt(22)

    r = cover.add_run("\nBDR-FLEET-READINESS\n")
    r.font.bold = True
    r.font.size = Pt(18)

    r = cover.add_run("\nFive-level analog-first progression of gamified decision-rehearsal scenarios for U.S. Navy Battle Damage Assessment and Repair (BDAR) teams and the fleet commanders they support\n")
    r.font.size = Pt(13)
    r.font.italic = True

    r = cover.add_run(f"\nDraft date: {date.today().isoformat()}  |  For CACI executive review (open-source-only).\n")
    r.font.size = Pt(10)

    r = cover.add_run("\nv0.4 changes from v0.3: tone tightened away from confirmatory framing; every acronym expanded on first appearance; single-author voice (the \"operator\" and \"analyst\" labels removed).\n")
    r.font.size = Pt(9)
    r.font.italic = True

    _pb(doc)

    # 1. BLUF
    _h1(doc, "1. Bottom Line Up Front")

    _body(doc,
        "The U.S. Navy is publicly signaling demand for better wartime ship-repair readiness. The rest of the "
        "Department of Defense (DoD) buys related contractor work at multi-hundred-million-dollar scale. The Navy "
        "could plausibly buy this kind of work too. As of this draft date, public sources do not show a Navy "
        "procurement line for the specific sub-product proposed here. That gap is the opportunity this brief asks "
        "CACI leadership to consider.")

    _body(doc,
        "The proposed product is a five-level progression of gamified decision-rehearsal scenarios. Level 1 is a "
        "five- to fifteen-minute flash drill on a printed scenario card. Level 2 is a thirty- to forty-five-minute "
        "linked-decision sequence. Level 3 is the full one-hour gamified wardroom session that integrates the six "
        "operational-decision moments. Level 4 is a two- to four-hour multi-session campaign across a deployment "
        "cycle. Level 5 is software-driven scenario injection into live fleet exercises. The progression starts "
        "analog and moves to software at Level 4, where artificial intelligence (AI) scenario generation is "
        "justified by scale. The same scenario content serves two audiences: the wardroom at a Navy repair "
        "activity preparing to receive a damaged ship, and the fleet commander deciding where to send her. Port "
        "selection is the canonical scenario serving both audiences.")

    _body(doc,
        "The proposed procurement path uses contracting vehicles CACI already holds. CACI National Security "
        "Solutions, LLC (CACI NSS, LLC) is a CACI subsidiary that holds three governmentwide acquisition vehicles "
        "open to any federal customer including any U.S. Navy command: General Services Administration (GSA) "
        "One Acquisition Solution for Integrated Services (OASIS); the National Institutes of Health Information "
        "Technology Acquisition and Assessment Center Chief Information Officer–Solutions and Partners 3 contract "
        "(NITAAC CIO-SP3); and GSA Multiple Award Schedule (MAS). The Navy could issue a task order against any of "
        "these vehicles without standing up a new Navy-specific contract. The $194 million task order CACI NSS, LLC "
        "currently holds for support to U.S. Africa Command (AFRICOM) is the precedent.")

    _body(doc,
        "The recommendation, summarized at section 7, asks CACI executive leadership to consider initiating "
        "engagement with one of the named Navy commands and pitching a Level 1 pilot task order against the "
        "existing OASIS or CIO-SP3 vehicle. The Navy is not running a Request for Proposal (RFP) for this work "
        "today. The path to procurement is likely to require proactive executive-level engagement. An alternative "
        "entry path uses the Defense Innovation Unit (DIU), Small Business Innovation Research (SBIR), or Small "
        "Business Technology Transfer (STTR) pilot mechanisms at smaller per-engagement scale.")

    _pb(doc)

    # 2. Demand signal
    _h1(doc, "2. Demand signal")

    _body(doc,
        "Four independent signals point at the corrected-scope demand. Senior Navy leadership is publicly "
        "committing dollar volume to wartime ship-repair readiness. The Navy is publicly executing the most "
        "directly relevant exercise type in the most directly relevant venue. The Navy professional community is "
        "publicly arguing the gap is real. And the Department of Defense procures related contractor work-types "
        "at multi-hundred-million-dollar scale, though under three different work-type categories rather than one.")

    _h2(doc, "2.1 Senior-leadership signal")

    _body(doc,
        "Admiral Daryl L. Caudle, Chief of Naval Operations (CNO), delivered a Statement on the Posture of the "
        "United States Navy before the House Armed Services Committee (HASC) on 14 May 2026, naming maintenance a "
        "warfighting requirement and announcing a planned deliberate study of Navy yard capacity. The Department "
        "of the Navy Fiscal Year 2027 (FY27) President's Budget Press Brief funds Ship Maintenance at $17.0 "
        "billion explicitly to drive the fleet toward an 80-percent Combat Surge Ready (CSR) posture. The same "
        "brief includes a separate $0.6 billion Contested Logistics line item, justified because adversaries will "
        "target supply lines, ports, and communications. Thirty-seven ships are listed as Forward Deployed Naval "
        "Forces (FDNF).")

    _h2(doc, "2.2 Navy exercise events — the Ship Wartime Repair and Maintenance Exercise at Cebu")

    _body(doc,
        "The Ship Repair Facility–Japan Regional Maintenance Center (SRF-JRMC), U.S. Seventh Fleet, U.S. "
        "Indo-Pacific Command (USINDOPACOM), and U.S. Pacific Fleet (COMPACFLT) all published parallel releases "
        "announcing that USS Ashland (Landing Ship Dock 48) completed Ship Wartime Repair and Maintenance in Cebu, "
        "Philippines (SWARMEX-Cebu) in coordination with Philippine Navy partners and local Philippine contractors. "
        "The exercise involved an amphibious warfare ship operating at a foreign port, host-nation contractor "
        "coordination for repair work, and Philippine Naval Sea Systems Command participation. The exercise "
        "structure aligns with three of the six operational-decision moments the proposed product addresses: "
        "forward team mobilization, foreign-port emergency contracting, and host-nation legal framework "
        "navigation.")

    _h2(doc, "2.3 Policy-community pressure")

    _body(doc,
        "The U.S. Naval Institute (USNI) Proceedings published \"Fix the Navy's Expeditionary Repair,\" arguing "
        "current expeditionary repair capability is inadequate to support fleet operations under wartime "
        "contestation. The Center for International Maritime Security (CIMSEC) published \"If the U.S. Navy Can't "
        "Repair Ships in Peacetime, How Will It Do So in War?\" — arguing current peacetime repair throughput is "
        "the floor and wartime throughput needs to be higher. Both publications reach the customer organizations.")

    _h2(doc, "2.4 Three distinct procurement precedents")

    _body(doc,
        "Earlier drafts of this brief lumped three federal contracts together as evidence the DoD buys exercise "
        "design at scale. That was imprecise. The three contracts split into three distinct work-types, and each "
        "tells the reader something different.")

    _body(doc,
        "**CACI NSS, LLC at AFRICOM.** $194 million task order for broad professional services to AFRICOM with "
        "exercise, training, and operations support as named scope. Department of Defense / Department of the "
        "Army funder. Awarded by the GSA Federal Acquisition Service against parent contract GS00Q14OADU121, "
        "an OASIS-equivalent governmentwide vehicle. This is the closest analog to the proposed product's "
        "work-type. It is also the strongest single piece of evidence that CACI has executed the same kind of "
        "work at multi-hundred-million-dollar scale at a sister Combatant Command.")

    _body(doc,
        "**Parsons Government Services Inc. at the Office of the Secretary of Defense (OSD).** $556.8 million "
        "Command, Control, Communications, Computers, Combat Systems, Intelligence, Surveillance, Reconnaissance, "
        "and Targeting / Exercises / Operations / Information Services (CEOIS) task order, providing near "
        "real-time situational awareness and decision support to the DoD Combatant Commands. This is decision "
        "support for active operations, not exercise design. Software and contractor staff embedded at the "
        "Combatant Commands. A different work-type than the proposed product, but adjacent.")

    _body(doc,
        "**Axient LLC at the Missile Defense Agency (MDA).** $233.8 million Test, Exercise and Wargames Support "
        "contract. Exercise and wargames work for missile defense testing and operator training. Same work-type "
        "as the proposed product, but in a different domain. Closest single precedent for a DoD agency procuring "
        "a large exercise-and-wargames contract through one named vehicle.")

    _body(doc,
        "What the three together suggest. The DoD procures the proposed-product work-type at scale via CACI NSS, "
        "LLC at AFRICOM and via Axient at MDA. The DoD procures the adjacent decision-support work-type at scale "
        "via Parsons at OSD. The Navy could plausibly procure either or both work-types. As of this draft date, "
        "the Navy has not done so under a named vehicle visible in the public record. A targeted source-finding "
        "pass for Navy fleet-command exercise-planning contracts was initiated 2026-05-26; results will inform "
        "later brief versions.")

    _pb(doc)

    # 3. Customer landscape
    _h1(doc, "3. Customer landscape")

    _body(doc,
        "The candidate customer set for the proposed product is layered. Strategic direction sits with the "
        "Secretary of the Navy (SECNAV) and CNO. They set priorities and authorize budget but do not procure the "
        "product directly. The repair-activity layer is the four public naval shipyards: Norfolk Naval Shipyard "
        "(NNSY), Puget Sound Naval Shipyard (PSNS), Pearl Harbor Naval Shipyard (PHNS), and Portsmouth Naval "
        "Shipyard (PNSY); plus the Regional Maintenance Centers (RMCs), including the forward-deployed SRF-JRMC; "
        "plus Commander, Navy Regional Maintenance Center (CNRMC) as the parent command. The doctrinal authority "
        "is Naval Warfare Development Center (NWDC). The fleet-command training-authority layer is Commander, U.S. "
        "Fleet Forces Command (FLTFORCOM); COMPACFLT; Commander, Naval Surface Forces (COMNAVSURFOR); and the U.S. "
        "Indo-Pacific Command Joint Operations Directorate (INDOPACOM J7).")

    _h2(doc, "3.1 Two audiences for the same scenario base")

    _body(doc,
        "The proposed product serves two audiences with the same scenario content. The repair-activity wardroom "
        "(the commanding officer of the shipyard, the executive officer of the RMC, the SRF-JRMC commander, plus "
        "department heads and visiting fleet liaisons) prepares to receive a damaged ship and mobilize against the "
        "work. The fleet commander (the tactical operator who owns the ship) decides where to send the damaged "
        "ship in the first place. Same scenario content. Different decision-maker focus.")

    _body(doc,
        "Port selection is the canonical dual-audience scenario. A destroyer takes shock damage 800 miles from "
        "three friendly ports. The fleet commander runs through the variables — proximity, repair capability, "
        "security posture, alliance access — and picks the port. The repair-activity wardroom at the chosen port "
        "then runs the same scenario from the receiving end — receive the ship, stage the repair team, hire local "
        "welders, handle customs. The product can run as a wardroom exercise at the receiving end (typically "
        "Level 3) or as a fleet-commander decision rehearsal at the sending end (Level 1 or Level 3 depending on "
        "the customer's preferred format).")

    _h2(doc, "3.2 Procurement model — end-user separate from contracting authority")

    _body(doc,
        "The operationally-most-exposed end-user (SRF-JRMC) is not necessarily the contracting authority. "
        "SRF-JRMC operates with Operations and Maintenance (O&M) dollars to fix forward-deployed ships. It does "
        "not typically possess the Research, Development, Test, and Evaluation (RDT&E) or Other Procurement, Navy "
        "(OPN) budget authority needed for a net-new software platform. The capture model in this brief decouples "
        "the two layers. SRF-JRMC sponsors the operational requirement. PACFLT N7 or N4, or FLTFORCOM N7 or N4, "
        "holds the contracting authority and the right budget color.")

    _h3(doc, "End-user / sponsor ranking")
    _bullet(doc, "SRF-JRMC wartime readiness cell — forward-deployed, audience and venue match the product.")
    _bullet(doc, "CNRMC plus subordinate RMCs plus Surface Team One (ST1) — Navy-wide repair-activity layer.")
    _bullet(doc, "Public naval shipyards (NNSY, PSNS, PHNS, PNSY) — heaviest infrastructure investment, longer procurement cycles.")

    _h3(doc, "Contracting-authority ranking — task order against an existing CACI vehicle")
    _bullet(doc, "PACFLT N7 / N4 or FLTFORCOM N7 / N4 — owns the fleet-exercise rhythm; could issue a task order against CACI NSS, LLC's existing OASIS or CIO-SP3 holdings.")
    _bullet(doc, "INDOPACOM J7 — owns joint multi-domain exercise procurement; same task-order path against the same vehicles.")
    _bullet(doc, "DIU, SBIR, or STTR — alternative entry if direct task-order pitch does not produce procurement inside the planning horizon.")
    _bullet(doc, "Naval Sea Systems Command Logistics, Maintenance, and Industrial Operations Directorate (NAVSEA 04), which sits above PAE Industrial Operations (PAE-IO), or Program Executive Office Integrated Warfare Systems (PEO IWS) — Navy Systems Command (SYSCOM) level authority; longer cycle but applicable if the product enters as part of a program of record.")

    _pb(doc)

    # 4. Competitive landscape
    _h1(doc, "4. Competitive landscape")

    _body(doc,
        "The competitive picture has four layers. Cross-customer exercise, wargames, and decision-support "
        "incumbents at federal scale. Navy training-systems incumbents adjacent but not directly competitive. "
        "Fleet-command exercise-planning incumbents under-mapped in public sources but the most-likely direct-"
        "competition layer. Navy professional-services incumbents at the proposed-product customer.")

    _h2(doc, "4.1 Three distinct adjacent-work-type incumbents")

    _body(doc,
        "Per section 2.4 above, the three federal-contract precedents are three distinct work-types. CACI itself "
        "is the dominant incumbent in the closest analog via CACI NSS, LLC at AFRICOM. Parsons holds the OSD "
        "decision-support work. Axient holds the MDA missile-defense wargames work. None of the three are direct "
        "competitors for the proposed product at the Navy repair-activity level. CACI NSS, LLC is the same-vendor "
        "capability-lineage anchor. Parsons and Axient are procurement-precedents only.")

    _h2(doc, "4.2 Navy training-systems incumbents — adjacent scope, not direct competition")

    _body(doc,
        "On 31 March 2026 the Navy modified the Naval Air Warfare Center Training Systems Division (NAWCTSD) "
        "training-systems support Indefinite Delivery / Indefinite Quantity (IDIQ) contract to add $1.2 billion "
        "in ceiling, raising total contract value to $2.51 billion. Nine prime contractors compete for task "
        "orders: BGI–Aero Simulation Joint Venture (JV), CAE USA, Delaware Resource Group of Oklahoma, "
        "Engineering Support Personnel, Fidelity Technologies, FlightSafety Defense, LB&B Associates, LTSS JV, "
        "and Valiant Global Defense Services. The work is simulator sustainment and instructional services, not "
        "the proposed-product work-type. The nine primes have Navy training relationships, contracting vehicles, "
        "and installation access, but are adjacent incumbents rather than direct competitors at the proposed-"
        "product level.")

    _h2(doc, "4.3 Fleet-command exercise-planning incumbents — under-mapped layer")

    _body(doc,
        "The biggest research gap in this section. The fleet-command exercise authorities (FLTFORCOM N7, PACFLT "
        "N7, INDOPACOM J7, and SRF-JRMC's wartime-readiness cell) procure exercise-planning support for Composite "
        "Training Unit Exercise (COMPTUEX), the Rim of the Pacific Exercise (RIMPAC), the Navy Large-Scale "
        "Exercise (LSE) series, the SWARMEX program, and unit-level workups. The specific contractor incumbents "
        "are not yet source-grounded in the ingested public material. The typical Master Scenario Events List "
        "(MSEL) prime base under similar vehicles at adjacent Combatant Commands likely includes Booz Allen "
        "Hamilton, Science Applications International Corporation (SAIC), and Huntington Ingalls Industries "
        "Mission Technologies (HII Mission Technologies), plus specialized sub-tier firms. These are candidates "
        "pending source-grounding. A targeted source-finding pass on this question was initiated 2026-05-26; "
        "results will inform v0.5 of this brief.")

    _h2(doc, "4.4 Navy professional-services incumbents at CNRMC — right customer, wrong vehicle")

    _body(doc,
        "Invictus Associates LLC holds Navy delivery order N0016424F3006 against parent contract N0017819D7883 "
        "for Professional Support Services (PSS) covering Fleet Readiness Support for CNRMC, subordinate RMCs, "
        "and Surface Team One. Total obligation $19,384,385. Place of performance Norfolk, Virginia. Department "
        "of the Navy funder. The parent contract is SeaPort Next-Generation (SeaPort-NxG), the Navy's generic-"
        "services vehicle. PSS through SeaPort-NxG buys staff augmentation and administrative support, not "
        "wargaming or scenario design. The Invictus task order is evidence the right customer procures contractor "
        "support generally, but the vehicle category is the wrong color of money for the proposed product.")

    _pb(doc)

    # 5. Our fit
    _h1(doc, "5. Our fit — case for CACI entry")

    _body(doc,
        "The case for CACI entry rests on three capability pillars and one procurement-infrastructure asset. "
        "The pillars: existing federal contract footprint in the work-type at scale (section 5.1); applied "
        "capability lineage from CACI's INDOPACOM J7 multi-domain exercise-design work (section 5.2); the ARKA "
        "acquisition signature libraries for threat-environment realism (section 5.3). The infrastructure asset: "
        "CACI NSS, LLC already holds the GSA-side governmentwide vehicles a Navy fleet command could issue task "
        "orders against (section 5.4).")

    _h2(doc, "5.1 CACI NSS, LLC at AFRICOM")

    _body(doc,
        "CACI NSS, LLC holds task order 47QFCA20F0042 against parent contract GS00Q14OADU121 for \"Plans, "
        "Operations, Logistics, Engagement, Training, Exercise, and Assessment Support to AFRICOM.\" Total "
        "obligation $194,034,792. Army funder via the GSA Federal Acquisition Service. The task order scope "
        "covers the full work-type stack named in the proposed-product description.")

    _body(doc,
        "What CACI NSS, LLC actually is. CACI NSS, LLC is a CACI subsidiary that holds multiple federal "
        "contracting vehicles. It is a general-purpose CACI services subsidiary, not a Navy-specific subsidiary. "
        "It provides program management, systems engineering, scientific services, logistics services, "
        "information technology solutions, and custom software development across DoD and the federal civil "
        "agencies. The \"NSS\" designation is historically tied to a National Security Solutions line.")

    _h2(doc, "5.2 PMTEC INDOPACOM J7 capability lineage")

    _body(doc,
        "CACI participates in the USINDOPACOM J7 Pacific Multi-Domain Training and Experimentation Capability "
        "(PMTEC) ecosystem, tracked separately in the PMTEC-USINDOPACOM opportunity in this vault. The "
        "capabilities CACI demonstrates under PMTEC are Live-Virtual-Constructive (LVC) integration, scenario "
        "design for multi-domain decision-making, AI and digital-twin technologies, and Regional Joint Training "
        "Infrastructure (RJTI). These transfer down-vertical to the naval-repair domain. The proposed product is "
        "a focused sub-application of the broader multi-domain exercise-design capability CACI demonstrates at "
        "PMTEC.")

    _h2(doc, "5.3 ARKA — threat-environment realism. Hull, Mechanical, and Electrical bridge needed separately.")

    _body(doc,
        "The March 2026 ARKA acquisition gives CACI electro-optical (EO), infrared (IR), and hyperspectral sensor "
        "signature libraries from intelligence, surveillance, and reconnaissance (ISR) work. For scenario "
        "realism, ARKA addresses the threat side: what a damaged ship's sensors would report, what threat "
        "indicators the team would see in the moments after a hit, what the operational picture looks like with "
        "degraded sensors. ARKA does not address repair-side realism: pump-out status, bulkhead compromise, "
        "drydock blocking, supply-chain routing for replacement valves.")

    _body(doc,
        "Repair-side realism requires a bridge to Navy Hull, Mechanical, and Electrical (HM&E) engineering and "
        "logistics data. There are three candidate closures for the bridge: partner with a prime that holds HM&E "
        "data; leverage CACI's existing naval IT footprint; or negotiate Government Furnished Information (GFI) "
        "from a Navy data-holder. The most likely GFI source is Naval Surface Warfare Center Carderock Division "
        "(NSWC Carderock), with possible additional sources at another Naval Surface Warfare Center or somewhere "
        "in PAE Industrial Operations (PAE-IO). The GFI path carries a 12-to-18-month Authority to Operate (ATO) "
        "schedule risk for a gamified non-system-of-record environment. The bridge decision should be made early "
        "so the executive pitch can offer a concrete data plan rather than a placeholder.")

    _h2(doc, "5.4 CACI NSS, LLC vehicle holdings — the procurement-infrastructure asset")

    _body(doc,
        "CACI NSS, LLC already holds three governmentwide acquisition vehicles open to any federal customer "
        "including any Navy command:")

    _bullet(doc, "GSA One Acquisition Solution for Integrated Services (OASIS) — Federal-wide Professional Services contract.")
    _bullet(doc, "NITAAC Chief Information Officer–Solutions and Partners 3 (CIO-SP3) — Information Technology services contract.")
    _bullet(doc, "GSA Multiple Award Schedule, contract number GS-35F-349CA.")

    _body(doc,
        "These are governmentwide acquisition contracts open to any federal customer. A Navy command could issue "
        "a task order against any of them without standing up a new Navy-specific contract. The AFRICOM "
        "$194 million task order is the precedent — Army funded, GSA awarded, CACI executes against an "
        "OASIS-equivalent vehicle CACI already held. The Navy fleet-command equivalent would be PACFLT N7 funds, "
        "GSA awards, CACI executes against the same vehicle.")

    _body(doc,
        "Implication for the pitch. CACI executives engaging PACFLT N7, FLTFORCOM N7, INDOPACOM J7, or SRF-JRMC "
        "leadership have a defined starting point. The work-type is named. The AFRICOM precedent is published. "
        "The vehicles are already held. The customer can issue a Level 1 pilot task order without committing to "
        "new procurement infrastructure.")

    _pb(doc)

    # 6. Working hypothesis
    _h1(doc, "6. Working hypothesis and falsifying legs")

    _body(doc,
        "Working hypothesis. The Navy has demand for contractor-supplied operational-decision-scenario content "
        "for BDAR teams and the fleet commanders they support. The proposed product is the five-level "
        "progression. CACI's case for entry rests on the four assets in section 5. The most-likely procurement "
        "pathway is a task order against CACI NSS, LLC's existing OASIS, CIO-SP3, or MAS holdings, sponsored "
        "operationally by SRF-JRMC and contracted by PACFLT N7 or FLTFORCOM N7.")

    _h2(doc, "6.1 Six falsifying legs")

    _bullet(doc,
        "Leg 1 — Navy operational-decision-scenario demand gap. Killed by senior signals turning ceremonial AND "
        "existing fleet exercises already injecting equivalent content organically.")
    _bullet(doc,
        "Leg 2 — Sub-product moat. Killed by incumbents replicating the analog form factor at the same time "
        "durations, or by the AI scenario-generation capability maturing rapidly enough among incumbents that "
        "CACI's PMTEC head start does not translate to a moat by procurement time.")
    _bullet(doc,
        "Leg 3 — CACI capability transferability from PMTEC to naval repair. Killed by PMTEC's multi-domain "
        "Combatant Command staff focus being structurally different from naval-repair small-unit decision "
        "rehearsal in ways that prevent direct capability transfer.")
    _bullet(doc,
        "Leg 4 — ARKA signature-library differentiator viability. Killed by intellectual-property posture "
        "blocking release, OR comparable data available to incumbents, OR the HM&E-data bridge cannot be closed, "
        "OR the GFI / ATO friction introduces a 12-to-18-month schedule delay that neutralizes the head start.")
    _bullet(doc,
        "Leg 5 — Customer-access pathway feasibility. Killed by fleet-command incumbents locking out CACI entry "
        "timeline, or by contractual right-of-first-refusal language preventing new-entrant access.")
    _bullet(doc,
        "Leg 6 — Sub-product viability as a fundable line. Killed by the sub-product being too small for "
        "standalone procurement (so it can only ride embedded in multi-day exercise contracts, which would mean "
        "partnering rather than priming), OR too specialized to attract a procurement vehicle at all, in which "
        "case the entry path is DIU, SBIR, or STTR pilot first.")

    _pb(doc)

    # 7. Recommendation
    _h1(doc, "7. Recommendations for executive action")

    _body(doc,
        "Four recommendations for CACI executive consideration.")

    _h2(doc, "7.1 Pursue the five-level progression at one initial customer")

    _body(doc,
        "Consider positioning CACI as a candidate prime for the five-level progression. Initial entry through "
        "Level 1 or Level 2 pilot. Operational sponsor: SRF-JRMC. Contracting authority: PACFLT N7 or N4, or "
        "FLTFORCOM N7 or N4. Contracting vehicle: CACI NSS, LLC's existing OASIS, CIO-SP3, or MAS holding, with "
        "no new Navy-specific vehicle required.")

    _h2(doc, "7.2 Initiate CACI executive engagement with the named Navy commands")

    _body(doc,
        "The Navy is not running an RFP for this work today. The path to procurement is likely to require "
        "proactive engagement with PACFLT N7, FLTFORCOM N7, INDOPACOM J7, NAVSEA 04 (the headquarters above "
        "PAE-IO), or SRF-JRMC leadership. The pitch could open with the work-type, the AFRICOM precedent, the "
        "GSA-wide vehicles CACI already holds, and a request for a Level 1 pilot task order against the existing "
        "OASIS or CIO-SP3 holding.")

    _h2(doc, "7.3 Decide the HM&E data bridge approach early")

    _body(doc,
        "Repair-side scenario realism needs Navy HM&E engineering and logistics data. ARKA covers threat-side "
        "only. Three candidate closures. Partner with a prime that holds HM&E data. Leverage CACI's existing "
        "naval IT footprint. Negotiate GFI from a Navy data-holder, most likely NSWC Carderock, with possible "
        "additional sources at another Naval Surface Warfare Center or somewhere in PAE-IO. The GFI path carries "
        "12-to-18-month ATO schedule risk. The decision is best made before pitch meetings begin.")

    _h2(doc, "7.4 Prepare an alternative entry path via DIU, SBIR, or STTR")

    _body(doc,
        "If direct executive engagement does not produce a task order inside the planning horizon, the "
        "alternative entry path is a DIU, SBIR, or STTR pilot for Level 1 or Level 2 scope. Lower per-engagement "
        "scale. Unblocks the no-existing-Navy-vehicle problem. Produces a reference customer for follow-on "
        "task-order competition. A parallel pilot strategy is worth having ready.")

    _pb(doc)

    # 8. Risks
    _h1(doc, "8. Risks and flagged concerns")

    _h2(doc, "8.1 No existing Navy task order")
    _body(doc,
        "The biggest risk. The Navy has not procured this sub-product as of this draft date. The OASIS / CIO-SP3 "
        "vehicle-holdings finding reduces this risk because no new contracting vehicle would be required, but the "
        "Navy still has to decide to issue the task order. If executive engagement fails to produce one, entry "
        "would collapse to the DIU, SBIR, or STTR pilot path, or to a subcontract on someone else's vehicle.")

    _h2(doc, "8.2 Incumbent lockout")
    _body(doc,
        "The fleet-command exercise-planning incumbent base is real but under-mapped in the public sources "
        "ingested for this research. Candidate names (Booz Allen Hamilton, SAIC, HII Mission Technologies, "
        "sub-tier firms) are not yet source-grounded. The actual subset active at FLTFORCOM N7, PACFLT N7, "
        "INDOPACOM J7, and SRF-JRMC could be crowded enough to require partnering, or sparse enough to allow "
        "first-mover prime entry. A targeted source-finding pass on this question was initiated 2026-05-26; "
        "results will inform v0.5 of this brief.")

    _h2(doc, "8.3 HM&E-data bridge schedule risk")
    _body(doc,
        "If the chosen bridge approach is the GFI path, securing the ATO for a gamified non-system-of-record "
        "environment can take 12 to 18 months. Engagement sequencing should account for that lead time. "
        "Mitigations include partnering with a prime that already holds HM&E data, leveraging CACI's existing "
        "naval IT footprint, or starting the ATO process in parallel with the engagement work.")

    _h2(doc, "8.4 Federal Acquisition Regulation Subpart 9.5 — Organizational Conflict of Interest")

    _body(doc,
        "Flagged here for executive review. The analytical direction of this research is informed in part by "
        "non-public knowledge of contractor personnel currently working at SRF-JRMC's wartime readiness group. "
        "This non-public knowledge has not entered any vault content as a citable claim. Its existence, however, "
        "may trigger an Organizational Conflict of Interest (OCI) analysis under Federal Acquisition Regulation "
        "(FAR) Subpart 9.5 if CACI advances this opportunity. CACI legal and contracts should review the OCI "
        "question before any pursuit decision or executive engagement with the named commands.")

    _h2(doc, "8.5 Source-grounding caveat")
    _body(doc,
        "Public sources support the hypothesis that the Navy has a wartime-repair problem and that the rest of "
        "DoD buys related work-types from contractors at scale. Public sources do not yet demonstrate a Navy "
        "procurement line for the specific proposed product. The case for entry rests on the adjacent-precedent "
        "argument plus CACI's existing capability lineage and vehicle holdings, not on a confirmed Navy demand "
        "signal.")

    _pb(doc)

    # 9. BD team next actions
    _h1(doc, "9. Business Development team next actions")

    _body(doc,
        "Operating tasks the BD team is ready to execute on green-light from CACI leadership. These do not "
        "require individual executive authorization — they are the workflow that follows from a decision to "
        "advance the recommendations in section 7.")

    _bullet(doc, "BD team gate review of this brief at v0.4 to confirm the proposed-scope frame.")
    _bullet(doc, "Run the targeted source-finding pass on fleet-command-exercise-planning incumbents at PACFLT N7, FLTFORCOM N7, INDOPACOM J7, and SRF-JRMC (initiated 2026-05-26). Results will be reflected in brief v0.5.")
    _bullet(doc, "Build the engagement-surface inventory targeting the specific contracting offices and active solicitations at the named commands, plus the SYSCOM authorities (NAVSEA 04 and PEO IWS).")
    _bullet(doc, "Develop the Level 1 pilot scenario portfolio — initial scenario cards for the port-selection, foreign-port emergency contracting, and forward team mobilization decision moments — ready for executive pitch meetings.")
    _bullet(doc, "Hold pursuit of any procurement action until CACI executive leadership commits to recommendations 7.1, 7.2, and 7.3, and the FAR Subpart 9.5 OCI consideration in section 8.4 is resolved by CACI leadership.")

    _pb(doc)

    # 10. Source ledger reference
    _h1(doc, "10. Source ledger reference")

    _body(doc,
        "Load-bearing factual claims trace to primary or quasi-primary sources in the opportunity's `01_sources/` "
        "directory. The full source ledger is at section 8 of `00_research-file.md` in the BDR-FLEET-READINESS "
        "opportunity folder. Key sources used in this brief:")

    _bullet(doc, "Department of the Navy FY 2027 President's Budget Press Brief, 28 April 2026.")
    _bullet(doc, "CNO Caudle Statement on the Posture of the United States Navy, HASC, 14 May 2026.")
    _bullet(doc, "USS Ashland SWARMEX-Cebu primary press releases — Navy.mil, Commander, U.S. Seventh Fleet (C7F), USINDOPACOM, COMPACFLT, NavalNews, NavalToday.")
    _bullet(doc, "USNI Proceedings, \"Fix the Navy's Expeditionary Repair.\"")
    _bullet(doc, "CIMSEC, \"If the U.S. Navy Can't Repair Ships in Peacetime, How Will It Do So in War?\"")
    _bullet(doc, "FY27 DoD Comptroller Justification Book — MAGTF Tactical Warfare Simulation, Joint Exercise Design (JExD) workforce, AI scenario generation funding lines.")
    _bullet(doc, "USAspending.gov award records: CACI NSS, LLC $194M AFRICOM; Parsons Government Services $556.8M CEOIS; Axient LLC $233.8M MDA wargames; Invictus Associates $19.4M CNRMC PSS.")
    _bullet(doc, "GovConWire article — Navy nine-contractors $1.2B NAWCTSD training-systems IDIQ modification, 31 March 2026.")
    _bullet(doc, "Norfolk Naval Shipyard $442M drydock modernization announcement (workboat.com).")
    _bullet(doc, "Pacific NAVFAC Shipyard Infrastructure Optimization Program (SIOP) Brief — PSNS and Subordinate Bases.")
    _bullet(doc, "CACI NSS, LLC vehicle holdings — GSA OASIS, NITAAC CIO-SP3, GSA MAS GS-35F-349CA (web-verified 2026-05-26).")

    _italic(doc,
        f"End of capture brief v0.4 draft. Build date: {date.today().isoformat()}. v0.4 changes from v0.3: tone "
        "tightened to non-confirmatory framing; every acronym expanded on first appearance; single-author voice "
        "throughout.")

    doc.save(str(OUT))
    print(f"✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    build()
