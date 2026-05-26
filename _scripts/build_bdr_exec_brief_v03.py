#!/usr/bin/env python3
"""
build_bdr_exec_brief_v03.py — Build BDR-FLEET-READINESS executive brief v0.3.

Updates from v0.2 per operator feedback 2026-05-26 (second round):
  1. Lead with the five-level progression, not 1-hour as the whole product.
  2. Differentiate Parsons / OSD / Axient — they are three different work-types, not one.
  3. Surface the OASIS / CIO-SP3 / GSA MAS finding. CACI NSS LLC already holds the vehicles
     the Navy can task-order against. Procurement path is sharper than v0.2 said.
  4. Make the dual-audience point explicit (repair-activity wardroom + fleet commander
     deciding where to send the damaged ship).
"""

from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

VAULT_ROOT = Path(__file__).parent.parent
OUT = VAULT_ROOT / "opportunities/BDR-FLEET-READINESS/04_artifacts/executive-brief-v0.3-draft.docx"


def _set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _set_tight_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)


def _body_compact(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def _bullet_compact(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"• {text}")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_default_font(doc)
    _set_tight_margins(doc)

    # Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run("BDR-FLEET-READINESS — Executive Brief v0.3 (Draft)")
    r.font.bold = True
    r.font.size = Pt(13)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run(f"Navy BDAR decision-rehearsal. Five-level progression. Task-order against CACI's existing GSA-wide vehicles. {date.today().isoformat()}. Internal exec review, OSI-only.")
    r.font.size = Pt(9)
    r.font.italic = True

    # BLUF
    _section_heading(doc, "Bottom Line Up Front")

    _body_compact(doc,
        "The Navy has a real wartime-repair problem. The rest of the Defense Department buys related contractor work "
        "at multi-hundred-million-dollar scale. The Navy should buy this work too. Today it has not issued a task "
        "order for it. That gap is the opportunity.")

    _body_compact(doc,
        "Recommendation. CACI should pursue a five-level progression of gamified decision-rehearsal scenarios, "
        "starting analog and moving to software at Level 4. The same scenario content serves two audiences: the "
        "wardroom at a Navy repair activity preparing to receive a damaged ship, AND the fleet commander deciding "
        "where to send her. Port selection is the canonical dual-audience scenario.")

    _body_compact(doc,
        "Procurement path is concrete. CACI NSS, LLC already holds GSA OASIS, NITAAC CIO-SP3, and GSA Multiple Award "
        "Schedule. The Navy can issue a task order against any of those vehicles. No new contracting infrastructure "
        "required. The AFRICOM $194M task order is the precedent — Army funded, GSA awarded, vehicle CACI already "
        "held. The Navy equivalent is the same shape with a Navy funder. CACI executives have to pitch the concept "
        "directly to PACFLT N7, FLTFORCOM N7, INDOPACOM J7, or SRF-JRMC leadership. The Navy is not running an RFP "
        "for this today.")

    # Why It Matters
    _section_heading(doc, "Why It Matters")

    _bullet_compact(doc,
        "Senior leadership demand signal. CNO Caudle named maintenance a warfighting requirement at HASC on 14 May "
        "2026. The FY27 budget funds $17B for ship maintenance toward an 80% Combat Surge Ready posture. SRF-JRMC "
        "ran a wartime-repair exercise with the Philippine Navy at Cebu.")

    _bullet_compact(doc,
        "DoD buys related contractor work-types at scale. AFRICOM via CACI NSS holds $194M for broad professional "
        "services with exercise scope. MDA via Axient holds $234M for missile-defense exercise and wargames. OSD via "
        "Parsons holds $557M for real-time decision support to Combatant Commanders. The Navy could plausibly "
        "procure any of these. Today it has not.")

    _bullet_compact(doc,
        "Navy professional community sees the gap. USNI Proceedings published \"Fix the Navy's Expeditionary "
        "Repair.\" CIMSEC published \"If the U.S. Navy Can't Repair Ships in Peacetime, How Will It Do So in War?\" "
        "Both reach the customer organizations.")

    _bullet_compact(doc,
        "Source-grounding caveat. The sources prove the Navy has the problem and the rest of DoD buys related work "
        "at scale. The sources do NOT yet prove the Navy has a named procurement line for this sub-product today. "
        "The hypothesis is solid. The salesmanship step is real BD work.")

    # The Progression
    _section_heading(doc, "The Five-Level Progression")

    _body_compact(doc,
        "Start analog. Move to software when scale justifies it. Each level can be procured separately. Each level "
        "builds a reference for the next.")

    _bullet_compact(doc,
        "Level 1. Flash drills. 5 to 15 minutes. One decision on a printed scenario card. Pattern recognition.")
    _bullet_compact(doc,
        "Level 2. Linked-decision sequences. 30 to 45 minutes. Three to five decisions in a row, each shaping the next.")
    _bullet_compact(doc,
        "Level 3. Full 1-hour gamified wardroom session. Multi-discipline audience. Turn-based. All six operational-"
        "decision moments integrated.")
    _bullet_compact(doc,
        "Level 4. Multi-session campaign. 2 to 4 hours across a deployment cycle. AI scenario generation pays for "
        "itself starting at this level.")
    _bullet_compact(doc,
        "Level 5. Live exercise injection into COMPTUEX, RIMPAC, Large-Scale Exercise, SWARMEX. Software-driven, "
        "federated into Navy synthetic training architecture.")

    # Procurement Path
    _section_heading(doc, "Procurement Path — Task Order Against Existing CACI Vehicles")

    _body_compact(doc,
        "This is the sharpest finding from the 2026-05-26 corpus update. CACI NSS, LLC already holds three "
        "governmentwide acquisition vehicles open to any federal customer including any Navy command. The Navy can "
        "issue a task order against any of these vehicles without standing up a new Navy-specific contract.")

    _bullet_compact(doc, "GSA OASIS — Federal-wide Professional Services contract.")
    _bullet_compact(doc, "NITAAC CIO-SP3 — Information Technology Acquisition and Assessment Center IT services contract.")
    _bullet_compact(doc, "GSA Multiple Award Schedule (GS-35F-349CA).")

    _body_compact(doc,
        "The pitch to PACFLT N7 or Fleet Forces N7 is concrete on day one. Here is the work-type. Here is the "
        "AFRICOM precedent. Here are the vehicles we already hold. Issue a Level 1 pilot task order. The customer "
        "commits no new procurement infrastructure.")

    # Recommendation
    _section_heading(doc, "Recommendation")

    _bullet_compact(doc,
        "Pursue the five-level progression. Start at SRF-JRMC if access permits. Otherwise PACFLT or Fleet Forces N7.")

    _bullet_compact(doc,
        "CACI executives pitch the concept directly. Path runs through proactive engagement with PACFLT N7, FLTFORCOM "
        "N7, INDOPACOM J7, NAVSEA 04, or SRF-JRMC leadership. Pitch is a task order against existing OASIS or CIO-SP3 "
        "holdings, not a request for a new vehicle.")

    _bullet_compact(doc,
        "Alternative entry if direct pitching does not produce a task order inside the planning horizon. DIU, SBIR, "
        "or STTR pilot for Level 1 or Level 2.")

    _bullet_compact(doc,
        "Close the FAR 9.5 OCI question before pitch meetings. Decide the HM&E-data bridge approach (partner, "
        "internal naval IT, or GFI / ATO). ARKA covers threat-side only.")

    # Top Risks
    _section_heading(doc, "Top Risks")

    _bullet_compact(doc,
        "No existing Navy task order. The OASIS / CIO-SP3 finding lowers this risk but does not eliminate it. The "
        "Navy still has to decide to issue the task order. If salesmanship fails, entry collapses to DIU / SBIR or "
        "subcontract path.")

    _bullet_compact(doc,
        "Incumbent lockout at the fleet-command exercise-planning layer. Candidate names (Booz Allen, SAIC, HII "
        "Mission Technologies, sub-tier players) not yet source-grounded. The next find_sources pass is targeting "
        "this gap.")

    _bullet_compact(doc,
        "FAR 9.5 OCI given operator-side knowledge of contractor exercise planners at SRF-JRMC. Owner: operator. "
        "Outcome may change the recommendation shape.")

    # Asks
    _section_heading(doc, "Asks")

    _bullet_compact(doc,
        "Authorize next find_sources pass to map fleet-command exercise-planning incumbents at PACFLT N7, FLTFORCOM "
        "N7, INDOPACOM J7, and SRF-JRMC.")

    _bullet_compact(doc,
        "Decide the HM&E-data bridge approach. Partner, internal naval IT, or GFI / ATO.")

    _bullet_compact(doc,
        "Close the FAR 9.5 OCI analysis.")

    _bullet_compact(doc,
        "Authorize CACI executive engagement with Navy leadership at the named commands. Pitch is task-order against "
        "OASIS or CIO-SP3, not new-vehicle creation.")

    # Footer
    f = doc.add_paragraph()
    f.paragraph_format.space_before = Pt(5)
    r = f.add_run(
        f"v0.3 built {date.today().isoformat()} from sealed §§1, 3, 4, 5, 6, 7 and §11.3 five-level progression of "
        f"00_research-file.md. Updates from v0.2: five-level progression made primary; three federal-contract "
        f"precedents differentiated by work-type; OASIS / CIO-SP3 / GSA MAS task-order procurement path surfaced; "
        f"dual-audience scenario design made explicit. Capture brief v0.2 is the long-form companion."
    )
    r.font.size = Pt(8)
    r.font.italic = True

    doc.save(str(OUT))
    print(f"✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    build()
