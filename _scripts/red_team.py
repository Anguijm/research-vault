#!/usr/bin/env python3
"""
red_team.py — Cross-AI red-team review of a capture or exec brief.

Pulls the operator's `_meta/prompts/cross-ai-red-team.md` prompt body, extracts
the brief text from the .docx, attaches FACT-level supporting quotes from the
latest verification report, sends to Gemini, and writes a red-team report into
the opportunity folder. Per SOP §2.1 rule 5 (adversarial review before
shipping) and the prompt note: red-team with the OTHER model from drafting.

Usage:
    python red_team.py --opportunity PMTEC-USINDOPACOM
    python red_team.py --opportunity PMTEC-USINDOPACOM --brief executive
    python red_team.py --opportunity PMTEC-USINDOPACOM --model gemini-2.5-pro
    python red_team.py --opportunity PMTEC-USINDOPACOM --no-quotes
"""

import argparse
import os
import re
import sys
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv

SCRIPT_DIR = Path(__file__).parent
VAULT_ROOT = SCRIPT_DIR.parent
OPP_ROOT = VAULT_ROOT / "opportunities"
PROMPT_FILE = VAULT_ROOT / "_meta" / "prompts" / "cross-ai-red-team.md"

sys.path.insert(0, str(SCRIPT_DIR))
load_dotenv(SCRIPT_DIR / ".env")

DEFAULT_MODEL = "gemini-3.1-pro-preview"
MAX_OUTPUT_TOKENS = 8000   # red-team responses are long — 3 passes × specific critiques


# ── Brief extraction ────────────────────────────────────────────────────

def _read_brief(path: Path) -> str:
    """Extract all paragraph text from a .docx, stopping at the auto-generated
    Sources sentinel so the red-team isn't graded on our own Sources appendix."""
    from docx import Document
    doc = Document(path)
    SENTINEL = "── Sources (auto-generated"
    lines = []
    for p in doc.paragraphs:
        if SENTINEL in p.text:
            break
        if p.text.strip():
            lines.append(p.text)
    # Tables too
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _find_brief(opp_dir: Path, kind: str) -> Path:
    artifacts = opp_dir / "04_artifacts"
    pattern = f"{kind}-brief-*.docx"
    matches = sorted(artifacts.glob(pattern))
    if not matches:
        raise FileNotFoundError(f"No {kind} brief found in {artifacts}")
    # Latest version if multiple
    return matches[-1]


# ── Source-quote extraction (from latest verification report) ───────────

def _collect_supporting_quotes(opp_dir: Path) -> str:
    """Pull FACT-level supporting quotes from the latest _verification-*.md.

    Returns a markdown block with one FACT per entry: claim → supporting quote.
    Used to ground the red-team in what the brief's evidence actually says,
    so Gemini can flag cases where the brief overclaims relative to sources.
    """
    reports = sorted(opp_dir.glob("_verification-*.md"))
    if not reports:
        return "(No verification report found — red-team has no per-FACT quote ground truth.)"
    latest = reports[-1]
    text = latest.read_text(encoding="utf-8")

    out = [f"## FACT-level supporting quotes (from {latest.name})", ""]
    # Each FACT block: claim line + supporting_quote lines (markdown blockquotes)
    blocks = re.split(r'^### ', text, flags=re.MULTILINE)
    for b in blocks[1:]:
        header_m = re.match(r'(?P<icon>\S+)\s+FACT #(?P<n>\d+)\s+§\d+\s+—\s+\*\*(?P<status>\w+)\*\*', b)
        if not header_m:
            continue
        n = header_m.group("n")
        status = header_m.group("status")
        claim_m = re.search(r'\*\*Claim:\*\*\s*\n>\s*(.+?)(?:\n\n|\n\*\*)', b, re.DOTALL)
        claim = claim_m.group(1).strip().replace("\n> ", " ") if claim_m else "(no claim extracted)"
        # All supporting quotes in this FACT block
        quote_matches = re.findall(r'^\s*>\s*_"(.+?)"_\s*$', b, re.MULTILINE)
        out.append(f"### FACT #{n} ({status})")
        out.append(f"**Claim:** {claim[:400]}{'…' if len(claim) > 400 else ''}")
        if quote_matches:
            for q in quote_matches[:2]:   # cap at 2 quotes per FACT to control prompt size
                out.append(f"> {q[:600]}{'…' if len(q) > 600 else ''}")
        else:
            out.append("> _(no supporting quote in report — verifier could not anchor)_")
        out.append("")
    return "\n".join(out)


# ── Prompt assembly ─────────────────────────────────────────────────────

def _load_red_team_prompt() -> str:
    """Extract the prompt body from inside the ``` fence in the SOP prompt file."""
    text = PROMPT_FILE.read_text(encoding="utf-8")
    m = re.search(r'```\n(.+?)\n```', text, re.DOTALL)
    if not m:
        raise RuntimeError(f"Could not find fenced prompt body in {PROMPT_FILE}")
    return m.group(1).strip()


_OUTPUT_SCHEMA_INSTRUCTIONS = """
Run all three passes (Customer / Competitor / Skeptical Exec) per the SOP prompt
above. Be specific and adversarial. Then collapse your findings into the JSON
schema below — one entry per distinct issue. If the same problem surfaces in
two passes, write ONE entry and list both passes in the "passes" array.

Calibrate severity:
  - critical  → factual error, brief literally cannot ship until fixed
  - high      → exec or customer would reject the brief on this alone
  - medium    → significantly weakens the position; competitor will exploit
  - low       → editorial or polish; doesn't change the recommendation

Return ONLY valid JSON matching this schema (no prose before or after, no
markdown fencing):

{
  "summary": "2-3 sentence TL;DR — what's the headline issue with this brief?",
  "verdict": "ship-ready | iterate | rebuild",
  "findings": [
    {
      "id": 1,
      "severity": "critical | high | medium | low",
      "sections": ["§7.1", "§4.1"],
      "passes": ["customer", "competitor", "exec"],
      "title": "≤80 char headline of the issue",
      "issue": "1-2 sentence problem statement",
      "why_it_matters": "1 sentence on the impact / who would notice",
      "fix": "1-2 sentence concrete recommended action",
      "evidence_quote": "Optional exact quote from the brief (≤200 chars), empty string if none"
    }
  ]
}
"""


def _assemble_prompt(brief_text: str, source_quotes: str, include_quotes: bool) -> str:
    prompt_body = _load_red_team_prompt()
    parts = [
        prompt_body,
        "",
        "---",
        "",
        "## CAPTURE BRIEF (full text)",
        "",
        brief_text,
    ]
    if include_quotes:
        parts.extend(["", "---", "", source_quotes])
    parts.extend([
        "",
        "---",
        "",
        "If a claim in the brief looks supportable from the FACT-level "
        "supporting quotes above, do not flag it — focus your ammunition on "
        "places where the brief reaches beyond what the source actually says.",
        "",
        _OUTPUT_SCHEMA_INSTRUCTIONS.strip(),
    ])
    return "\n".join(parts)


# ── Gemini call ─────────────────────────────────────────────────────────

def _call_gemini(prompt: str, model: str) -> str:
    from google import genai
    from google.genai import types
    client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
    response = client.models.generate_content(
        model=model,
        contents=prompt,
        config=types.GenerateContentConfig(
            max_output_tokens=MAX_OUTPUT_TOKENS,
            temperature=0.4,
            response_mime_type="application/json",   # demand structured output
        ),
    )
    return response.text or ""


# ── JSON parsing + scannable markdown rendering ─────────────────────────

import json

_SEVERITY_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3}
_SEVERITY_BADGE = {
    "critical": "🔥 CRITICAL",
    "high":     "🔴 HIGH",
    "medium":   "🟡 MEDIUM",
    "low":      "🟢 LOW",
}
_PASS_LABEL = {
    "customer":   "J7 staffer",
    "competitor": "Competitor",
    "exec":       "Skeptical exec",
}


def _parse_response(text: str) -> dict | None:
    """Parse Gemini's JSON response. Returns None if unparseable."""
    if not text:
        return None
    s = text.strip()
    # Strip code fences if present (some models still add them)
    m = re.match(r'^```(?:json)?\s*(\{.*\})\s*```\s*$', s, re.DOTALL)
    if m:
        s = m.group(1)
    try:
        data = json.loads(s)
    except json.JSONDecodeError:
        # Last-ditch: find the outermost braces
        start, end = s.find("{"), s.rfind("}")
        if start != -1 and end > start:
            try:
                data = json.loads(s[start:end + 1])
            except json.JSONDecodeError:
                return None
        else:
            return None
    if not isinstance(data, dict) or "findings" not in data:
        return None
    return data


def _render_markdown(data: dict, model: str, brief_path: Path, opp_id: str) -> str:
    findings = sorted(
        data.get("findings", []),
        key=lambda f: (_SEVERITY_ORDER.get(str(f.get("severity", "low")).lower(), 99),
                       int(f.get("id", 999))),
    )
    counts = {s: sum(1 for f in findings if str(f.get("severity", "")).lower() == s)
              for s in _SEVERITY_ORDER}
    total = len(findings)

    summary = data.get("summary", "(no summary returned)")
    verdict = data.get("verdict", "iterate").lower()
    verdict_badge = {
        "ship-ready": "✅ SHIP-READY",
        "iterate":    "🔁 ITERATE",
        "rebuild":    "🛑 REBUILD",
    }.get(verdict, f"❓ {verdict.upper()}")

    lines = [
        f"# Red-Team Review — {opp_id}",
        "",
        f"**Subject:** `{brief_path.name}`  ·  **Reviewer:** `{model}`  ·  "
        f"**Generated:** {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## At a glance",
        "",
        f"| Verdict | Total | 🔥 Critical | 🔴 High | 🟡 Medium | 🟢 Low |",
        f"|---|---|---|---|---|---|",
        f"| **{verdict_badge}** | {total} | {counts['critical']} | "
        f"{counts['high']} | {counts['medium']} | {counts['low']} |",
        "",
        f"**SOP threshold:** more than 2 findings → brief not yet ready "
        f"(see `_meta/prompts/cross-ai-red-team.md`).",
        "",
        "## TL;DR",
        "",
        f"> {summary}",
        "",
        "---",
        "",
    ]

    # Group by severity
    by_sev: dict[str, list[dict]] = {s: [] for s in _SEVERITY_ORDER}
    for f in findings:
        sev = str(f.get("severity", "low")).lower()
        if sev not in by_sev:
            sev = "low"
        by_sev[sev].append(f)

    for sev in _SEVERITY_ORDER:
        items = by_sev[sev]
        if not items:
            continue
        lines.append(f"## {_SEVERITY_BADGE[sev]}  ({len(items)})")
        lines.append("")
        for f in items:
            secs = " · ".join(f.get("sections", []) or ["?"])
            passes = " · ".join(_PASS_LABEL.get(p, p) for p in f.get("passes", []) or [])
            lines.append(f"### #{f.get('id','?')}  ·  {secs}  ·  {f.get('title','(no title)')}")
            lines.append("")
            lines.append(f"*Flagged by: {passes}*" if passes else "*Flagged by: (unspecified)*")
            lines.append("")
            issue = (f.get("issue") or "").strip()
            why = (f.get("why_it_matters") or "").strip()
            fix = (f.get("fix") or "").strip()
            quote = (f.get("evidence_quote") or "").strip()
            lines.append(f"**Issue.** {issue}")
            lines.append("")
            if why:
                lines.append(f"**Why it matters.** {why}")
                lines.append("")
            if fix:
                lines.append(f"**Fix.** {fix}")
                lines.append("")
            if quote:
                lines.append(f"> _“{quote}”_ — from the brief")
                lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def _render_fallback(raw: str, model: str, brief_path: Path, opp_id: str) -> str:
    """If JSON parsing failed, dump the raw response with a warning header."""
    return (
        f"# Red-Team Review — {opp_id}\n\n"
        f"**Subject:** `{brief_path.name}`  ·  **Reviewer:** `{model}`  ·  "
        f"**Generated:** {datetime.now().isoformat(timespec='seconds')}\n\n"
        f"> ⚠ The model did not return valid JSON. Raw output below.\n\n"
        f"---\n\n{raw}\n"
    )


# ── Report writing ──────────────────────────────────────────────────────

def _extract_brief_version(brief_path: Path) -> str:
    """Pull 'v0.3.1' or 'v0.3' or similar out of a brief filename
    like 'capture-brief-v0.3.1-draft.docx'. Greedy on the dotted-version part."""
    m = re.search(r'(v\d+(?:\.\d+)+(?:[a-z]+)?)', brief_path.stem, re.IGNORECASE)
    return m.group(1).lower() if m else "vX"


def _write_report(opp_dir: Path, model: str, brief_path: Path, response: str) -> tuple[Path, dict | None]:
    """Write both the rendered markdown report AND the raw JSON.

    Returns (markdown_path, parsed_data_or_None).
    """
    date = datetime.now().strftime("%Y-%m-%d")
    safe_model = model.replace("/", "_").replace(":", "_")
    version = _extract_brief_version(brief_path)
    md_path = opp_dir / f"_red-team-{date}-{safe_model}-{version}.md"
    json_path = opp_dir / f"_red-team-{date}-{safe_model}-{version}.json"

    # Persist raw JSON regardless of parse success
    json_path.write_text(response, encoding="utf-8")

    data = _parse_response(response)
    frontmatter = (
        f"---\n"
        f"type: red_team_report\n"
        f"opportunity: {opp_dir.name}\n"
        f"reviewed: {brief_path.name}\n"
        f"model: {model}\n"
        f"date: {date}\n"
        f"prompt_source: _meta/prompts/cross-ai-red-team.md\n"
        f"raw_json: {json_path.name}\n"
        f"---\n\n"
    )
    if data is None:
        body = _render_fallback(response, model, brief_path, opp_dir.name)
    else:
        body = _render_markdown(data, model, brief_path, opp_dir.name)
    md_path.write_text(frontmatter + body, encoding="utf-8")
    return md_path, data


def _append_decision_log(opp_dir: Path, model: str, brief_path: Path,
                        report_path: Path, finding_count: int | None) -> None:
    log = opp_dir / "05_decision-log.md"
    if not log.exists():
        return
    date = datetime.now().strftime("%Y-%m-%d")
    finding_str = f"{finding_count} substantive findings" if finding_count is not None else "findings TBD by operator"
    entry = (
        f"\n### {date} — Cross-AI red-team via {model}\n\n"
        f"**By:** operator + `_scripts/red_team.py`\n"
        f"**Rationale:** SOP §2.1 rule 5 adversarial review before shipping. "
        f"Drafting used Claude; red-team uses Gemini per prompt note.\n"
        f"**What changed:** Report written to `{report_path.name}`. "
        f"Reviewed `{brief_path.name}` ({finding_str}). "
        f"Operator decides whether to address findings before next gate.\n\n"
        f"---\n"
    )
    with open(log, "a", encoding="utf-8") as f:
        f.write(entry)


def _count_findings(data: dict | None) -> int | None:
    """Count findings from parsed JSON. Returns None if parse failed."""
    if data is None:
        return None
    findings = data.get("findings", []) or []
    return len(findings)


# ── Main ────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--opportunity", required=True, metavar="ID")
    parser.add_argument("--brief", default="capture",
        choices=["capture", "executive"],
        help="Which brief to red-team (default: capture)")
    parser.add_argument("--model", default=DEFAULT_MODEL,
        help=f"Gemini model ID (default: {DEFAULT_MODEL})")
    parser.add_argument("--no-quotes", action="store_true",
        help="Skip the FACT-level supporting quotes context (smaller prompt)")
    args = parser.parse_args()

    if not os.environ.get("GEMINI_API_KEY"):
        print("ERROR: GEMINI_API_KEY not set in .env")
        sys.exit(1)

    opp_dir = OPP_ROOT / args.opportunity
    if not opp_dir.is_dir():
        print(f"ERROR: Opportunity folder not found: {opp_dir}")
        sys.exit(1)

    brief_path = _find_brief(opp_dir, args.brief)
    print(f"Brief: {brief_path.relative_to(VAULT_ROOT)}")

    print(f"Extracting brief text...")
    brief_text = _read_brief(brief_path)
    print(f"  {len(brief_text):,} chars")

    include_quotes = not args.no_quotes
    quotes = _collect_supporting_quotes(opp_dir) if include_quotes else ""
    if include_quotes:
        print(f"Collected FACT-level supporting quotes: {len(quotes):,} chars")

    prompt = _assemble_prompt(brief_text, quotes, include_quotes)
    print(f"Total prompt: {len(prompt):,} chars  ·  model: {args.model}")
    print(f"Calling Gemini (this can take 30-60s for a 3-pass red-team)...")

    response = _call_gemini(prompt, args.model)
    print(f"  ↩ {len(response):,} chars of red-team output\n")

    report, data = _write_report(opp_dir, args.model, brief_path, response)
    rel = report.relative_to(VAULT_ROOT)

    findings = _count_findings(data)
    _append_decision_log(opp_dir, args.model, brief_path, report, findings)

    print(f"✓ Report written: {rel}")
    if data is None:
        print("  ⚠ JSON parsing failed — see raw output dumped in the report")
    elif findings is not None:
        counts = {s: sum(1 for f in data.get("findings", [])
                         if str(f.get("severity", "")).lower() == s)
                  for s in ("critical", "high", "medium", "low")}
        print(f"  Findings: {findings} total  ·  "
              f"🔥{counts['critical']}  🔴{counts['high']}  "
              f"🟡{counts['medium']}  🟢{counts['low']}")
        print(f"  Verdict: {data.get('verdict', '(none)')}")
    print(f"  Decision log updated: {(opp_dir / '05_decision-log.md').relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    main()
