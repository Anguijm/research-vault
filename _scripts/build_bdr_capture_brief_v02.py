#!/usr/bin/env python3
"""
build_bdr_capture_brief_v02.py — Build BDR-FLEET-READINESS capture brief v0.2.

Three corrections from v0.1, per operator feedback 2026-05-26:
  1. Lead with the five-level progression. Stop framing the 1-hour session as the whole product.
  2. Differentiate the three federal contracts. CACI NSS at AFRICOM is broad services with exercise
     scope. Parsons at OSD is decision-support for active commanders. Axient at MDA is missile-defense
     exercise/wargames. Lumping them as "exercise design at scale" was imprecise.
  3. Add the OASIS / CIO-SP3 / GSA MAS finding. CACI NSS, LLC already holds the governmentwide vehicles
     a Navy fleet command can task-order against. The Navy procurement path is "issue a task order
     against an existing CACI vehicle," not "find a new Navy vehicle."

Plus the dual-audience point made explicit: same scenario content serves the repair-activity wardroom
AND the fleet commander deciding where to send the damaged ship.
"""

from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

VAULT_ROOT = Path(__file__).parent.parent
OUT = VAULT_ROOT / "opportunities/BDR-FLEET-READINESS/04_artifacts/capture-brief-v0.2-draft.docx"


def _set_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(18)


def _add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(14)


def _add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(12)


def _add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def _add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(f"• {text}")


def _add_italic(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.italic = True


def _add_page_break(doc):
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_styles(doc)

    # Cover
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(36)
    r = cover.add_run("CAPTURE BRIEF — v0.2 DRAFT\n")
    r.font.bold = True
    r.font.size = Pt(22)

    r = cover.add_run("\nBDR-FLEET-READINESS\n")
    r.font.bold = True
    r.font.size = Pt(18)

    r = cover.add_run("\nFive-level analog-first progression of gamified decision-rehearsal scenarios for Navy BDAR teams and fleet commanders\n")
    r.font.size = Pt(13)
    r.font.italic = True

    r = cover.add_run(f"\nDraft date: {date.today().isoformat()}  |  Classification: For BD-team review (OSI-only).\n")
    r.font.size = Pt(10)

    r = cover.add_run("\nv0.2 revisions: (1) lead with five-level progression; (2) differentiate the three federal-contract precedents by work-type; (3) procurement path is task-order against existing CACI NSS LLC OASIS / CIO-SP3 / GSA MAS holdings; (4) dual-audience point made explicit.\n")
    r.font.size = Pt(9)
    r.font.italic = True

    _add_page_break(doc)

    # 1. BLUF
    _add_heading_1(doc, "1. Bottom Line Up Front")

    _add_body(doc,
        "The Navy has a real wartime-repair readiness problem. The rest of the Defense Department already buys "
        "exercise and decision-rehearsal content from contractors at multi-hundred-million-dollar scale. The Navy "
        "should buy this kind of work too. Today the Navy has not issued a task order for it. That gap is the "
        "opportunity, and the path to closing it is sharper than v0.1 implied.")

    _add_body(doc,
        "The product is a five-level progression of gamified decision-rehearsal scenarios. Level 1 is 5 to 15 minute "
        "flash drills on printed cards. Level 2 is 30 to 45 minute linked-decision sequences. Level 3 is the 1-hour "
        "full gamified wardroom session. Level 4 is a 2 to 4 hour multi-session campaign across a deployment cycle. "
        "Level 5 is software-driven scenario injection into live fleet exercises. Start analog. Move to software at "
        "Level 4 when the customer commits to scale. The same scenario content serves two audiences: the repair-"
        "activity wardroom at a shipyard, RMC, or SRF-JRMC, AND the fleet commander deciding where to send a damaged "
        "ship. Port selection is the canonical dual-audience scenario.")

    _add_body(doc,
        "The procurement path is concrete. CACI NSS, LLC already holds GSA OASIS, NITAAC CIO-SP3, and GSA Multiple "
        "Award Schedule. The AFRICOM $194M task order ($194 million for Plans, Operations, Logistics, Engagement, "
        "Training, Exercise, and Assessment Support) is the precedent. Army funded it. GSA awarded it. CACI executes "
        "it on a vehicle CACI already held. The Navy can do the same thing. Any Navy fleet command (PACFLT N7, "
        "FLTFORCOM N7, INDOPACOM J7) or SYSCOM authority (NAVSEA 04, PEO IWS) can issue a task order against CACI's "
        "existing OASIS or CIO-SP3 vehicles without setting up a new Navy-specific contract.")

    _add_body(doc,
        "Recommendation. CACI executives pitch the concept directly to PACFLT N7 or FLTFORCOM N7 or SRF-JRMC "
        "leadership. The pitch is concrete on day one. Here is the work-type. Here is the AFRICOM precedent. Here are "
        "the GSA-wide vehicles we already hold. Issue a Level 1 pilot task order. No new procurement infrastructure "
        "required. The salesmanship step is real BD work, not waiting for an RFP. The alternative entry path if the "
        "task-order pitch does not produce procurement inside the planning horizon is DIU, SBIR, or STTR pilot for "
        "Level 1 or Level 2 scope.")

    _add_page_break(doc)

    # 2. Demand signal
    _add_heading_1(doc, "2. Demand signal")

    _add_body(doc,
        "Four signals converge on the corrected-scope demand. Senior Navy leadership is publicly committing to "
        "wartime repair readiness. The Navy is publicly running the most directly relevant exercise type in the most "
        "directly relevant venue. The Navy professional community is publicly arguing the gap is real. And the "
        "Defense Department broadly procures contractor-delivered exercise design and decision-support content at "
        "multi-hundred-million-dollar scale.")

    _add_heading_2(doc, "2.1 Senior-leadership signal")

    _add_body(doc,
        "Admiral Daryl L. Caudle, the Chief of Naval Operations, delivered a Statement on the Posture of the United "
        "States Navy before the House Armed Services Committee on 14 May 2026, naming maintenance as a warfighting "
        "requirement and announcing a planned deliberate study of Navy yard capacity. The DON FY27 Press Brief funds "
        "Ship Maintenance at $17.0 billion to drive an 80% Combat Surge Ready posture, plus a $0.6B Contested "
        "Logistics line because adversaries will target supply lines, ports, and communications. 37 ships are listed "
        "as Forward Deployed Naval Forces.")

    _add_heading_2(doc, "2.2 Navy public exercise events — SWARMEX-Cebu")

    _add_body(doc,
        "SRF-JRMC, U.S. Seventh Fleet, U.S. Pacific Command, and U.S. Pacific Fleet published parallel releases "
        "announcing that USS Ashland (LSD 48) completed Ship Wartime Repair and Maintenance in Cebu, Philippines, "
        "with Philippine Navy partners and local Philippine contractors. The exercise involved an amphibious warfare "
        "ship operating at a foreign port, host-nation contractor coordination, and Philippine Naval Sea Systems "
        "Command participation. The exercise structure maps directly onto three of the six operational-decision "
        "moments: forward team mobilization, foreign-port emergency contracting, and host-nation framework navigation.")

    _add_heading_2(doc, "2.3 Policy-community pressure")

    _add_body(doc,
        "USNI Proceedings published \"Fix the Navy's Expeditionary Repair.\" CIMSEC published \"If the U.S. Navy "
        "Can't Repair Ships in Peacetime, How Will It Do So in War?\" Both publications argue current expeditionary "
        "repair capability is inadequate for wartime contestation. Both reach the customer organizations.")

    _add_heading_2(doc, "2.4 Three distinct procurement precedents — keep them separate")

    _add_body(doc,
        "v0.1 of this brief lumped three federal contracts together as evidence the Defense Department buys exercise "
        "design at scale. That was imprecise. The three contracts are three distinct work-types. Each one is a "
        "different procurement precedent for different reasons.")

    _add_body(doc,
        "**CACI NSS, LLC at AFRICOM.** $194 million task order for broad professional services to AFRICOM with "
        "exercise, training, and operations support as named scope. Army funder via GSA Federal Acquisition Service. "
        "Awarded against parent IDC GS00Q14OADU121, an OASIS-equivalent governmentwide vehicle. This is the closest "
        "analog to the corrected-scope work-type and the single most consequential capability-lineage anchor for "
        "CACI's right to win.")

    _add_body(doc,
        "**Parsons Government Services at OSD.** $556.8 million CEOIS task order. CEOIS expands to Command, Control, "
        "Communications, Computers, Combat Systems, Intelligence, Surveillance, Reconnaissance, and Targeting / "
        "Exercises / Operations / Information Services. The work is decision-support for active Combatant Command "
        "operations — software platforms and contractor staff embedded at INDOPACOM, EUCOM, AFRICOM, CENTCOM "
        "supporting the four-star and the J-staff in real time. This is NOT exercise design. This is decision-"
        "support for ongoing operations. It is a different work-type than the corrected-scope sub-product but "
        "adjacent.")

    _add_body(doc,
        "**Axient LLC at MDA.** $233.8 million Test, Exercise and Wargames Support contract for missile defense "
        "testing and operator training. MDA funder. Place of performance Huntsville, Alabama. This IS exercise and "
        "wargames work but for a different domain (missile defense, not naval repair). It is the closest precedent "
        "for a single DoD agency procuring a multi-hundred-million-dollar exercise-and-wargames contract in one "
        "place — relevant precedent for a Navy fleet command or SYSCOM authority doing similar work for the naval-"
        "repair domain.")

    _add_body(doc,
        "What the three together tell the BD team: the Defense Department definitely procures the corrected-scope "
        "work-type from contractors at scale (AFRICOM via CACI NSS, MDA via Axient). DoD also procures the adjacent "
        "decision-support work-type at scale (OSD via Parsons). The Navy could plausibly procure either or both. "
        "Today the Navy has not. The next find_sources pass is trying to identify whether the Navy is procuring "
        "anything similar through fleet-command vehicles we have not yet surfaced.")

    _add_page_break(doc)

    # 3. Customer landscape
    _add_heading_1(doc, "3. Customer landscape")

    _add_body(doc,
        "The corrected-scope customer landscape is layered. Strategic direction sits with SECNAV and CNO. They set "
        "priorities and authorize budget but do not procure the product directly. The repair-activity layer is the "
        "four public naval shipyards (Norfolk, Puget Sound, Pearl Harbor, Portsmouth), the Regional Maintenance "
        "Centers including the forward-deployed SRF-JRMC, and Commander, Navy Regional Maintenance Center (CNRMC) "
        "as parent command. The doctrinal authority is Naval Warfare Development Center (NWDC). The fleet-command "
        "training-authority layer is FLTFORCOM, COMPACFLT, COMNAVSURFOR, and INDOPACOM J7.")

    _add_heading_2(doc, "3.1 Two audiences for the same scenario base")

    _add_body(doc,
        "The corrected-scope product serves two audiences with the same scenario content. The repair-activity "
        "wardroom — Captain of the shipyard, Executive Officer of the RMC, SRF-JRMC commander, plus department "
        "heads and visiting fleet liaisons — prepares to receive a damaged ship and mobilize against the work. The "
        "fleet commander — the tactical operator who owns the ship — decides where to send the damaged ship in the "
        "first place. Same scenario content. Different decision-maker focus.")

    _add_body(doc,
        "Port-selection is the canonical dual-audience scenario. A destroyer takes shock damage 800 miles from three "
        "friendly ports. The fleet commander runs through the variables: proximity, repair capability, security "
        "posture, alliance access. He picks the port. The repair-activity wardroom at the chosen port then runs the "
        "same scenario from the receiving end: receive the ship, stage the repair team, hire local welders, handle "
        "customs. The product as designed can run as a wardroom exercise at the receiving end (Level 3) or as a "
        "fleet-commander decision rehearsal at the sending end (Level 1 flash drill or Level 3 full session "
        "depending on the customer).")

    _add_heading_2(doc, "3.2 Procurement model — end-user separate from contracting authority")

    _add_body(doc,
        "The operationally-most-exposed end-user (SRF-JRMC) is not necessarily the contracting authority. SRF-JRMC "
        "has Operations and Maintenance dollars to fix forward-deployed ships. It does not typically possess the "
        "RDT&E or Other Procurement, Navy budget authority needed for a net-new gamified software platform. The "
        "capture model decouples the two layers. SRF-JRMC sponsors the operational requirement. PACFLT N7 / N4 or "
        "FLTFORCOM N7 / N4 holds the contracting authority and the right budget color.")

    _add_heading_3(doc, "End-user / sponsor ranking")
    _add_bullet(doc, "SRF-JRMC wartime readiness cell — forward-deployed, audience and venue match the product.")
    _add_bullet(doc, "CNRMC plus subordinate RMCs plus Surface Team One — Navy-wide repair-activity layer.")
    _add_bullet(doc, "Public naval shipyards (NNSY, PSNS, PNSY, PHNS) — heaviest investment, longer cycles.")

    _add_heading_3(doc, "Contracting-authority ranking — task order against existing CACI vehicles")
    _add_bullet(doc, "PACFLT N7 / N4 or FLTFORCOM N7 / N4 — owns the fleet-exercise rhythm; can issue a task order against CACI NSS LLC's existing OASIS or CIO-SP3 holdings.")
    _add_bullet(doc, "INDOPACOM J7 — owns joint multi-domain exercise procurement; same task-order path.")
    _add_bullet(doc, "DIU, SBIR, or STTR — alternative entry if direct task-order pitch does not produce procurement inside the planning horizon.")
    _add_bullet(doc, "NAVSEA 04 or PEO IWS — SYSCOM-level authority for training-systems work; longer cycle but applicable if the product enters as part of a program of record.")

    _add_page_break(doc)

    # 4. Competitive landscape
    _add_heading_1(doc, "4. Competitive landscape")

    _add_body(doc,
        "Four competitive layers. Cross-customer Exercise / Wargames / Decision-Support incumbents at federal scale. "
        "Navy training-systems incumbents adjacent but not directly competitive. Fleet-command exercise-planning "
        "incumbents under-mapped in public sources but the most-likely direct-competition layer. Navy professional-"
        "services incumbents at the corrected-scope customer.")

    _add_heading_2(doc, "4.1 Three distinct adjacent-work-type incumbents")

    _add_body(doc,
        "Per §2.4 above, the three federal-contract precedents are three distinct work-types. CACI itself is the "
        "dominant incumbent in the closest analog (broad services with exercise scope via NSS LLC at AFRICOM). "
        "Parsons holds the OSD decision-support work. Axient holds the MDA wargames work. None of the three are "
        "direct competitors for the corrected-scope sub-product at the Navy repair-activity level. CACI NSS is the "
        "same-vendor capability lineage. The other two are procurement-precedents only.")

    _add_heading_2(doc, "4.2 Navy training-systems incumbents — adjacent, not direct")

    _add_body(doc,
        "On 31 March 2026 the Navy modified the NAWCTSD-managed training-systems IDIQ to add $1.2 billion in "
        "ceiling, raising total contract value to $2.51 billion. Nine prime contractors compete for task orders: "
        "BGI-Aero Simulation JV, CAE USA, Delaware Resource Group of Oklahoma, Engineering Support Personnel, "
        "Fidelity Technologies, FlightSafety Defense, LB&B Associates, LTSS JV, and Valiant Global Defense Services. "
        "The work is simulator sustainment and instructional services, NOT the corrected-scope sub-product. The nine "
        "are adjacent incumbents (Navy training relationships, contracting vehicles, installation access) but not "
        "direct competitors. Three (CAE USA, FlightSafety Defense, BGI-Aero Simulation JV) have scenario-design "
        "capability from aviation training and could pivot if a procurement vehicle existed.")

    _add_heading_2(doc, "4.3 Fleet-command exercise-planning incumbents — under-mapped layer")

    _add_body(doc,
        "The biggest research gap in §4. Fleet-command exercise authorities (FLTFORCOM N7, COMPACFLT N7, INDOPACOM "
        "J7, SRF-JRMC's wartime-readiness cell) procure exercise-planning support for COMPTUEX, RIMPAC, Large-Scale "
        "Exercise, SWARMEX, and unit-level workups. The specific contractor incumbencies are not yet source-grounded "
        "in our ingested material. The typical Master Scenario Events List prime base under similar vehicles at "
        "adjacent Combatant Commands likely includes Booz Allen Hamilton, SAIC, and HII Mission Technologies plus "
        "sub-tier players. Pending source-grounding. The next find_sources pass against fleet-command exercise-"
        "planning queries is the load-bearing research task.")

    _add_heading_2(doc, "4.4 Navy PSS incumbents at CNRMC — right customer, wrong vehicle")

    _add_body(doc,
        "Invictus Associates LLC holds Navy delivery order N0016424F3006 against parent IDC N0017819D7883 for "
        "Professional Support Services — Fleet Readiness Support for CNRMC, subordinate RMCs, and Surface Team One. "
        "Total obligation $19,384,385. Norfolk, VA. Navy funder. The parent IDC is SeaPort-NxG, the Navy's generic-"
        "services vehicle. PSS through SeaPort-NxG buys staff augmentation and admin support, not wargaming or "
        "scenario design. Invictus is evidence the right customer procures contractor support generally, but the "
        "vehicle category is the wrong color of money for the corrected-scope product.")

    _add_page_break(doc)

    # 5. Our fit
    _add_heading_1(doc, "5. Our fit — CACI right to win")

    _add_body(doc,
        "CACI's right to win on the corrected-scope progression rests on three pillars plus a fourth procurement-"
        "infrastructure asset that v0.1 of this brief did not surface. The pillars: existing federal contract "
        "footprint in the work-type at multi-hundred-million-dollar scale (§5.1); applied capability lineage from "
        "PMTEC INDOPACOM J7 (§5.2); ARKA acquisition signature libraries for threat-environment realism with HM&E-"
        "data bridge required separately (§5.3). The procurement-infrastructure asset: CACI NSS, LLC already holds "
        "the GSA-side governmentwide vehicles a Navy fleet command can issue task orders against (§5.4).")

    _add_heading_2(doc, "5.1 CACI NSS, LLC at AFRICOM — work-type at scale")

    _add_body(doc,
        "CACI NSS, LLC holds task order PIID 47QFCA20F0042 against parent IDC GS00Q14OADU121 for \"Plans, "
        "Operations, Logistics, Engagement, Training, Exercise, and Assessment Support to AFRICOM.\" Total "
        "obligation $194,034,792. Army funder via GSA Federal Acquisition Service. The task order scope covers the "
        "full work-type stack named in the corrected-scope product description.")

    _add_body(doc,
        "What CACI NSS, LLC actually is. The CACI NSS, LLC subsidiary holds multiple federal contracting vehicles. "
        "It is a general-purpose CACI services subsidiary, not a naval-specific subsidiary. It provides program "
        "management, systems engineering, scientific services, logistics services, IT solutions, and custom software "
        "development across DoD and federal civil agencies. The NSS designation in the entity name is not formally "
        "expanded in legal filings but historically traces to a National Security Solutions line.")

    _add_heading_2(doc, "5.2 PMTEC INDOPACOM J7 capability lineage")

    _add_body(doc,
        "CACI is an active engaged participant in the USINDOPACOM J7 Pacific Multi-Domain Training and "
        "Experimentation Capability (PMTEC) ecosystem, tracked separately in the PMTEC-USINDOPACOM opportunity in "
        "this vault. The capabilities CACI demonstrates under PMTEC are Live-Virtual-Constructive (LVC) integration, "
        "scenario design for multi-domain decision-making, AI and digital-twin technologies, and Regional Joint "
        "Training Infrastructure. These transfer down-vertical to the naval-repair domain. The corrected-scope "
        "progression is a focused sub-application of the broader multi-domain exercise-design capability CACI "
        "demonstrates at PMTEC.")

    _add_heading_2(doc, "5.3 ARKA — threat-environment realism, not repair-side realism")

    _add_body(doc,
        "The March 2026 ARKA acquisition gives CACI electro-optical, infrared, and hyperspectral sensor signature "
        "libraries from intelligence, surveillance, and reconnaissance work. For scenario realism, ARKA drives the "
        "threat side: what a damaged ship's sensors report, what threat indicators the team sees, what the "
        "operational picture looks like with degraded sensors. ARKA does NOT drive repair-side realism: pump-out "
        "status, bulkhead compromise, drydock blocking, supply-chain for replacement valves. Repair-side realism "
        "requires a bridge to NAVSEA / SYSCOM technical-authority data, ship-class engineering documents, or HM&E "
        "databases (MFOM, NMD). Three candidate closures for the bridge: partner with an HM&E-data-holding prime; "
        "leverage CACI's existing naval-IT footprint; negotiate Government Furnished Information as part of any "
        "task order. The GFI path carries 12 to 18 month schedule risk for ATO on a gamified non-system-of-record "
        "environment.")

    _add_heading_2(doc, "5.4 CACI NSS LLC vehicle holdings — the procurement-infrastructure asset")

    _add_body(doc,
        "This is the asset v0.1 of this brief did not surface clearly. CACI NSS, LLC already holds three "
        "governmentwide acquisition vehicles a Navy fleet command can issue task orders against directly:")

    _add_bullet(doc, "GSA OASIS — General Services Administration Federal-wide Professional Services contract.")
    _add_bullet(doc, "NITAAC CIO-SP3 — National Institutes of Health Information Technology Acquisition and Assessment Center IT services contract.")
    _add_bullet(doc, "GSA Multiple Award Schedule (GS-35F-349CA).")

    _add_body(doc,
        "These are governmentwide acquisition contracts open to any federal customer. The Navy can issue task orders "
        "against them without standing up a new Navy-specific contract. The AFRICOM $194M task order is the "
        "precedent — Army funded, GSA awarded, CACI executes against an OASIS-equivalent vehicle CACI already held. "
        "The Navy fleet-command equivalent is \"PACFLT N7 funds, GSA awards, CACI executes against the same vehicle.\" "
        "Same shape. Different funder.")

    _add_body(doc,
        "Implication for the pitch. CACI executives walking into PACFLT N7, FLTFORCOM N7, INDOPACOM J7, or SRF-JRMC "
        "leadership can offer a concrete contracting path on day one. \"You don't need to invent a new vehicle. We "
        "already hold OASIS and CIO-SP3. Issue a Level 1 pilot task order against our existing holdings. The customer "
        "commits no new procurement infrastructure.\" That is materially easier salesmanship than asking the customer "
        "to fund a new contract from scratch.")

    _add_heading_2(doc, "5.5 OCI and contact-protection")

    _add_body(doc,
        "The operator's research-origin context triggers FAR 9.5 OCI considerations and §9.3 contact-protection "
        "discipline. The OCI analysis is operator-owned and runs separately from this brief. The §9.3 discipline "
        "keeps operator-side knowledge out of vault analytical content. The recommendation here uses only OSI-"
        "sourced claims as load-bearing.")

    _add_page_break(doc)

    # 6. Working hypothesis
    _add_heading_1(doc, "6. Working hypothesis and falsifying legs")

    _add_body(doc,
        "Hypothesis: the Navy has real demand for contractor-supplied operational-decision-scenario content. The "
        "specific products are the five levels of the progression at §11.3 of the research file. CACI's right to "
        "win rests on the four assets in §5. The most-likely procurement-pathway is task order against CACI NSS "
        "LLC's existing OASIS, CIO-SP3, or GSA MAS holdings, sponsored operationally by SRF-JRMC and contracted by "
        "PACFLT N7 or FLTFORCOM N7. CACI executives have to pitch the concept directly. The Navy is not running an "
        "RFP today.")

    _add_heading_2(doc, "6.1 Six falsifying legs")

    _add_bullet(doc,
        "Leg 1 — Navy operational-decision-scenario demand gap. Killed by senior signals turning ceremonial AND "
        "existing fleet exercises already injecting equivalent content organically.")
    _add_bullet(doc,
        "Leg 2 — Sub-product moat. Killed by incumbents replicating the analog form factor at the same time durations "
        "OR the AI scenario-generation capability maturing rapidly enough among incumbents that CACI's PMTEC head "
        "start does not translate to a moat by procurement time.")
    _add_bullet(doc,
        "Leg 3 — CACI capability transferability (PMTEC to naval-repair). Killed by PMTEC's multi-domain Combatant "
        "Command staff focus being structurally different from naval-repair small-unit decision-rehearsal.")
    _add_bullet(doc,
        "Leg 4 — ARKA signature-library differentiator viability. Killed by (a) IP posture blocking release; (b) "
        "comparable data available to incumbents; (c) HM&E-data bridge cannot be closed; (d) GFI / ATO friction "
        "introduces 12-18 month schedule delay neutralizing the head start.")
    _add_bullet(doc,
        "Leg 5 — Customer-access pathway feasibility. Killed by fleet-command incumbents locking out CACI entry "
        "timeline, or contractual right-of-first-refusal preventing new-entrant access.")
    _add_bullet(doc,
        "Leg 6 — Sub-product viability as fundable line. Killed by sub-product too small for standalone procurement "
        "(partner not prime) OR too specialized to attract any vehicle (need DIU / SBIR / STTR pilot first).")

    _add_page_break(doc)

    # 7. Recommendation
    _add_heading_1(doc, "7. Recommendation (draft)")

    _add_body(doc,
        "Position CACI as a candidate prime for the five-level progression. Initial entry through Level 1 or Level 2 "
        "pilot. Customer end-user is SRF-JRMC. Contracting authority is PACFLT N7 / N4 or FLTFORCOM N7 / N4. "
        "Contracting vehicle is CACI NSS LLC's existing OASIS, CIO-SP3, or GSA MAS holding — no new Navy vehicle "
        "required.")

    _add_body(doc,
        "CACI executives have to pitch the concept directly. The Navy is not running an RFP for this today. The "
        "path to procurement runs through proactive executive engagement with PACFLT N7, FLTFORCOM N7, INDOPACOM J7, "
        "NAVSEA 04, or SRF-JRMC leadership. The pitch is concrete on day one: here is the work-type, here is the "
        "AFRICOM precedent, here are the GSA-wide vehicles we already hold, issue a Level 1 pilot task order against "
        "our existing OASIS or CIO-SP3 holdings.")

    _add_body(doc,
        "Alternative entry path if the direct task-order pitch does not produce procurement inside the planning "
        "horizon. DIU, SBIR, or STTR pilot for Level 1 or Level 2 scope. Lower per-engagement scale. Unblocks the "
        "no-RFP problem. Produces a reference customer for follow-on task-order competition.")

    _add_body(doc,
        "Close the FAR 9.5 OCI analysis before the pitch meetings happen. Decide the HM&E-data bridge approach: "
        "partner, internal naval-IT, or GFI / ATO. ARKA covers threat-side only.")

    _add_body(doc,
        "Reassessment trigger. The recommendation shape will be reassessed after (a) the next find_sources pass "
        "against fleet-command-exercise-planning queries completes and produces source-grounded incumbent "
        "identification at the §4.3 layer; (b) the §11.1 engagement-surface inventory maps the specific contracting "
        "offices and active solicitations at FLTFORCOM N7 / N4, PACFLT N7 / N4, SRF-JRMC, NAVSEA 04, and PEO IWS; "
        "and (c) the operator's OCI analysis closes per FAR 9.5.")

    _add_page_break(doc)

    # 8. Risks
    _add_heading_1(doc, "8. Risks")

    _add_heading_2(doc, "8.1 No existing Navy task order")
    _add_body(doc,
        "Biggest risk. The Navy is not procuring this sub-product today. We are pitching a task order that does not "
        "exist yet. If salesmanship fails, entry collapses to DIU / SBIR / STTR pilot or subcontract on someone "
        "else's vehicle. The OASIS / CIO-SP3 procurement-path finding reduces this risk but does not eliminate it. "
        "The Navy still has to decide to issue the task order.")

    _add_heading_2(doc, "8.2 Incumbent-lockout")
    _add_body(doc,
        "Fleet-command exercise-planning incumbent base is real but under-mapped. Candidate names (Booz Allen "
        "Hamilton, SAIC, HII Mission Technologies, sub-tier players) are Assessment-tagged pending source-grounding. "
        "The actual subset active at FLTFORCOM N7, COMPACFLT N7, INDOPACOM J7, and SRF-JRMC could be either more "
        "crowded (forcing partnering) or sparser (enabling first-mover prime entry).")

    _add_heading_2(doc, "8.3 HM&E-data bridge schedule risk")
    _add_body(doc,
        "Repair-side scenario realism needs HM&E, drydock, and MFOM / NMD data. ARKA does not directly address this. "
        "If the GFI / ATO path is chosen, securing Authority to Operate for a gamified non-system-of-record "
        "environment can take 12 to 18 months. Mitigations: partner with an HM&E-data-holding prime; leverage CACI's "
        "existing naval-IT footprint; start the ATO process now in parallel with engagement work.")

    _add_heading_2(doc, "8.4 OCI")
    _add_body(doc,
        "Operator's research-origin knowledge of contractor exercise planners at SRF-JRMC's wartime readiness group "
        "triggers FAR 9.5 OCI analysis. Owner: operator. If a finding limits CACI's competitive position, the "
        "recommendation may need to change before pursuit.")

    _add_heading_2(doc, "8.5 Source-grounding caveat")
    _add_body(doc,
        "The sources prove the Navy has a wartime-repair problem. The sources prove the rest of DoD buys related "
        "work-types at scale. The sources do NOT yet prove a Navy procurement line for the corrected-scope sub-"
        "product exists today. The hypothesis is that the Navy needs this work, the rest of DoD buys this kind of "
        "work, and therefore the Navy should buy this work too. The \"should\" is real. The \"is buying today\" is "
        "the gap CACI executive salesmanship has to help close.")

    _add_page_break(doc)

    # 9. Asks
    _add_heading_1(doc, "9. Asks")

    _add_heading_2(doc, "9.1 Asks of the BD team")
    _add_bullet(doc, "Confirm the corrected-scope frame at the BD-team gate review. Operational-decision-scenarios for the BDAR team and the fleet commanders they support. Not schoolhouse curriculum.")
    _add_bullet(doc, "Authorize the next find_sources pass against fleet-command-exercise-planning queries at the named Navy commands.")
    _add_bullet(doc, "Approve the engagement-surface inventory (§11.1) work targeting the named contracting offices and the SYSCOM authorities (NAVSEA 04, PEO IWS).")
    _add_bullet(doc, "Decide the HM&E-data bridge approach: partner, internal naval-IT, or GFI / ATO.")

    _add_heading_2(doc, "9.2 Asks of the operator")
    _add_bullet(doc, "Close the FAR 9.5 OCI analysis given operator-side knowledge of contractor exercise planners at SRF-JRMC's wartime readiness group.")
    _add_bullet(doc, "Confirm ARKA intellectual-property and release-authority posture for training-environment use of EO/IR/hyperspectral signature libraries.")
    _add_bullet(doc, "Authorize CACI executive engagement with Navy leadership at the named commands. This is the salesmanship step the procurement-vehicle question requires.")

    _add_page_break(doc)

    # 10. Source ledger
    _add_heading_1(doc, "10. Source ledger reference")

    _add_body(doc,
        "Load-bearing FACT claims trace to primary or quasi-primary sources in `01_sources/`. Full ledger lives in "
        "research-file §8. Key sources:")

    _add_bullet(doc, "DON FY 2027 President's Budget Press Brief, 28 Apr 2026.")
    _add_bullet(doc, "CNO Caudle HASC Statement, 14 May 2026.")
    _add_bullet(doc, "SWARMEX-Cebu primary releases — Navy.mil, C7F, INDOPACOM, COMPACFLT, NavalNews, NavalToday.")
    _add_bullet(doc, "USNI Proceedings \"Fix the Navy's Expeditionary Repair.\"")
    _add_bullet(doc, "CIMSEC \"If the U.S. Navy Can't Repair Ships in Peacetime, How Will It Do So in War?\"")
    _add_bullet(doc, "FY27 DoD Comptroller Justification Book — MAGTF Tactical Warfare Simulation, JExD, AI scenario generation funding lines.")
    _add_bullet(doc, "USAspending awards: CACI NSS LLC $194M AFRICOM, Parsons Government Services $556.8M CEOIS, Axient LLC $233.8M MDA Wargames, Invictus Associates $19.4M CNRMC PSS.")
    _add_bullet(doc, "GovConWire — Navy nine-contractors $1.2B NAWCTSD training-systems IDIQ modification.")
    _add_bullet(doc, "Norfolk Naval Shipyard $442M drydock modernization (workboat.com).")
    _add_bullet(doc, "Pacific NAVFAC SIOP Brief — PSNS-SBS infrastructure.")
    _add_bullet(doc, "CACI NSS, LLC vehicle holdings — GSA OASIS, NITAAC CIO-SP3, GSA MAS GS-35F-349CA (web-verified 2026-05-26).")

    _add_italic(doc,
        f"End of capture brief v0.2 draft. Build date: {date.today().isoformat()}. Built from sealed §§1, 3, 4, 5, 6, "
        f"7, and the five-level progression in §11.3 of 00_research-file.md.")

    doc.save(str(OUT))
    print(f"✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    build()
