#!/usr/bin/env python3
"""
build_bdr_exec_brief_v02.py — Build BDR-FLEET-READINESS executive brief v0.2.

v0.2 reasons for the rewrite:
  - Operator feedback on v0.1: "still a little difficult to read."
  - Operator added the 5-level progression (analog-first) per the 2026-05-26 brainstorm.
  - Operator wants the hypothesis stated cleanly without soft-pedaling.
  - Operator wants the salesmanship step explicit. Navy is not running an RFP today.
  - Source-grounding honesty: the brief now flags that the Navy procurement vehicle for
    the corrected-scope sub-product does NOT exist today.
"""

from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

VAULT_ROOT = Path(__file__).parent.parent
OUT = VAULT_ROOT / "opportunities/BDR-FLEET-READINESS/04_artifacts/executive-brief-v0.2-draft.docx"


def _set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _set_tight_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)


def _body_compact(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def _bullet_compact(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"• {text}")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_default_font(doc)
    _set_tight_margins(doc)

    # Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run("BDR-FLEET-READINESS — Executive Brief v0.2 (Draft)")
    r.font.bold = True
    r.font.size = Pt(13)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run(f"Navy BDAR decision-rehearsal training. Five-level analog-first progression. "
                    f"{date.today().isoformat()}. Internal exec review, OSI-only.")
    r.font.size = Pt(9)
    r.font.italic = True

    # BLUF
    _section_heading(doc, "Bottom Line Up Front")

    _body_compact(doc,
        "The Navy has a real wartime-repair readiness problem. The rest of the Defense Department already buys exercise "
        "design and decision-rehearsal content from contractors at multi-hundred-million-dollar scale. The Navy should "
        "buy this kind of work too. Today the Navy does not have a named procurement vehicle for it. That gap is the "
        "opportunity.")

    _body_compact(doc,
        "Recommendation. CACI should pursue a five-level progression that starts with simple analog wardroom drills "
        "and ends with software-driven scenario content injected into live fleet exercises. The first customer is most "
        "likely the Ship Repair Facility in Japan (SRF-JRMC). The contracting authority is most likely Pacific Fleet "
        "or U.S. Fleet Forces Command N7 or N4. The pitch has to come from CACI. This is not a wait-for-the-RFP play. "
        "CACI executives have to sell the concept to Navy leadership directly.")

    _body_compact(doc,
        "CACI has the right to pitch this. Three reasons. CACI NSS LLC runs the same kind of work for AFRICOM today "
        "at $194 million. CACI builds INDOPACOM PMTEC exercise scenarios already. The 2026 ARKA acquisition gives CACI "
        "signature-library content for threat-side scenario realism.")

    # Why It Matters
    _section_heading(doc, "Why It Matters")

    _bullet_compact(doc,
        "Navy senior leadership is signaling demand. CNO Caudle named maintenance a warfighting requirement at HASC on "
        "14 May 2026. The FY27 budget funds $17B for ship maintenance to drive an 80% Combat Surge Ready posture. "
        "SRF-JRMC just ran a wartime-repair exercise with the Philippine Navy at Cebu.")

    _bullet_compact(doc,
        "The Defense Department buys this work from contractors. CACI NSS at AFRICOM ($194M). Parsons at OSD ($557M). "
        "Axient at MDA ($234M). The Joint Staff funds AI-powered scenario generation as an FY27 budget line.")

    _bullet_compact(doc,
        "The Navy professional community sees the gap. USNI Proceedings published \"Fix the Navy's Expeditionary "
        "Repair.\" CIMSEC published \"If the U.S. Navy Can't Repair Ships in Peacetime, How Will It Do So in War?\" "
        "Both publications reach the customer.")

    _bullet_compact(doc,
        "What the sources do not yet prove. The Navy does not have a named procurement vehicle for 1-hour gamified "
        "decision rehearsal today. The hypothesis is that the Navy needs this work, the rest of the Defense Department "
        "buys this kind of work, and therefore the Navy should buy this work too. The \"should\" is real. The \"is "
        "buying today\" is the gap CACI executives have to help close through direct salesmanship.")

    # The Progression
    _section_heading(doc, "The Five-Level Progression")

    _body_compact(doc,
        "Start analog. Move to software when scale justifies it. Each level can be procured separately. Each level "
        "builds reference material for the next.")

    _bullet_compact(doc,
        "Level 1. Flash drills. Five to fifteen minutes. One decision per drill on a printed scenario card. Pattern "
        "recognition. Cheap to run. Easy to sell as a pilot.")

    _bullet_compact(doc,
        "Level 2. Linked-decision sequences. Thirty to forty-five minutes. Three to five decisions in a row. Each "
        "decision shapes the next. Wardroom learns that early choices box you in later.")

    _bullet_compact(doc,
        "Level 3. Full 1-hour gamified wardroom session. Multi-discipline audience. Turn-based. All six operational-"
        "decision moments in one integrated scenario. Analog or analog-plus-software.")

    _bullet_compact(doc,
        "Level 4. Multi-session campaign. Two to four hours across a deployment cycle. Decisions in session one carry "
        "into session two. This is where AI scenario generation pays for itself.")

    _bullet_compact(doc,
        "Level 5. Live exercise injection. Scenarios injected into COMPTUEX, RIMPAC, Large-Scale Exercise, SWARMEX. "
        "Software driven and federated into the Navy synthetic training architecture.")

    _body_compact(doc,
        "Capture path is sequential. Sell Level 1 to one customer first. Use the analog reference to justify Level 3 "
        "at the same customer. Use the Level 3 reference to justify the Level 4 software investment.")

    # Recommendation
    _section_heading(doc, "Recommendation")

    _bullet_compact(doc,
        "Pursue the five-level progression. Start at SRF-JRMC if access permits. Otherwise start at PACFLT or Fleet "
        "Forces N7.")

    _bullet_compact(doc,
        "CACI executives have to pitch the concept directly. The Navy is not running an RFP for this today. The path "
        "to procurement runs through proactive engagement with PACFLT N7, FLTFORCOM N7, INDOPACOM J7, NAVSEA 04, or "
        "SRF-JRMC leadership.")

    _bullet_compact(doc,
        "Alternative entry path if direct pitching does not produce a vehicle inside the planning horizon. DIU, SBIR, "
        "or STTR pilot. Lower per-engagement scale. Unblocks the no-vehicle problem. Produces a reference for follow-on.")

    _bullet_compact(doc,
        "Close the FAR 9.5 OCI question before pitch meetings happen. Decide the HM&E-data bridge approach for "
        "repair-side scenario realism. ARKA covers threat-side only.")

    # Top Risks
    _section_heading(doc, "Top Risks")

    _bullet_compact(doc,
        "No existing Navy vehicle. The biggest risk. We are pitching a sub-product the Navy is not procuring today. "
        "If salesmanship fails, entry collapses to DIU or SBIR pilot or subcontract on someone else's vehicle.")

    _bullet_compact(doc,
        "Incumbent lockout. Fleet-command exercise planners exist at PACFLT, FLTFORCOM, INDOPACOM J7, and SRF-JRMC. "
        "Likely names are Booz Allen, SAIC, HII Mission Technologies plus sub-tier players. We do not yet know the "
        "actual subset. The next find_sources pass is targeting that gap.")

    _bullet_compact(doc,
        "OCI. The operator-side knowledge of contractor exercise planners at SRF-JRMC may trigger FAR 9.5 analysis. "
        "That work is operator-owned and runs separately from this brief.")

    # Asks
    _section_heading(doc, "Asks")

    _bullet_compact(doc,
        "Authorize the next find_sources pass to map fleet-command exercise-planning incumbents at PACFLT N7, "
        "FLTFORCOM N7, INDOPACOM J7, and SRF-JRMC.")

    _bullet_compact(doc,
        "Decide the HM&E-data bridge approach. Three options. Partner with a prime that has it. Lean on CACI internal "
        "naval IT. Negotiate Government Furnished Information with the customer.")

    _bullet_compact(doc,
        "Close the FAR 9.5 OCI analysis given the operator-side knowledge of contractor exercise planners at SRF-JRMC.")

    _bullet_compact(doc,
        "Authorize CACI executive engagement with Navy leadership at the named commands. This is the salesmanship step "
        "the procurement-vehicle question requires.")

    # Footer
    f = doc.add_paragraph()
    f.paragraph_format.space_before = Pt(6)
    r = f.add_run(
        f"v0.2 built {date.today().isoformat()} from sealed §§1, 3, 4, 5, 6, 7 and the rebuilt §11.3 five-level "
        f"progression in 00_research-file.md. Source-grounding gap on Navy procurement vehicle for the corrected-"
        f"scope sub-product is flagged explicitly per operator's 2026-05-26 feedback. Capture brief v0.1 is the "
        f"long-form companion."
    )
    r.font.size = Pt(8)
    r.font.italic = True

    doc.save(str(OUT))
    print(f"✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    build()
