#!/usr/bin/env python3
"""
build_bdr_capture_brief_v03.py — Build BDR-FLEET-READINESS capture brief v0.3.

v0.3 reframe per operator feedback 2026-05-26 (third round):
  1. Drop "Asks of the operator" section. The operator is the analyst delivering this brief,
     not the decision-maker. The brief is going up the chain to CACI executives. Items the brief
     previously asked the operator to do now move to either "Recommendations for executive action"
     or to "Flagged concerns for executive review."
  2. FAR 9.5 OCI moves to a flagged concern for executive review. Not an operator action item.
  3. CACI executive engagement at Navy leadership becomes a recommendation, not an ask.
  4. HM&E-data bridge updated with operator-side direction: most likely GFI source is NSWC Carderock,
     with possible additional sources at another Warfare Center or somewhere in PAE Industrial Operations.
     Framed as "operator-side scoping points toward..." — directional input, not a locked decision.
"""

from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

VAULT_ROOT = Path(__file__).parent.parent
OUT = VAULT_ROOT / "opportunities/BDR-FLEET-READINESS/04_artifacts/capture-brief-v0.3-draft.docx"


def _set_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(18)


def _h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(14)


def _h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(12)


def _body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def _bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(f"• {text}")


def _italic(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.italic = True


def _pb(doc):
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_styles(doc)

    # Cover
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(36)
    r = cover.add_run("CAPTURE BRIEF — v0.3 DRAFT\n")
    r.font.bold = True
    r.font.size = Pt(22)

    r = cover.add_run("\nBDR-FLEET-READINESS\n")
    r.font.bold = True
    r.font.size = Pt(18)

    r = cover.add_run("\nFive-level analog-first progression of gamified decision-rehearsal scenarios for Navy BDAR teams and fleet commanders\n")
    r.font.size = Pt(13)
    r.font.italic = True

    r = cover.add_run(f"\nDraft date: {date.today().isoformat()}  |  For CACI executive review (OSI-only).\n")
    r.font.size = Pt(10)

    r = cover.add_run("\nv0.3 revisions: structural reframe to executive-facing format. \"Asks of the operator\" section dropped — operator is the analyst delivering this brief, not the decision-maker. FAR 9.5 OCI moves to flagged concern. CACI executive engagement moves to recommendation. HM&E-data bridge updated with operator-side directional input toward NSWC Carderock and other Warfare Centers / PAE Industrial Operations.\n")
    r.font.size = Pt(9)
    r.font.italic = True

    _pb(doc)

    # 1. BLUF
    _h1(doc, "1. Bottom Line Up Front")

    _body(doc,
        "The Navy has a real wartime-repair readiness problem. The rest of the Defense Department already buys "
        "exercise and decision-rehearsal content from contractors at multi-hundred-million-dollar scale. The Navy "
        "should buy this kind of work too. Today the Navy has not issued a task order for it. That gap is the "
        "opportunity, and the path to closing it is concrete.")

    _body(doc,
        "The product is a five-level progression of gamified decision-rehearsal scenarios. Level 1 is 5 to 15 minute "
        "flash drills on printed cards. Level 2 is 30 to 45 minute linked-decision sequences. Level 3 is the 1-hour "
        "full gamified wardroom session. Level 4 is a 2 to 4 hour multi-session campaign across a deployment cycle. "
        "Level 5 is software-driven scenario injection into live fleet exercises. Start analog. Move to software at "
        "Level 4 when the customer commits to scale. The same scenario content serves two audiences: the repair-"
        "activity wardroom at a shipyard, RMC, or SRF-JRMC, AND the fleet commander deciding where to send a damaged "
        "ship. Port selection is the canonical dual-audience scenario.")

    _body(doc,
        "The procurement path is concrete. CACI NSS, LLC already holds GSA OASIS, NITAAC CIO-SP3, and GSA Multiple "
        "Award Schedule. The AFRICOM $194M task order is the precedent — Army funded, GSA awarded, CACI executes on "
        "a vehicle CACI already held. The Navy can do the same thing. Any Navy fleet command (PACFLT N7, FLTFORCOM "
        "N7, INDOPACOM J7) or SYSCOM authority (NAVSEA 04, PEO IWS) can issue a task order against CACI's existing "
        "OASIS or CIO-SP3 vehicles without setting up a new Navy-specific contract.")

    _body(doc,
        "Recommendation. CACI executive leadership should initiate engagement with PACFLT N7, FLTFORCOM N7, "
        "INDOPACOM J7, NAVSEA 04, or SRF-JRMC leadership and pitch the concept directly. The Navy is not running an "
        "RFP for this today. The pitch is concrete on day one: here is the work-type, here is the AFRICOM precedent, "
        "here are the GSA-wide vehicles CACI already holds, issue a Level 1 pilot task order. The salesmanship step "
        "is real BD work. The alternative entry path if the direct task-order pitch does not produce procurement "
        "inside the planning horizon is DIU, SBIR, or STTR pilot for Level 1 or Level 2 scope.")

    _pb(doc)

    # 2. Demand signal
    _h1(doc, "2. Demand signal")

    _body(doc,
        "Four signals converge on the corrected-scope demand. Senior Navy leadership is publicly committing to "
        "wartime repair readiness. The Navy is publicly running the most directly relevant exercise type. The Navy "
        "professional community is publicly arguing the gap is real. And the Defense Department broadly procures "
        "contractor-delivered exercise and decision-support content at multi-hundred-million-dollar scale.")

    _h2(doc, "2.1 Senior-leadership signal")

    _body(doc,
        "Admiral Daryl L. Caudle, the Chief of Naval Operations, delivered a Statement on the Posture of the United "
        "States Navy before the House Armed Services Committee on 14 May 2026, naming maintenance as a warfighting "
        "requirement and announcing a planned deliberate study of Navy yard capacity. The DON FY27 Press Brief funds "
        "Ship Maintenance at $17.0 billion to drive an 80% Combat Surge Ready posture, plus a $0.6B Contested "
        "Logistics line because adversaries will target supply lines, ports, and communications. 37 ships are listed "
        "as Forward Deployed Naval Forces.")

    _h2(doc, "2.2 SWARMEX-Cebu")

    _body(doc,
        "SRF-JRMC, U.S. Seventh Fleet, U.S. Pacific Command, and U.S. Pacific Fleet published parallel releases "
        "announcing that USS Ashland (LSD 48) completed Ship Wartime Repair and Maintenance in Cebu, Philippines, "
        "with Philippine Navy partners and local Philippine contractors. The exercise involved an amphibious warfare "
        "ship operating at a foreign port, host-nation contractor coordination, and Philippine Naval Sea Systems "
        "Command participation. The exercise structure maps directly onto three of the six operational-decision "
        "moments the corrected-scope product addresses: forward team mobilization, foreign-port emergency "
        "contracting, and host-nation framework navigation.")

    _h2(doc, "2.3 Policy-community pressure")

    _body(doc,
        "USNI Proceedings published \"Fix the Navy's Expeditionary Repair.\" CIMSEC published \"If the U.S. Navy "
        "Can't Repair Ships in Peacetime, How Will It Do So in War?\" Both publications argue current expeditionary "
        "repair capability is inadequate for wartime contestation. Both reach the customer organizations.")

    _h2(doc, "2.4 Three distinct procurement precedents — keep them separate")

    _body(doc,
        "Three federal contracts give the Defense Department's procurement pattern at scale. They are three distinct "
        "work-types, not one. Each tells the BD team something different.")

    _body(doc,
        "**CACI NSS, LLC at AFRICOM.** $194 million task order for broad professional services to AFRICOM with "
        "exercise, training, and operations support as named scope. Army funder via GSA Federal Acquisition Service. "
        "Awarded against parent IDC GS00Q14OADU121, an OASIS-equivalent governmentwide vehicle. This is the closest "
        "analog to the corrected-scope work-type and the single most consequential capability-lineage anchor for "
        "CACI's right to win.")

    _body(doc,
        "**Parsons Government Services at OSD.** $556.8 million CEOIS task order. CEOIS expands to Command, Control, "
        "Communications, Computers, Combat Systems, Intelligence, Surveillance, Reconnaissance, and Targeting / "
        "Exercises / Operations / Information Services. The work is decision-support for active Combatant Command "
        "operations — software platforms and contractor staff embedded at INDOPACOM, EUCOM, AFRICOM, CENTCOM "
        "supporting the four-star and the J-staff in real time. This is NOT exercise design. It is decision-support "
        "for ongoing operations. Different work-type from the corrected-scope sub-product but adjacent.")

    _body(doc,
        "**Axient LLC at MDA.** $233.8 million Test, Exercise and Wargames Support contract for missile defense "
        "testing and operator training. MDA funder. Place of performance Huntsville, Alabama. This IS exercise and "
        "wargames work, but for a different domain (missile defense, not naval repair). Closest precedent for a "
        "single DoD agency procuring a multi-hundred-million-dollar exercise-and-wargames contract in one place.")

    _body(doc,
        "What the three together tell the BD team. The Defense Department definitely procures the corrected-scope "
        "work-type from contractors at scale (AFRICOM via CACI NSS, MDA via Axient). DoD also procures the adjacent "
        "decision-support work-type at scale (OSD via Parsons). The Navy could plausibly procure either or both. "
        "Today the Navy has not. The next find_sources pass is targeting Navy fleet-command vehicles to see if "
        "something similar has been procured through pathways the current research has not yet surfaced.")

    _pb(doc)

    # 3. Customer landscape
    _h1(doc, "3. Customer landscape")

    _body(doc,
        "The corrected-scope customer landscape is layered. Strategic direction sits with SECNAV and CNO. They set "
        "priorities and authorize budget but do not procure the product directly. The repair-activity layer is the "
        "four public naval shipyards (Norfolk, Puget Sound, Pearl Harbor, Portsmouth), the Regional Maintenance "
        "Centers including the forward-deployed SRF-JRMC, and Commander, Navy Regional Maintenance Center (CNRMC) "
        "as parent command. The doctrinal authority is Naval Warfare Development Center (NWDC). The fleet-command "
        "training-authority layer is FLTFORCOM, COMPACFLT, COMNAVSURFOR, and INDOPACOM J7.")

    _h2(doc, "3.1 Two audiences for the same scenario base")

    _body(doc,
        "The corrected-scope product serves two audiences with the same scenario content. The repair-activity "
        "wardroom — Captain of the shipyard, Executive Officer of the RMC, SRF-JRMC commander, plus department "
        "heads and visiting fleet liaisons — prepares to receive a damaged ship and mobilize against the work. The "
        "fleet commander — the tactical operator who owns the ship — decides where to send the damaged ship in the "
        "first place. Same scenario content. Different decision-maker focus.")

    _body(doc,
        "Port selection is the canonical dual-audience scenario. A destroyer takes shock damage 800 miles from "
        "three friendly ports. The fleet commander runs through the variables: proximity, repair capability, "
        "security posture, alliance access. He picks the port. The repair-activity wardroom at the chosen port then "
        "runs the same scenario from the receiving end: receive the ship, stage the repair team, hire local welders, "
        "handle customs. The product can run as a wardroom exercise at the receiving end (Level 3) or as a fleet-"
        "commander decision rehearsal at the sending end (Level 1 flash drill or Level 3 full session depending on "
        "the customer).")

    _h2(doc, "3.2 Procurement model — end-user separate from contracting authority")

    _body(doc,
        "The operationally-most-exposed end-user (SRF-JRMC) is not necessarily the contracting authority. SRF-JRMC "
        "has Operations and Maintenance dollars to fix forward-deployed ships. It does not typically possess the "
        "RDT&E or Other Procurement, Navy budget authority needed for a net-new gamified software platform. The "
        "capture model decouples the two layers. SRF-JRMC sponsors the operational requirement. PACFLT N7 / N4 or "
        "FLTFORCOM N7 / N4 holds the contracting authority and the right budget color.")

    _h3(doc, "End-user / sponsor ranking")
    _bullet(doc, "SRF-JRMC wartime readiness cell — forward-deployed, audience and venue match the product.")
    _bullet(doc, "CNRMC plus subordinate RMCs plus Surface Team One — Navy-wide repair-activity layer.")
    _bullet(doc, "Public naval shipyards (NNSY, PSNS, PNSY, PHNS) — heaviest investment, longer cycles.")

    _h3(doc, "Contracting-authority ranking — task order against existing CACI vehicles")
    _bullet(doc, "PACFLT N7 / N4 or FLTFORCOM N7 / N4 — owns the fleet-exercise rhythm; can issue a task order against CACI NSS LLC's existing OASIS or CIO-SP3 holdings.")
    _bullet(doc, "INDOPACOM J7 — owns joint multi-domain exercise procurement; same task-order path.")
    _bullet(doc, "DIU, SBIR, or STTR — alternative entry if direct task-order pitch does not produce procurement inside the planning horizon.")
    _bullet(doc, "NAVSEA 04 or PEO IWS — SYSCOM-level authority; longer cycle but applicable if the product enters as part of a program of record.")

    _pb(doc)

    # 4. Competitive landscape
    _h1(doc, "4. Competitive landscape")

    _body(doc,
        "Four competitive layers. Cross-customer Exercise / Wargames / Decision-Support incumbents at federal scale. "
        "Navy training-systems incumbents adjacent but not directly competitive. Fleet-command exercise-planning "
        "incumbents under-mapped in public sources but the most-likely direct-competition layer. Navy professional-"
        "services incumbents at the corrected-scope customer.")

    _h2(doc, "4.1 Three distinct adjacent-work-type incumbents")

    _body(doc,
        "Per §2.4 above, the three federal-contract precedents are three distinct work-types. CACI itself is the "
        "dominant incumbent in the closest analog via NSS LLC at AFRICOM. Parsons holds the OSD decision-support "
        "work. Axient holds the MDA missile-defense wargames work. None of the three are direct competitors for "
        "the corrected-scope sub-product at the Navy repair-activity level. CACI NSS is the same-vendor capability "
        "lineage. The other two are procurement-precedents only.")

    _h2(doc, "4.2 Navy training-systems incumbents — adjacent, not direct")

    _body(doc,
        "On 31 March 2026 the Navy modified the NAWCTSD-managed training-systems IDIQ to add $1.2 billion in "
        "ceiling, raising total contract value to $2.51 billion. Nine prime contractors compete for task orders: "
        "BGI-Aero Simulation JV, CAE USA, Delaware Resource Group of Oklahoma, Engineering Support Personnel, "
        "Fidelity Technologies, FlightSafety Defense, LB&B Associates, LTSS JV, and Valiant Global Defense Services. "
        "The work is simulator sustainment and instructional services, NOT the corrected-scope sub-product. The "
        "nine are adjacent incumbents but not direct competitors. Three (CAE USA, FlightSafety Defense, BGI-Aero "
        "Simulation JV) have scenario-design capability from aviation training and could pivot if a procurement "
        "vehicle existed.")

    _h2(doc, "4.3 Fleet-command exercise-planning incumbents — under-mapped layer")

    _body(doc,
        "The biggest research gap in §4. Fleet-command exercise authorities (FLTFORCOM N7, COMPACFLT N7, INDOPACOM "
        "J7, SRF-JRMC's wartime-readiness cell) procure exercise-planning support for COMPTUEX, RIMPAC, Large-Scale "
        "Exercise, SWARMEX, and unit-level workups. The specific contractor incumbencies are not yet source-grounded "
        "in our ingested material. The typical Master Scenario Events List prime base under similar vehicles at "
        "adjacent Combatant Commands likely includes Booz Allen Hamilton, SAIC, and HII Mission Technologies plus "
        "sub-tier players. Pending source-grounding. A new find_sources pass against fleet-command exercise-planning "
        "queries was authorized on 2026-05-26; results will be reflected in v0.4 of this brief.")

    _h2(doc, "4.4 Navy PSS incumbents at CNRMC — right customer, wrong vehicle")

    _body(doc,
        "Invictus Associates LLC holds Navy delivery order N0016424F3006 against parent IDC N0017819D7883 for "
        "Professional Support Services — Fleet Readiness Support for CNRMC, subordinate RMCs, and Surface Team One. "
        "Total obligation $19,384,385. Norfolk, VA. Navy funder. The parent IDC is SeaPort-NxG, the Navy's generic-"
        "services vehicle. PSS through SeaPort-NxG buys staff augmentation and admin support, not wargaming or "
        "scenario design. Invictus is evidence the right customer procures contractor support generally, but the "
        "vehicle category is the wrong color of money for the corrected-scope product.")

    _pb(doc)

    # 5. Our fit
    _h1(doc, "5. Our fit — CACI right to win")

    _body(doc,
        "CACI's right to win rests on three capability pillars plus a fourth procurement-infrastructure asset. "
        "Existing federal contract footprint in the work-type at multi-hundred-million-dollar scale (§5.1). Applied "
        "capability lineage from PMTEC INDOPACOM J7 (§5.2). ARKA acquisition signature libraries for threat-"
        "environment realism with HM&E-data bridge required separately (§5.3). And the procurement-infrastructure "
        "asset: CACI NSS, LLC already holds the GSA-side governmentwide vehicles a Navy fleet command can issue task "
        "orders against (§5.4).")

    _h2(doc, "5.1 CACI NSS, LLC at AFRICOM — work-type at scale")

    _body(doc,
        "CACI NSS, LLC holds task order PIID 47QFCA20F0042 against parent IDC GS00Q14OADU121 for \"Plans, "
        "Operations, Logistics, Engagement, Training, Exercise, and Assessment Support to AFRICOM.\" Total "
        "obligation $194,034,792. Army funder via GSA Federal Acquisition Service. The task order scope covers the "
        "full work-type stack named in the corrected-scope product description.")

    _body(doc,
        "What CACI NSS, LLC actually is. The CACI NSS, LLC subsidiary holds multiple federal contracting vehicles. "
        "It is a general-purpose CACI services subsidiary, not a naval-specific subsidiary. It provides program "
        "management, systems engineering, scientific services, logistics services, IT solutions, and custom "
        "software development across DoD and federal civil agencies. The NSS designation in the entity name is "
        "historically tied to a National Security Solutions line.")

    _h2(doc, "5.2 PMTEC INDOPACOM J7 capability lineage")

    _body(doc,
        "CACI is an active engaged participant in the USINDOPACOM J7 Pacific Multi-Domain Training and "
        "Experimentation Capability (PMTEC) ecosystem, tracked separately in the PMTEC-USINDOPACOM opportunity in "
        "this vault. The capabilities CACI demonstrates under PMTEC are Live-Virtual-Constructive (LVC) integration, "
        "scenario design for multi-domain decision-making, AI and digital-twin technologies, and Regional Joint "
        "Training Infrastructure. These transfer down-vertical to the naval-repair domain. The corrected-scope "
        "progression is a focused sub-application of the broader multi-domain exercise-design capability CACI "
        "demonstrates at PMTEC.")

    _h2(doc, "5.3 ARKA — threat-environment realism. HM&E bridge required separately.")

    _body(doc,
        "The March 2026 ARKA acquisition gives CACI electro-optical, infrared, and hyperspectral sensor signature "
        "libraries from intelligence, surveillance, and reconnaissance work. For scenario realism, ARKA drives the "
        "threat side: what a damaged ship's sensors report, what threat indicators the team sees, what the "
        "operational picture looks like with degraded sensors. ARKA does NOT drive repair-side realism: pump-out "
        "status, bulkhead compromise, drydock blocking, supply-chain for replacement valves. Repair-side realism "
        "requires a bridge to Navy Hull, Mechanical, and Electrical (HM&E) engineering and logistics data.")

    _body(doc,
        "Operator-side scoping direction. The most likely source for the HM&E data is Government Furnished "
        "Information from NSWC Carderock, with possible additional sources at another Naval Surface Warfare Center "
        "or somewhere in PAE Industrial Operations. This is directional input from the operator-side scoping work; "
        "the specific GFI sourcing decision is part of the executive-action set in §9. The Government Furnished "
        "Information path carries 12 to 18 month schedule risk on Authority to Operate for a gamified non-system-"
        "of-record environment, so the executive decision on the bridge approach should be made early in the "
        "engagement work, not late.")

    _h2(doc, "5.4 CACI NSS LLC vehicle holdings — the procurement-infrastructure asset")

    _body(doc,
        "CACI NSS, LLC already holds three governmentwide acquisition vehicles a Navy fleet command can issue task "
        "orders against directly:")

    _bullet(doc, "GSA OASIS — General Services Administration Federal-wide Professional Services contract.")
    _bullet(doc, "NITAAC CIO-SP3 — National Institutes of Health Information Technology Acquisition and Assessment Center IT services contract.")
    _bullet(doc, "GSA Multiple Award Schedule (GS-35F-349CA).")

    _body(doc,
        "These are governmentwide acquisition contracts open to any federal customer. The Navy can issue task "
        "orders against them without standing up a new Navy-specific contract. The AFRICOM $194M task order is the "
        "precedent — Army funded, GSA awarded, CACI executes against an OASIS-equivalent vehicle CACI already held. "
        "The Navy fleet-command equivalent is \"PACFLT N7 funds, GSA awards, CACI executes against the same vehicle.\" "
        "Same shape. Different funder.")

    _body(doc,
        "Implication for the pitch. CACI executives engaging PACFLT N7, FLTFORCOM N7, INDOPACOM J7, or SRF-JRMC "
        "leadership can offer a concrete contracting path on day one. \"You don't need to invent a new vehicle. We "
        "already hold OASIS and CIO-SP3. Issue a Level 1 pilot task order against our existing holdings. The "
        "customer commits no new procurement infrastructure.\" That is materially easier salesmanship than asking "
        "the customer to fund a new contract from scratch.")

    _pb(doc)

    # 6. Working hypothesis
    _h1(doc, "6. Working hypothesis and falsifying legs")

    _body(doc,
        "Hypothesis: the Navy has real demand for contractor-supplied operational-decision-scenario content for "
        "BDAR teams and the fleet commanders they support. The product is the five-level progression at §11.3 of "
        "the research file. CACI's right to win rests on the four assets in §5. The most-likely procurement-pathway "
        "is task order against CACI NSS LLC's existing OASIS, CIO-SP3, or GSA MAS holdings, sponsored operationally "
        "by SRF-JRMC and contracted by PACFLT N7 or FLTFORCOM N7.")

    _h2(doc, "6.1 Six falsifying legs")

    _bullet(doc,
        "Leg 1 — Navy operational-decision-scenario demand gap. Killed by senior signals turning ceremonial AND "
        "existing fleet exercises already injecting equivalent content organically.")
    _bullet(doc,
        "Leg 2 — Sub-product moat. Killed by incumbents replicating the analog form factor at the same time "
        "durations OR the AI scenario-generation capability maturing rapidly enough among incumbents that CACI's "
        "PMTEC head start does not translate to a moat by procurement time.")
    _bullet(doc,
        "Leg 3 — CACI capability transferability (PMTEC to naval-repair). Killed by PMTEC's multi-domain Combatant "
        "Command staff focus being structurally different from naval-repair small-unit decision-rehearsal.")
    _bullet(doc,
        "Leg 4 — ARKA signature-library differentiator viability. Killed by (a) IP posture blocking release; (b) "
        "comparable data available to incumbents; (c) HM&E-data bridge cannot be closed; (d) GFI / ATO friction "
        "introduces 12-18 month schedule delay neutralizing the head start.")
    _bullet(doc,
        "Leg 5 — Customer-access pathway feasibility. Killed by fleet-command incumbents locking out CACI entry "
        "timeline, or contractual right-of-first-refusal preventing new-entrant access.")
    _bullet(doc,
        "Leg 6 — Sub-product viability as fundable line. Killed by sub-product too small for standalone procurement "
        "(partner not prime) OR too specialized to attract any vehicle (need DIU / SBIR / STTR pilot first).")

    _pb(doc)

    # 7. Recommendation
    _h1(doc, "7. Recommendations for executive action")

    _body(doc,
        "Three recommendations for CACI executive decision-makers. These are not asks of the analyst who prepared "
        "the brief. They are the action items the brief is requesting CACI leadership consider.")

    _h2(doc, "7.1 Pursue the five-level progression at one initial customer")

    _body(doc,
        "Position CACI as a candidate prime for the five-level progression. Initial entry through Level 1 or Level 2 "
        "pilot. Customer end-user is SRF-JRMC. Contracting authority is PACFLT N7 / N4 or FLTFORCOM N7 / N4. "
        "Contracting vehicle is CACI NSS LLC's existing OASIS, CIO-SP3, or GSA MAS holding — no new Navy vehicle "
        "required.")

    _h2(doc, "7.2 Initiate CACI executive engagement with the named Navy commands")

    _body(doc,
        "The Navy is not running an RFP for this today. The path to procurement runs through proactive CACI "
        "executive engagement with PACFLT N7, FLTFORCOM N7, INDOPACOM J7, NAVSEA 04 (PAE-IO), or SRF-JRMC "
        "leadership. The pitch is concrete on day one: here is the work-type, here is the AFRICOM precedent "
        "demonstrating CACI executes this work at scale, here are the GSA-wide vehicles we already hold, issue a "
        "Level 1 pilot task order against our existing OASIS or CIO-SP3 holdings.")

    _h2(doc, "7.3 Decide the HM&E-data bridge approach early")

    _body(doc,
        "Repair-side scenario realism requires Navy Hull, Mechanical, and Electrical engineering and logistics "
        "data. ARKA covers threat-side only. Three candidate closures: partner with a prime that holds HM&E data; "
        "leverage CACI's existing naval-IT footprint; negotiate Government Furnished Information from a Navy source. "
        "Operator-side scoping points toward NSWC Carderock as the most likely GFI source, with possible additional "
        "sources at another Naval Surface Warfare Center or somewhere in PAE Industrial Operations. The GFI path "
        "carries 12-18 months of Authority to Operate schedule risk. CACI leadership should decide the bridge "
        "approach before pitch meetings happen, so the pitch can offer a concrete data plan rather than a "
        "placeholder.")

    _h2(doc, "7.4 Alternative entry path if direct task-order pitch fails")

    _body(doc,
        "If executive engagement does not produce a task order inside the planning horizon, the alternative entry "
        "path is DIU, SBIR, or STTR pilot for Level 1 or Level 2 scope. Lower per-engagement scale. Unblocks the "
        "no-RFP problem. Produces a reference customer for follow-on task-order competition. CACI leadership should "
        "have a parallel pilot strategy ready in case the direct task-order path stalls.")

    _pb(doc)

    # 8. Risks and flagged concerns
    _h1(doc, "8. Risks and flagged concerns")

    _h2(doc, "8.1 No existing Navy task order")
    _body(doc,
        "Biggest risk. The Navy is not procuring this sub-product today. The OASIS / CIO-SP3 finding reduces this "
        "risk because no new vehicle is required, but the Navy still has to decide to issue the task order. If "
        "executive engagement fails to produce a task-order commitment, entry collapses to DIU / SBIR / STTR pilot "
        "path or subcontract on someone else's vehicle.")

    _h2(doc, "8.2 Incumbent-lockout")
    _body(doc,
        "Fleet-command exercise-planning incumbent base is real but under-mapped. Candidate names (Booz Allen "
        "Hamilton, SAIC, HII Mission Technologies, sub-tier players) are not yet source-grounded. The actual subset "
        "active at FLTFORCOM N7, COMPACFLT N7, INDOPACOM J7, and SRF-JRMC could be either crowded enough to force "
        "partnering or sparse enough to enable first-mover prime entry. A find_sources pass authorized on 2026-05-26 "
        "is targeting this gap; results will inform v0.4.")

    _h2(doc, "8.3 HM&E-data bridge schedule risk")
    _body(doc,
        "If the executive decision on the bridge is to take the Government Furnished Information path (the operator-"
        "side direction in §5.3 points toward NSWC Carderock plus possibly another Warfare Center or PAE Industrial "
        "Operations), securing Authority to Operate for a gamified non-system-of-record environment can take 12 to "
        "18 months. Mitigations: partner with an HM&E-data-holding prime; leverage CACI's existing naval-IT "
        "footprint; or start the ATO process now in parallel with engagement work.")

    _h2(doc, "8.4 FAR 9.5 OCI consideration — flagged for executive review")

    _body(doc,
        "Operator-side knowledge of contractor exercise planners currently working at SRF-JRMC's wartime readiness "
        "group informed the analytical direction of this research without entering the vault as a citable claim "
        "(per the §9.3 contact-protection discipline documented in the research file). That awareness may trigger "
        "Organizational Conflict of Interest analysis under FAR 9.5 depending on whether CACI's competitive position "
        "in any future pursuit would unfairly benefit from operator-side knowledge of incumbent personnel.")

    _body(doc,
        "This is flagged here for executive review. The OCI analysis is a decision for CACI leadership and CACI "
        "legal / contracts, not for the analyst preparing this brief or the BD team. If executive leadership "
        "decides to advance this opportunity, the OCI question should be closed before any pursuit decision or "
        "executive engagement with the named commands. The vault's OCI primer is at `_meta/oci-primer.md`.")

    _h2(doc, "8.5 Source-grounding caveat")
    _body(doc,
        "The public sources prove the Navy has a wartime-repair problem. The public sources prove the rest of DoD "
        "buys related work-types at scale. The public sources do NOT yet prove a Navy procurement line for the "
        "corrected-scope sub-product exists today. The hypothesis is that the Navy needs this work, the rest of DoD "
        "buys this kind of work, and therefore the Navy should buy this work too. The \"should\" is real. The \"is "
        "buying today\" is the gap CACI executive salesmanship has to help close. Recommendation 7.2 is the "
        "salesmanship step.")

    _pb(doc)

    # 9. BD team next actions
    _h1(doc, "9. BD team next actions")

    _body(doc,
        "These are the operating tasks for the BD team if CACI leadership decides to advance the recommendations "
        "in §7. They do not require executive authorization individually — they are the workflow the BD team is "
        "ready to execute on green-light.")

    _bullet(doc, "BD-team gate review of this brief at v0.3 to confirm the corrected-scope frame.")
    _bullet(doc, "Run the find_sources pass against fleet-command-exercise-planning queries (authorized 2026-05-26) to map incumbents at PACFLT N7, FLTFORCOM N7, INDOPACOM J7, and SRF-JRMC. Results will be reflected in brief v0.4.")
    _bullet(doc, "Build the engagement-surface inventory (§11.1 in the research file) targeting the specific contracting offices and active solicitations at the named commands plus SYSCOM authorities (NAVSEA 04, PEO IWS).")
    _bullet(doc, "Develop the Level 1 pilot scenario portfolio — initial scenario cards for the port-selection, foreign-port contracting, and forward-team mobilization decision moments — ready for executive pitch meetings.")
    _bullet(doc, "Hold pursuit of any procurement-action until (a) CACI executive leadership commits to recommendation 7.1, 7.2, and 7.3, and (b) the FAR 9.5 OCI consideration in §8.4 is resolved by CACI leadership.")

    _pb(doc)

    # 10. Source ledger
    _h1(doc, "10. Source ledger reference")

    _body(doc,
        "Load-bearing FACT claims trace to primary or quasi-primary sources in `01_sources/`. Full ledger lives in "
        "research-file §8. Key sources:")

    _bullet(doc, "DON FY 2027 President's Budget Press Brief, 28 Apr 2026.")
    _bullet(doc, "CNO Caudle HASC Statement, 14 May 2026.")
    _bullet(doc, "SWARMEX-Cebu primary releases — Navy.mil, C7F, INDOPACOM, COMPACFLT, NavalNews, NavalToday.")
    _bullet(doc, "USNI Proceedings \"Fix the Navy's Expeditionary Repair.\"")
    _bullet(doc, "CIMSEC \"If the U.S. Navy Can't Repair Ships in Peacetime, How Will It Do So in War?\"")
    _bullet(doc, "FY27 DoD Comptroller Justification Book — MAGTF Tactical Warfare Simulation, JExD, AI scenario generation funding lines.")
    _bullet(doc, "USAspending awards: CACI NSS LLC $194M AFRICOM, Parsons Government Services $556.8M CEOIS, Axient LLC $233.8M MDA Wargames, Invictus Associates $19.4M CNRMC PSS.")
    _bullet(doc, "GovConWire — Navy nine-contractors $1.2B NAWCTSD training-systems IDIQ modification.")
    _bullet(doc, "Norfolk Naval Shipyard $442M drydock modernization (workboat.com).")
    _bullet(doc, "Pacific NAVFAC SIOP Brief — PSNS-SBS infrastructure.")
    _bullet(doc, "CACI NSS, LLC vehicle holdings — GSA OASIS, NITAAC CIO-SP3, GSA MAS GS-35F-349CA (web-verified 2026-05-26).")

    _italic(doc,
        f"End of capture brief v0.3 draft. Build date: {date.today().isoformat()}. Reframed to executive-facing "
        "structure per operator feedback. Asks-of-the-operator section removed; OCI moves to flagged concern for "
        "executive review; executive engagement moves to recommendation for executive action; HM&E data bridge "
        "updated with operator-side directional input toward NSWC Carderock and other Warfare Centers / PAE "
        "Industrial Operations.")

    doc.save(str(OUT))
    print(f"✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    build()
