#!/usr/bin/env python3
"""
build_bdr_exec_brief_v04.py — Build BDR-FLEET-READINESS executive brief v0.4.

v0.4 reframe per operator feedback 2026-05-26 (third round):
  1. Drop "Asks" framing where the operator is the implicit decision-maker. The operator is the
     analyst delivering this brief, not the decision-maker. The brief is going up the chain to
     CACI executives.
  2. FAR 9.5 OCI moves to a flagged concern for executive review.
  3. CACI executive engagement becomes a recommendation, not an ask.
  4. HM&E-data bridge updated with operator-side direction: most likely GFI source is NSWC Carderock,
     possibly another Warfare Center or PAE Industrial Operations.
"""

from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

VAULT_ROOT = Path(__file__).parent.parent
OUT = VAULT_ROOT / "opportunities/BDR-FLEET-READINESS/04_artifacts/executive-brief-v0.4-draft.docx"


def _set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _set_tight_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)


def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)


def _body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def _bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"• {text}")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_default_font(doc)
    _set_tight_margins(doc)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run("BDR-FLEET-READINESS — Executive Brief v0.4 (Draft)")
    r.font.bold = True
    r.font.size = Pt(13)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run(f"Navy BDAR decision-rehearsal. Five-level progression. Task-order against CACI's existing GSA-wide vehicles. {date.today().isoformat()}. For CACI executive review (OSI-only).")
    r.font.size = Pt(9)
    r.font.italic = True

    # BLUF
    _heading(doc, "Bottom Line Up Front")

    _body(doc,
        "The Navy has a real wartime-repair problem. The rest of the Defense Department buys related contractor "
        "work at multi-hundred-million-dollar scale. The Navy should buy this work too. Today it has not issued a "
        "task order for it. That gap is the opportunity.")

    _body(doc,
        "Recommendation. CACI pursues a five-level progression of gamified decision-rehearsal scenarios, starting "
        "analog and moving to software at Level 4. The same scenario content serves two audiences: the wardroom at "
        "a Navy repair activity preparing to receive a damaged ship, AND the fleet commander deciding where to send "
        "her. Port selection is the canonical dual-audience scenario.")

    _body(doc,
        "Procurement path is concrete. CACI NSS, LLC already holds GSA OASIS, NITAAC CIO-SP3, and GSA Multiple Award "
        "Schedule. The Navy can issue a task order against any of those vehicles. No new contracting infrastructure "
        "required. The AFRICOM $194M task order is the precedent. CACI executive leadership has to initiate the "
        "engagement with PACFLT N7, FLTFORCOM N7, INDOPACOM J7, or SRF-JRMC leadership directly. The Navy is not "
        "running an RFP for this today.")

    # Why It Matters
    _heading(doc, "Why It Matters")

    _bullet(doc,
        "Senior leadership demand signal. CNO Caudle named maintenance a warfighting requirement at HASC on 14 May "
        "2026. The FY27 budget funds $17B for ship maintenance toward an 80% Combat Surge Ready posture. SRF-JRMC "
        "ran a wartime-repair exercise with the Philippine Navy at Cebu.")

    _bullet(doc,
        "DoD buys related contractor work-types at scale. AFRICOM via CACI NSS holds $194M for broad professional "
        "services with exercise scope. MDA via Axient holds $234M for missile-defense exercise and wargames. OSD "
        "via Parsons holds $557M for real-time decision support to Combatant Commanders. The Navy could plausibly "
        "procure any of these. Today it has not.")

    _bullet(doc,
        "Navy professional community sees the gap. USNI Proceedings published \"Fix the Navy's Expeditionary "
        "Repair.\" CIMSEC published \"If the U.S. Navy Can't Repair Ships in Peacetime, How Will It Do So in War?\" "
        "Both reach the customer organizations.")

    _bullet(doc,
        "Source-grounding caveat. Public sources prove the Navy has the problem and the rest of DoD buys related "
        "work at scale. Public sources do NOT yet prove the Navy has a named procurement line for this sub-product "
        "today. The hypothesis is solid. The salesmanship step is real BD work.")

    # The Progression
    _heading(doc, "The Five-Level Progression")

    _body(doc,
        "Start analog. Move to software when scale justifies it. Each level can be procured separately. Each level "
        "builds a reference for the next.")

    _bullet(doc,
        "Level 1. Flash drills. 5 to 15 minutes. One decision on a printed scenario card. Pattern recognition.")
    _bullet(doc,
        "Level 2. Linked-decision sequences. 30 to 45 minutes. Three to five decisions in a row, each shaping the next.")
    _bullet(doc,
        "Level 3. Full 1-hour gamified wardroom session. Multi-discipline audience. All six operational-decision "
        "moments integrated.")
    _bullet(doc,
        "Level 4. Multi-session campaign. 2 to 4 hours across a deployment cycle. AI scenario generation pays for "
        "itself starting at this level.")
    _bullet(doc,
        "Level 5. Live exercise injection into COMPTUEX, RIMPAC, Large-Scale Exercise, SWARMEX. Software-driven, "
        "federated into Navy synthetic training architecture.")

    # Procurement Path
    _heading(doc, "Procurement Path — Task Order Against Existing CACI Vehicles")

    _body(doc,
        "CACI NSS, LLC already holds three governmentwide acquisition vehicles open to any federal customer "
        "including any Navy command. The Navy can issue a task order against any of these without standing up a "
        "new Navy-specific contract.")

    _bullet(doc, "GSA OASIS — Federal-wide Professional Services contract.")
    _bullet(doc, "NITAAC CIO-SP3 — IT services contract.")
    _bullet(doc, "GSA Multiple Award Schedule (GS-35F-349CA).")

    _body(doc,
        "The pitch is concrete on day one. Here is the work-type. Here is the AFRICOM precedent. Here are the "
        "vehicles we already hold. Issue a Level 1 pilot task order. The customer commits no new procurement "
        "infrastructure.")

    # Recommendations
    _heading(doc, "Recommendations for Executive Action")

    _bullet(doc,
        "Pursue the five-level progression at one initial customer. End-user: SRF-JRMC if access permits, "
        "otherwise PACFLT or Fleet Forces N7. Contracting vehicle: CACI NSS LLC's existing OASIS or CIO-SP3 "
        "holdings — no new Navy vehicle required.")

    _bullet(doc,
        "Initiate CACI executive engagement with the named Navy commands. The Navy is not running an RFP for this "
        "today. The path to procurement runs through proactive executive-level engagement at PACFLT N7, FLTFORCOM "
        "N7, INDOPACOM J7, NAVSEA 04 (PAE-IO), or SRF-JRMC leadership.")

    _bullet(doc,
        "Decide the HM&E-data bridge approach early. Repair-side scenario realism needs Navy Hull, Mechanical, and "
        "Electrical engineering and logistics data. ARKA covers threat-side only. Three options: partner with a "
        "prime that holds HM&E data; leverage CACI's existing naval IT footprint; negotiate Government Furnished "
        "Information. Operator-side direction points toward NSWC Carderock as the most likely GFI source, with "
        "possible additional sources at another Warfare Center or somewhere in PAE Industrial Operations. The GFI "
        "path carries 12-18 month Authority to Operate schedule risk. Decision is needed before pitch meetings.")

    _bullet(doc,
        "Prepare an alternative entry path through DIU, SBIR, or STTR pilot for Level 1 or Level 2 scope, ready to "
        "run if direct executive engagement does not produce a task order inside the planning horizon.")

    # Risks and Flagged Concerns
    _heading(doc, "Risks and Flagged Concerns")

    _bullet(doc,
        "No existing Navy task order. The OASIS / CIO-SP3 finding reduces this risk because no new vehicle is "
        "required, but the Navy still has to decide to issue the task order. If executive engagement fails to "
        "produce one, entry collapses to DIU / SBIR pilot path or subcontract on someone else's vehicle.")

    _bullet(doc,
        "Incumbent lockout. Fleet-command exercise-planning incumbent base is real but under-mapped. Candidate "
        "names (Booz Allen Hamilton, SAIC, HII Mission Technologies, sub-tier players) not yet source-grounded. "
        "A find_sources pass authorized on 2026-05-26 is targeting this gap; results will inform v0.5.")

    _bullet(doc,
        "FAR 9.5 OCI consideration — flagged for executive review. Operator-side knowledge of contractor exercise "
        "planners at SRF-JRMC's wartime readiness group informed the analytical direction of this research without "
        "entering the vault as a citable claim (per the §9.3 contact-protection discipline). That awareness may "
        "trigger Organizational Conflict of Interest analysis under FAR 9.5. This is a decision for CACI leadership "
        "and CACI legal / contracts, not for the analyst preparing the brief. The OCI question should be closed "
        "before any pursuit decision.")

    _bullet(doc,
        "HM&E-data bridge schedule risk. If the executive decision is the GFI path, securing Authority to Operate "
        "for a gamified non-system-of-record environment can take 12 to 18 months. Plan for this in the engagement "
        "sequencing.")

    # Footer
    f = doc.add_paragraph()
    f.paragraph_format.space_before = Pt(5)
    r = f.add_run(
        f"v0.4 built {date.today().isoformat()} from sealed §§1, 3, 4, 5, 6, 7 and §11.3 five-level progression of "
        f"00_research-file.md. Reframed to executive-facing structure per operator feedback. Asks-of-the-operator "
        f"section removed; OCI moves to flagged concern; executive engagement moves to recommendation; HM&E "
        f"updated with operator direction toward Carderock and other Warfare Centers / PAE Industrial Operations. "
        f"Capture brief v0.3 is the long-form companion."
    )
    r.font.size = Pt(8)
    r.font.italic = True

    doc.save(str(OUT))
    print(f"✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    build()
