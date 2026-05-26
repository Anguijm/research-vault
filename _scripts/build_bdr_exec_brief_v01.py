#!/usr/bin/env python3
"""
build_bdr_exec_brief_v01.py — Build BDR-FLEET-READINESS executive brief v0.1
distilled from the capture brief v0.1 per the SOP §5.5 distillation prompt at
`_meta/prompts/distillation-capture-to-exec.md`.

1-page Word artifact for a defense pure-play exec who reads on a phone.
Constraints from the distillation prompt: BLUF (3-5 sentences), Why It Matters
(3-5 bullets), Recommendation (2-3 bullets), Top Risks (2-3 bullets), Asks (2-3
bullets). No methodology, no bibliography, no hedging.
"""

from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

VAULT_ROOT = Path(__file__).parent.parent
OUT = VAULT_ROOT / "opportunities/BDR-FLEET-READINESS/04_artifacts/executive-brief-v0.1-draft.docx"


def _set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _set_tight_margins(doc):
    """1-page brief — narrow margins."""
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)


def _body_compact(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def _bullet_compact(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"• {text}")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_default_font(doc)
    _set_tight_margins(doc)

    # ── Header ────────────────────────────────────────────────────────
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run("BDR-FLEET-READINESS — Executive Brief v0.1 (Draft)")
    r.font.bold = True
    r.font.size = Pt(13)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run(f"Navy BDAR / BDAT operational-decision scenarios for repair-activity teams  |  {date.today().isoformat()}  |  For internal exec review (OSI-only)")
    r.font.size = Pt(9)
    r.font.italic = True

    # ── BLUF ──────────────────────────────────────────────────────────
    _section_heading(doc, "Bottom Line Up Front")

    _body_compact(doc,
        "The U.S. Navy is publicly signaling demand for contractor-supplied 1-hour gamified decision-rehearsal sessions "
        "for the staff cells and wardrooms of its Battle Damage Assessment and Repair (BDAR) teams at the four public "
        "naval shipyards, the Regional Maintenance Centers (RMCs), and the forward-deployed Ship Repair Facility–Japan "
        "(SRF-JRMC). CACI's right to win rests on (a) the $194M NSS LLC AFRICOM contract that demonstrates the exact "
        "work-type at sister Combatant Command scale, (b) the INDOPACOM PMTEC exercise-design capability lineage that "
        "translates down-vertical to naval repair, and (c) the March 2026 ARKA acquisition signature libraries for "
        "threat-environment realism. Recommendation: pursue initial entry as a subcontractor on a Pacific Fleet or "
        "Fleet Forces N7 / N4 exercise-planning vehicle with SRF-JRMC as operational sponsor, with a DIU / SBIR / "
        "STTR pilot as the alternative entry path.")

    # ── Why It Matters ────────────────────────────────────────────────
    _section_heading(doc, "Why It Matters")

    _bullet_compact(doc,
        "FY27 DON budget: $17B for Ship Maintenance targeting 80% Combat Surge Ready posture; $0.6B for Contested "
        "Logistics. CNO Caudle named maintenance a warfighting requirement at HASC on 14 May 2026.")

    _bullet_compact(doc,
        "Navy already executing the right exercise shape in the right venue: SRF-JRMC ran SWARMEX-Cebu on USS "
        "Ashland with Philippine Navy and host-nation contractor coordination — directly rehearsing the foreign-port "
        "emergency repair decision moments the corrected-scope product addresses.")

    _bullet_compact(doc,
        "DoD procures contractor-delivered exercise design at scale: $194M CACI NSS at AFRICOM, $556.8M Parsons CEOIS "
        "at OSD, $233.8M Axient Wargames at MDA. The FY27 comptroller justification book funds AI-powered scenario "
        "generation and a Joint Exercise Design workforce as named capability lines.")

    _bullet_compact(doc,
        "Policy-community pressure is real: USNI Proceedings (\"Fix the Navy's Expeditionary Repair\") and CIMSEC "
        "(\"If the U.S. Navy Can't Repair Ships in Peacetime, How Will It Do So in War?\") both argue wartime repair "
        "throughput is undersized — both publications read inside the customer organizations.")

    # ── Recommendation ────────────────────────────────────────────────
    _section_heading(doc, "Recommendation")

    _bullet_compact(doc,
        "Pursue the 1-hour turn-based gamified decision-rehearsal sub-product for BDAR / BDAT staff-cell and wardroom "
        "audiences. Differentiate on gamified-software engine + AI scenario generation, not 1-hour duration alone "
        "(incumbents can replicate analog tabletops at the same duration).")

    _bullet_compact(doc,
        "Two-layer customer-procurement model: SRF-JRMC wartime readiness cell as end-user / sponsor; PACFLT N7 / N4 "
        "or FLTFORCOM N7 / N4 as the contracting authority. CNRMC PSS-on-SeaPort-NxG vehicle is wrong color of money. "
        "Alternative entry path: DIU / SBIR / STTR pilot.")

    _bullet_compact(doc,
        "Close the repair-side scenario realism gap (HM&E, drydock, MFOM / NMD data) via partnering with an HM&E-data-"
        "holding prime, leveraging CACI naval-IT footprint, or negotiating Government Furnished Information — start "
        "the ATO process now in parallel with engagement work to avoid 12-18 month schedule risk.")

    # ── Top Risks ─────────────────────────────────────────────────────
    _section_heading(doc, "Top Risks")

    _bullet_compact(doc,
        "Procurement-vehicle risk — specific Navy fleet-command vehicle for the sub-product is not yet identified; if "
        "no vehicle materializes inside the planning horizon, entry collapses to pilot or subcontract paths with "
        "longer paths to material revenue.")

    _bullet_compact(doc,
        "Incumbent-lockout risk — fleet-command exercise-planning incumbent base (likely Booz Allen Hamilton, SAIC, "
        "HII Mission Technologies, plus sub-tier players) is real but under-mapped in OSI sources; whether the layer "
        "is partnering-friendly or new-entrant-hostile is the next find_sources-pass research question.")

    _bullet_compact(doc,
        "FAR 9.5 OCI risk — operator-side research origin triggers Organizational Conflict of Interest analysis (owner: "
        "operator); a finding that limits CACI's competitive position could change the recommendation shape before "
        "pursuit.")

    # ── Asks ──────────────────────────────────────────────────────────
    _section_heading(doc, "Asks")

    _bullet_compact(doc,
        "Authorize next find_sources pass targeting fleet-command-exercise-planning contracts at PACFLT N7 / N4, "
        "FLTFORCOM N7 / N4, INDOPACOM J7, and SRF-JRMC — closes the incumbent-identification research gap.")

    _bullet_compact(doc,
        "Decide on the repair-side data bridge: partner, internal naval-IT leverage, or GFI / ATO path. The GFI / "
        "ATO path requires starting now to avoid 12-18 month schedule risk.")

    _bullet_compact(doc,
        "Close the FAR 9.5 OCI analysis given operator-side research-origin knowledge of contractor exercise planners "
        "at SRF-JRMC wartime readiness group.")

    # ── Footer ────────────────────────────────────────────────────────
    f = doc.add_paragraph()
    f.paragraph_format.space_before = Pt(6)
    r = f.add_run(
        f"Derived from capture-brief-v0.1-draft.docx ({date.today().isoformat()}); both built from sealed §§1, 3, 4, 5, 6, 7 of "
        f"00_research-file.md under the 2026-05-26 corrected scope. Verifier: claude-sonnet-4-6. Cross-AI red-team: Gemini Pro."
    )
    r.font.size = Pt(8)
    r.font.italic = True

    doc.save(str(OUT))
    print(f"✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    build()
