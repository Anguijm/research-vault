#!/usr/bin/env python3
"""
build_pmtec_brief_v03.py — Roll v0.2 → v0.3 addressing 2026-05-14 v0.2 red-team.

Surgical fixes for 5 high-impact findings + 3 softer wording fixes:
1. STRIP all "AI PROMPT — …" paragraphs (template artifacts the v0.1 author
   never removed — Gemini caught this; embarrassing miss).
2. Restructure Play 3 — drop the self-licking 2-vendor consortium framing;
   recommend government run a parallel/independent OPFOR award.
3. Reconcile agentic-AI language: acknowledge CACI corporate marketing AND
   translate to operational specifics, in one paragraph rather than across
   §6.2 (still old wording) and §7.1 (denying the term).
4. Fix §10.3 ROI math: "7× ROI" → "revenue-to-B&P multiple," cost of
   delivery acknowledged.
5. Reconcile §9 ask with §10.3 Play 1 confidence: gate Play 1 capture
   activation on TAM-sizing deliverable rather than asking immediately.
6. Soften Spectral-Lite J6 over-promise.
7. Drop "likely APL" guess; replace with "JHU unit and scope to be
   confirmed via direct outreach in 30 days."
8. Acknowledge L3Harris/JFN incumbency rather than claim 30-day urgency
   will fix it.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt

VAULT_ROOT = Path(__file__).parent.parent
SRC = VAULT_ROOT / "opportunities/PMTEC-USINDOPACOM/04_artifacts/capture-brief-v0.2-draft.docx"
OUT = VAULT_ROOT / "opportunities/PMTEC-USINDOPACOM/04_artifacts/capture-brief-v0.3-draft.docx"


# ── Paragraphs to remove entirely ──────────────────────────────────────
# Match by paragraph start. All AI-PROMPT template artifacts.
REMOVE_IF_STARTS_WITH = [
    "AI PROMPT —",
    "AI PROMPT –",
    "AI PROMPT -",
]


# ── Inline text replacements (find → replace) ──────────────────────────

REPLACEMENTS = [
    # ── (1) BLUF Play 3 reframe — drop 2-vendor consortium framing ──
    (
        "lead as prime on (a) non-kinetic effects simulation using training-grade variants of Spectral/Trojan, (b) space integration via ARKA-supplied satellite sensors and ML models targeting DIU Hybrid Space Architecture in the next 30 days, and (c) C-UAS Blue-Cell Validation (CACI's SkyTracker/CORIAN/X-MADIS in their actual defensive role, with metrics tied to OPFOR drone-target performance — partnered with a swarm-target vendor like Shield AI or DroneShield).",
        "lead as prime on (a) non-kinetic effects simulation using training-grade variants of Spectral/Trojan (J6/J7 coordination required to scope feasibility), (b) space integration via ARKA-supplied satellite sensors and ML models targeting DIU Hybrid Space Architecture in the next 30 days, and (c) C-UAS Blue-Cell Validation — bid CACI's SkyTracker/CORIAN/X-MADIS in their defensive role only, and recommend the government issue a separate independent OPFOR swarm-target award (CACI explicitly does NOT propose to be both the test target and the test grader).",
    ),

    # ── (2) §7.1 Play 1 — soften Spectral-Lite J6 over-promise ──
    (
        'Plan addresses: (1) classification — partner with J6 to define a training-enclave UNCLASSIFIED reduced-fidelity dataset; (2) frequency allocation — leverage already-cleared exercise spectrum allotments at Pohakuloa and Palawan ranges; (3) deconfliction — pre-coordinate with PEO IWS for ship-borne Spectral platforms during VS28 windows.',
        'Plan acknowledges the hard constraints (and frames them as feasibility questions to be answered, not assumed): (1) classification — the government, not CACI, must define what UNCLASSIFIED reduced-fidelity dataset is releasable to a coalition training enclave; J6 partnership and IC declassification scoping is a precondition to the play, not a deliverable. (2) frequency allocation — confirm with the appropriate spectrum manager (J6/J3) whether existing exercise allotments at Pohakuloa and Palawan ranges can accommodate the planned signal injection. (3) deconfliction — PEO IWS controls ship-borne Spectral availability; coordination request, not commitment. Bottom line: this play depends on government feasibility decisions in the first 60-90 days, after which CACI builds.',
    ),

    # ── (3) §7.1 Play 3 — replace 2-vendor-consortium with independent-OPFOR model ──
    (
        "C-UAS Blue-Cell Validation: deploy SkyTracker/CORIAN/X-MADIS in their actual defensive role during PMTEC exercises, with quantified Blue-Cell detection + engagement metrics fed back to Bednarcik. White-paper proposes a 2-vendor consortium: CACI as defensive C-UAS provider + a UAS swarm vendor (Shield AI, DroneShield, AeroVironment) supplying actual OPFOR drone target swarms. This is honest — CACI's existing C-UAS products are detection/jamming systems, NOT targets, and the brief no longer claims otherwise.",
        "C-UAS Blue-Cell Validation: bid SkyTracker/CORIAN/X-MADIS strictly in their defensive role during PMTEC exercises, with quantified Blue-Cell detection + engagement metrics fed to Bednarcik. EXPLICITLY recommend the government issue a SEPARATE, INDEPENDENT OPFOR swarm-target award to a different vendor (Shield AI, DroneShield, AeroVironment, or via Replicator pathway) — CACI takes itself out of the conflict-of-interest of grading against a teammate. White paper to Bednarcik makes the independent-OPFOR recommendation explicit. CACI's offer is bounded: defensive C-UAS validation services, independent of who supplies the OPFOR swarm.",
    ),

    # ── (4) §6.2 capability map — reconcile agentic-AI language across the brief ──
    (
        "ARKA brings autonomous threat-classification models trained on satellite SIGINT — concrete inputs to digital-twin scenario generation, not generic 'agentic AI'",
        "ARKA's commercial materials market 'Agentic AI-based software' for space-based sensing; for PMTEC purposes the operational translation is: autonomous threat-classification models running over EO/IR/hyperspectral and SIGINT data, with model outputs (detections, classifications, confidence) usable as digital-twin scenario inputs and as automated training assessment signals",
    ),

    # ── (5) §7.1 Play 2 — match the same translation, drop self-contradiction ──
    (
        "Space integration / SDA PWSA pull-through: ARKA's commercial EO/IR/hyperspectral sensor datasets and autonomous threat-classification models are commercially available now — integration with legacy CACI software is NOT a prerequisite for engagement.",
        "Space integration / SDA PWSA pull-through: ARKA's commercial EO/IR/hyperspectral sensor datasets and the autonomous threat-classification software ARKA markets as 'Agentic AI-based' are commercially available now from ARKA's pre-acquisition customer base — integration with legacy CACI software is NOT a prerequisite for initial DIU engagement.",
    ),

    # ── (6) §7.2 — drop the "likely APL" guess; commit to direct verification ──
    (
        "Johns Hopkins University (announced PMTEC research partnership; specific JHU unit unverified — likely APL given defense context). APL is a UARC, not a commercial prime — they will NOT sign capture-team MOUs. The correct engagement: (1) position CACI SIGINT/EW telemetry as input data to APL's AI/digital-twin models via a non-binding data-sharing NDA; (2) engage APL as technical evaluator for any government-sponsored modeling assessment of CACI offerings; (3) DO NOT propose teaming MOU.",
        "Johns Hopkins University (Stridiron announced PMTEC AI/digital-twin research partnership at the quarterly meeting; specific JHU unit and partnership scope to be confirmed via direct outreach to Stridiron's office in the first 30 days). If the partnered unit is APL (the obvious defense-research candidate), CACI's engagement is non-traditional: APL is a UARC and will not sign capture-team MOUs, so the play is data-sharing NDA + positioning CACI SIGINT/EW telemetry as input to APL models, NOT a teaming MOU. If the partnered unit is one of JHU's other defense-relevant programs, CACI re-scopes accordingly. Either way, the 30-day outreach answers the question before we invest more capture effort here.",
    ),

    # ── (7) §7.2 Lockheed JFN — acknowledge L3Harris incumbency ──
    (
        "Lockheed JFN — open architecture, CACI EW/SIGINT plugs in cleanly. URGENCY: L3Harris demonstrated DiSCO at VS24 and is the obvious natively-integrated EW partner for primes; HII and SOSi may already be discussing exclusive teaming with L3Harris/Lockheed. CACI must initiate teaming conversations with HII (Mission Technologies leadership) and SOSi (Steve Robles, VP Coalition Solutions) within 30 days or risk being locked out of both LVC and IMN positions.",
        "Lockheed JFN — open architecture in principle, but L3Harris's DiSCO is the natively-integrated EW partner at this point (demonstrated at VS24, two years of integration reps now). CACI's realistic JFN play is NOT to displace L3Harris on the EW thread but to offer complementary SIGINT-to-targeting plug-ins on parts of the JFN architecture L3Harris isn't directly covering (e.g., post-ARKA space-based detection feeding JFN's targeting fusion). For LVC content (HII NCTE) and IMN (SOSi MPE), the teaming-letter outreach is real but acknowledges that L3Harris already has the inside track — CACI's value proposition must be additive, not substitutive.",
    ),

    # ── (8) §10.1 — narrow Play 1 TAM to PMTEC-direct ──
    (
        "Play 1 — Non-kinetic effects simulation (Spectral-Lite / Trojan-Lite): $20M–$100M TAM. Floor = single DIU OTA at PMTEC scale (~$10M, multi-year). Ceiling = multi-service expansion (Army DEVCOM EW training + Navy NIWC training + USAF 53rd EWG).",
        "Play 1 — Non-kinetic effects simulation (Spectral-Lite / Trojan-Lite): $10M–$40M TAM PMTEC-DIRECT (single OTA + 2-3 yr expansion). The earlier draft cited $100M including Army DEVCOM EW training and Navy NIWC training — those buy through separate procurement chains, NOT PMTEC, so they are properly classified as ADJACENT market opportunities (estimated additional $30M-$60M over 5 years if Play 1 PMTEC win establishes a reusable platform). The two numbers should not be conflated when sizing the PMTEC capture decision.",
    ),

    # ── (9) §10.3 ROI math — fix the language ──
    (
        'Total addressable for FY27–FY29: $95M–$525M (prime + sub). Realistic capture probability '
        'weighting (10–30% per play, varies by maturity): $20M–$140M FY27–FY29 expected revenue. '
        'Compare to B&P cost: plays 1–3 capture team activation estimated $1.5M–$3M over 12 months '
        '(2–4 full-time analyst equivalents + travel + white-paper development). ROI floor ~7× at '
        'lower-bound capture probability.',
        "Total addressable for FY27–FY29: $85M–$465M (prime + sub, after Play 1 PMTEC-direct narrowing). "
        "Realistic capture probability weighting (10–30% per play, varies by maturity): $15M–$130M "
        "FY27–FY29 expected revenue. B&P cost: plays 1–3 capture team activation estimated "
        "$1.5M–$3M over 12 months (2–4 full-time analyst equivalents + travel + white-paper development). "
        "Revenue-to-B&P multiple at the lower-bound capture probability is ~5×; this is a topline "
        "revenue multiplier, NOT a profit ROI — actual profit ROI depends on delivery cost structure "
        "for each program (typical defense-services gross margin 10-15%) and would need to be modeled "
        "in the 30-day TAM exercise. The capture decision should turn on win probability and strategic "
        "fit, not on the revenue multiple alone.",
    ),

    # ── (10) §10 confidence reconciliation note (LOW for Play 1 was the contradiction) ──
    (
        "Confidence is LOW for play 1 (PMTEC procurement model unclear), MEDIUM for plays 2 + 3 "
        "(established DIU OTA + multi-vendor partnership patterns), MEDIUM for sub plays (clear "
        "incumbent-driven sub-tier sizing).",
        "Confidence is LOW for Play 1 PMTEC procurement model (the customer has not signaled how "
        "they will buy non-kinetic effects sim — could be DIU OTA, OASIS+ task order, or new vehicle); "
        "MEDIUM for Plays 2 + 3 (established DIU OTA patterns + clear independent-OPFOR + defensive-"
        "CUAS bid structure); MEDIUM for sub plays (clear incumbent-driven sub-tier sizing). The §9 "
        "ask is correspondingly staged: Plays 2 + 3 capture activation immediately; Play 1 capture "
        "activation gated on the 30-day TAM-sizing exercise and a customer-procurement-model "
        "confirmation conversation with Brent Parker / Mary Ann Swendsen.",
    ),

    # ── (11) §9 decision asks — reconcile with §10 confidence ──
    (
        "Authorize capture team for plays 7.1.1 (non-kinetic sim) and 7.1.3 (C-UAS adversary) immediately.",
        "Authorize capture team activation for Plays 2 (ARKA-led space integration via DIU Hybrid Space Architecture) and 3 (C-UAS Blue-Cell Validation, with independent-OPFOR recommendation) IMMEDIATELY. Play 1 (non-kinetic effects sim) capture activation GATED on: (a) completion of the 30-day TAM-sizing exercise, AND (b) customer-procurement-model confirmation conversation with Brent Parker and Mary Ann Swendsen.",
    ),

    # ── (12) §8 ARKA risk — strengthen mitigation language ──
    (
        'Mitigation: pitch ARKA\'s commercially-available sensors and datasets (no internal integration required for first 6-12 months); reserve "fully integrated ARKA + CACI offering" claims until ARKA fiscal Q1 FY27 integration milestone; partner with ARKA legacy customers (e.g., NRO, SDA) to demonstrate operational continuity through the transition.',
        'Mitigation strategy is to make integration risk irrelevant to the FIRST 6-12 months of customer engagement: (1) ARKA\'s commercial sensor data and pre-acquisition customer programs continue operationally under ARKA branding through the transition; CACI does NOT present a "fully integrated CACI + ARKA" offering yet. (2) Partner with ARKA legacy government customers (NRO, SDA, Space Force) for direct operational continuity references — turn the integration question into a "no change for the customer" answer. (3) Reserve the integrated-offering pitch for after the ARKA fiscal Q1 FY27 integration milestone; require ARKA transition-lead sign-off (see §9) before any pitch claims "integrated."',
    ),

    # ── (13) §9 — add ARKA transition-lead sign-off requirement ──
    (
        "Authorize ARKA-led space play (7.1.2) — proceed IMMEDIATELY with DIU Hybrid Space Architecture engagement using ARKA's commercially-available capabilities; do NOT wait for Q3 2026 integration milestone (the demand signal is active now).",
        "Authorize ARKA-led space play (7.1.2) — proceed IMMEDIATELY with DIU Hybrid Space Architecture engagement using ARKA's commercially-available pre-acquisition capabilities; do NOT wait for the Q3 2026 integration milestone. PRECONDITION: capture-team pitch materials reviewed and signed off by the ARKA transition lead before any external use, to ensure CACI does NOT prematurely claim integration that is not yet real.",
    ),
]


# ── Helpers ─────────────────────────────────────────────────────────────

def _strip_ai_prompts(doc):
    """Remove paragraphs that start with 'AI PROMPT —' (template artifacts)."""
    body = doc.element.body
    removed = []
    for p in list(doc.paragraphs):
        text = p.text.strip()
        for prefix in REMOVE_IF_STARTS_WITH:
            if text.startswith(prefix):
                body.remove(p._p)
                removed.append(text[:80])
                break
    return removed


def _apply_replacements(doc):
    """Apply text replacements per-paragraph. First match per paragraph wins."""
    n = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for find, replace in REPLACEMENTS:
            if find in p.text:
                _set_text(p, p.text.replace(find, replace))
                n += 1
                break
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.text.strip():
                        continue
                    for find, replace in REPLACEMENTS:
                        if find in p.text:
                            _set_text(p, p.text.replace(find, replace))
                            n += 1
                            break
    return n


def _set_text(p, new_text):
    for r in list(p.runs):
        r.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


def _bump_version(doc):
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "v0.2 — draft" in p.text:
                        _set_text(p, p.text.replace(
                            "v0.2 — draft (post Gemini 3.1 Pro red-team, 2026-05-14)",
                            "v0.3 — draft (second red-team pass, 2026-05-14)",
                        ))


def main():
    print(f"Reading {SRC.name}...")
    doc = Document(SRC)

    print("Stripping AI PROMPT template artifacts...")
    removed = _strip_ai_prompts(doc)
    print(f"  removed {len(removed)} AI-PROMPT paragraph(s):")
    for line in removed:
        print(f"    · {line}…")

    print("\nApplying v0.3 corrections...")
    n = _apply_replacements(doc)
    print(f"  applied {n} text replacements")

    _bump_version(doc)

    doc.save(OUT)
    print(f"\n✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    main()
