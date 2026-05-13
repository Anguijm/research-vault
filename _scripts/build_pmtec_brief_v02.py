#!/usr/bin/env python3
"""
build_pmtec_brief_v02.py — One-shot: build capture-brief-v0.2-draft.docx from v0.1
by applying all corrections from the 2026-05-14 Gemini red-team review.

This is a single-use script for PMTEC v0.2. The "right" long-term move is to
make the brief markdown-source-of-truth and render to .docx automatically;
that's out of scope here. For now we surgically edit the v0.1 .docx.

Mapping of finding → fix is documented inline so the diff is auditable.
"""

from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VAULT_ROOT = Path(__file__).parent.parent
SRC = VAULT_ROOT / "opportunities/PMTEC-USINDOPACOM/04_artifacts/capture-brief-v0.1-draft.docx"
OUT = VAULT_ROOT / "opportunities/PMTEC-USINDOPACOM/04_artifacts/capture-brief-v0.2-draft.docx"


# ── Inline text replacements (find → replace) ──────────────────────────
# Each entry maps a substring in a v0.1 paragraph to its v0.2 replacement.
# Applied per-paragraph in order; the FIRST match wins for a paragraph.

REPLACEMENTS = [
    # ── #1 Funding inaccuracy: $9.9B → $10.0B (3 places) ──
    (
        "FY26 PDI request was $9.9B.",
        "FY26 PDI request is $10.0B (Grand Total $10,004,542K per the FY26 PDI budget book; FY25 baseline ~$9.4B).",
    ),
    (
        "FY26 PDI request: $9.9B (FY25 baseline). PMTEC is a named line item.",
        "FY26 PDI request: $10.0B ($10,004,542K Grand Total per the FY26 PDI budget book; FY25 baseline $9.4B). PMTEC is funded across Army, Navy, Air Force, and Joint Staff line items — no single consolidated PMTEC dollar.",
    ),
    (
        "PMTEC Live Fire Target Support and Joint Force Development funded across multiple PB sub-accounts (Army, Air Force, Joint exercises).",
        "Per FY26 PDI book: Navy ($588M category) funds PMTEC studies, live-fire target support, JFDD, and JDECC. Air Force ($752M) carries dedicated PMTEC Operations sub-category. Army ($851M) funds JPMRC + PMTEC range integration. Joint Staff ($310M) covers JTEEP/Theater Forces separately.",
    ),

    # ── BLUF update: drop "agentic AI" buzzword; reframe C-UAS play; bump PDI ──
    (
        "USINDOPACOM J7's Pacific Multi-Domain Training and Experimentation Capability (PMTEC) is the command's enterprise to link distributed Indo-Pacific ranges into a persistent, multi-domain training ecosystem. At its 13 March 2026 Quarterly Commercial Industry Update, the command publicly named eight technology gaps and is actively soliciting industry. CACI is already embedded in the AOR via the Deloitte-prime INDOPACOM Alpha task order (sub) and a $180M PACAF Base Area Network modernization (prime), and post-ARKA acquisition has materially stronger space and agentic-AI offerings.",
        "USINDOPACOM J7's Pacific Multi-Domain Training and Experimentation Capability (PMTEC) is the command's enterprise to link distributed Indo-Pacific ranges into a persistent, multi-domain training ecosystem. At a Q1 2026 PMTEC quarterly industry meeting, J7 named priority gaps spanning LVC integration, non-kinetic effects, multi-level secure info sharing, AI/digital-twin, realistic training targets, space integration, and RJTI. CACI is embedded in the AOR via the Deloitte-prime INDOPACOM Alpha task order (sub, $58.9M obligated of a $467M ceiling per HigherGov), the $180M PACAF AFBIM modernization (prime, awarded Sep 2025), and existing Spectral / Trojan EATS programs. The March 2026 ARKA acquisition ($2.6B close) adds satellite EO/IR/hyperspectral imaging and autonomous threat-classification software relevant to the SDA PWSA pull-through.",
    ),

    # ── §1 BLUF Recommendation: drop C-UAS-as-adversary; reframe to defensible plays ──
    (
        "Recommendation: PURSUE with a hybrid posture — lead as prime on (a) non-kinetic effects simulation, (b) space integration via ARKA, and (c) C-UAS-as-adversary-target; sub on (d) LVC content (HII), (e) coalition mission-network (SOSi), and (f) Joint Fires Network (Lockheed). Investment ask: capture team activation for plays (a)–(c), partner-positioning for (d)–(f).",
        "Recommendation: PURSUE with a hybrid posture — lead as prime on (a) non-kinetic effects simulation using training-grade variants of Spectral/Trojan, (b) space integration via ARKA-supplied satellite sensors and ML models targeting DIU Hybrid Space Architecture in the next 30 days, and (c) C-UAS Blue-Cell Validation (CACI's SkyTracker/CORIAN/X-MADIS in their actual defensive role, with metrics tied to OPFOR drone-target performance — partnered with a swarm-target vendor like Shield AI or DroneShield). Sub on (d) LVC content (HII), (e) coalition mission-network (SOSi), (f) Joint Fires Network (Lockheed). Investment ask: B&P for plays (a)–(c) capture activation; partner-positioning for (d)–(f). TAM range $80M–$450M across plays 1–3 (FY27–FY29); see §10.",
    ),

    # ── §2.2 AI/digital-twin gap: APL → Johns Hopkins University (no APL specifically) ──
    (
        "AI & digital twin technologies (new Johns Hopkins APL partnership announced March 2026)",
        "AI & digital twin technologies (new Johns Hopkins University research partnership announced by Stridiron — specific JHU unit unconfirmed in ingested sources)",
    ),

    # ── §6.1 footprint: tighten contract numerics with USAspending primary data ──
    (
        "Sub on Deloitte INDOPACOM Alpha — $467M GSA OASIS TO, period 1 Mar 2025 – 28 Feb 2030. Embedded staff at Camp H.M. Smith including PM/Exercise Planner SME.",
        "Sub on Deloitte INDOPACOM Alpha — GSA OASIS delivery order PIID 47QFCA25F0010, period start 1 Mar 2025, current end 31 Jul 2026 (Year 1), place of performance Hawaii. Obligated to date $58.9M per USAspending. $467M ceiling and 28 Feb 2030 final option-period end remain industry-reported (HigherGov) only — USAspending's base_and_all_options_value is null for this PIID.",
    ),
    (
        "ARKA acquisition closed March 2026 ($2.6B) — adds satellite sensors + agentic AI for SIGINT/EW.",
        "ARKA Group acquisition closed 9 March 2026 for $2,642.7M (net of cash acquired) per CACI 10-Q. ARKA brings satellite EO/IR and hyperspectral imaging plus autonomous threat-classification software for SIGINT/GEOINT — relevant to SDA PWSA pull-through. ARKA expected to contribute ~$150M revenue to remainder of CACI FY26 (5-month partial-year run-rate). ARKA was acquired from Blackstone Tactical Opportunities; ~1,100 ARKA employees joined CACI.",
    ),

    # ── §6.2 capability map: kill "pivot to UAS red-team stimulator" framing ──
    (
        "SkyTracker/CORIAN/X-MADIS — pivot to UAS red-team stimulator",
        "SkyTracker/CORIAN/X-MADIS are defensive C-UAS systems (detection + RF jamming) — appropriate role is Blue-Cell validation against external OPFOR target swarms, NOT as targets themselves",
    ),
    (
        "Realistic training targets",
        "Realistic training targets (defensive)",
    ),

    # ── §6.2 AI/digital twin: replace "agentic AI" with operational specifics ──
    (
        "Materially strengthened by ARKA agentic AI",
        "ARKA brings autonomous threat-classification models trained on satellite SIGINT — concrete inputs to digital-twin scenario generation, not generic 'agentic AI'",
    ),
    (
        "Post-ARKA: satellite sensors + agentic AI",
        "Post-ARKA: EO/IR/hyperspectral satellite sensors + autonomous threat-classification software (1,100 ARKA staff joined March 2026)",
    ),

    # ── §7.1 Play 1: add classification/deconfliction language ──
    (
        'Non-kinetic effects simulation: "live red-team-as-a-service" using Spectral/Trojan to inject realistic adversary EW/SIGINT/IO into PMTEC exercises (Cobra Gold, Valiant Shield, Balikatan). Direct hit on Swendsen (AI strategy) and Bednarcik (range/targets) gaps.',
        'Non-kinetic effects simulation: develop training-grade variants ("Spectral-Lite", "Trojan-Lite") that inject realistic adversary EW/SIGINT/IO signatures into PMTEC exercises (Cobra Gold, Valiant Shield, Balikatan) WITHOUT exposing classified operational waveforms. Plan addresses: (1) classification — partner with J6 to define a training-enclave UNCLASSIFIED reduced-fidelity dataset; (2) frequency allocation — leverage already-cleared exercise spectrum allotments at Pohakuloa and Palawan ranges; (3) deconfliction — pre-coordinate with PEO IWS for ship-borne Spectral platforms during VS28 windows. Direct hit on Swendsen (AI strategy) and Bednarcik (range/targets) gaps.',
    ),

    # ── §7.1 Play 2: replace "agentic AI" wording, accelerate ARKA timeline ──
    (
        "Space integration / SDA PWSA pull-through: leverage ARKA agentic AI + sensors for autonomous threat detection and missile-warning training stimulation. Engage Emslie and Hannah in parallel.",
        "Space integration / SDA PWSA pull-through: ARKA's commercial EO/IR/hyperspectral sensor datasets and autonomous threat-classification models are commercially available now — integration with legacy CACI software is NOT a prerequisite for engagement. Pursue DIU Hybrid Space Architecture demo meeting within 30 days (multiple demos underway across INDOPACOM/EUCOM/CENTCOM/SOUTHCOM). Engage Emslie and Hannah in parallel via Larry Jordan (DIU OnRamp Hub: Hawaii). Position ARKA capability as a non-traditional acquisition front-door.",
    ),

    # ── §7.1 Play 3: replace C-UAS-as-adversary nonsense with defensible play ──
    (
        "C-UAS-as-adversary: repackage SkyTracker/CORIAN/X-MADIS as adversary UAS swarm stimulator for training ranges. Engage Bednarcik directly with a white paper.",
        "C-UAS Blue-Cell Validation: deploy SkyTracker/CORIAN/X-MADIS in their actual defensive role during PMTEC exercises, with quantified Blue-Cell detection + engagement metrics fed back to Bednarcik. White-paper proposes a 2-vendor consortium: CACI as defensive C-UAS provider + a UAS swarm vendor (Shield AI, DroneShield, AeroVironment) supplying actual OPFOR drone target swarms. This is honest — CACI's existing C-UAS products are detection/jamming systems, NOT targets, and the brief no longer claims otherwise.",
    ),

    # ── §7.2 APL teaming reframe (UARC fix) ──
    (
        "Johns Hopkins APL — newly announced PMTEC partnership; position CACI data and SIGINT/EW as input layer to APL models.",
        "Johns Hopkins University (announced PMTEC research partnership; specific JHU unit unverified — likely APL given defense context). APL is a UARC, not a commercial prime — they will NOT sign capture-team MOUs. The correct engagement: (1) position CACI SIGINT/EW telemetry as input data to APL's AI/digital-twin models via a non-binding data-sharing NDA; (2) engage APL as technical evaluator for any government-sponsored modeling assessment of CACI offerings; (3) DO NOT propose teaming MOU.",
    ),

    # ── §7.2 Add competitor-locking urgency on HII/SOSi (after Lockheed JFN bullet) ──
    (
        "Lockheed JFN — open architecture, CACI EW/SIGINT plugs in cleanly.",
        "Lockheed JFN — open architecture, CACI EW/SIGINT plugs in cleanly. URGENCY: L3Harris demonstrated DiSCO at VS24 and is the obvious natively-integrated EW partner for primes; HII and SOSi may already be discussing exclusive teaming with L3Harris/Lockheed. CACI must initiate teaming conversations with HII (Mission Technologies leadership) and SOSi (Steve Robles, VP Coalition Solutions) within 30 days or risk being locked out of both LVC and IMN positions.",
    ),

    # ── §7.3 timeline: strip SAM.gov hygiene from exec asks; accelerate ARKA ──
    (
        "0–30 days: SAM.gov saved searches live; Brent Parker email; Larry Jordan coffee; Bednarcik white paper draft; INDOPACOM Alpha embeds socialize the eight gaps internally.",
        "0–30 days: Brent Parker email; Larry Jordan coffee for DIU Hybrid Space Architecture meeting; Bednarcik white paper draft (C-UAS Blue-Cell Validation + swarm-vendor consortium); HII Mission Technologies + SOSi (Robles) teaming conversations initiated; INDOPACOM Alpha embeds socialize the priority gaps internally.",
    ),
    (
        "30–90 days: white paper #1 (C-UAS as adversary target); pre-position with HII and SOSi for teaming letters; capture team activation for plays 1–3.",
        "30–90 days: HII/SOSi teaming letters signed (deadline-driven — see §7.2 urgency); white paper #2 (training-grade Spectral-Lite/Trojan-Lite concept of operations); DIU Hybrid Space Architecture demo pitch delivered; capture team activated for plays 1–3.",
    ),
    (
        "90–180 days: response to next PMTEC quarterly RFI; ARKA-led space white paper; APL teaming MOU.",
        "90–180 days: response to next PMTEC quarterly RFI; ARKA capability brief to SDA PWSA program office (NOT a teaming MOU with APL — see §7.2 UARC note); FY27 PMTEC procurement pre-positioning (RFP drop estimated Q1 FY27 based on quarterly-meeting cadence).",
    ),

    # ── §8 Risk: tighten ARKA integration risk with explicit mitigation ──
    (
        "ARKA integration risk: $2.6B acquisition just closed; cultural and technical integration not yet proven. Premature pitching of ARKA capabilities could backfire.",
        'ARKA integration risk: $2.6B acquisition closed 9 March 2026; cultural and technical integration not yet proven. Competitor-likely "ghosting" angle: "CACI is figuring out its own org chart while Lockheed/Northrop have legacy theater space presence." Mitigation: pitch ARKA\'s commercially-available sensors and datasets (no internal integration required for first 6-12 months); reserve "fully integrated ARKA + CACI offering" claims until ARKA fiscal Q1 FY27 integration milestone; partner with ARKA legacy customers (e.g., NRO, SDA) to demonstrate operational continuity through the transition.',
    ),

    # ── §9 decision asks: strip SAM.gov hygiene; add TAM-sized ask ──
    (
        "Allocate analyst time for monthly SAM.gov / DIU / DVIDS scan.",
        "Authorize TAM-sizing exercise for plays 1–3 (deliverable: 30-day numeric per-play revenue ranges with confidence intervals — see §10). Note: monthly SAM.gov / DIU / DVIDS scanning is operational hygiene already in motion via _scripts/find_sources.py — not an exec decision.",
    ),
    (
        "Authorize ARKA-led space play (7.1.2) pending ARKA integration milestone (Q3 2026).",
        "Authorize ARKA-led space play (7.1.2) — proceed IMMEDIATELY with DIU Hybrid Space Architecture engagement using ARKA's commercially-available capabilities; do NOT wait for Q3 2026 integration milestone (the demand signal is active now).",
    ),
]


# ── Structural additions ────────────────────────────────────────────────
# A new §10 TAM block to insert before the appendices, addressing the
# "no addressable market sized" exec critique.

NEW_SECTION_10_HEADING = "10. Total Addressable Market (TAM) per play — first-pass estimate"
NEW_SECTION_10_PARAGRAPHS = [
    ("BodyText", "Sensitivity: Internal"),
    ("BodyText",
     "Per-play addressable revenue estimates (FY27–FY29, 3-year horizon). Ranges reflect the gap "
     "between conservative single-vehicle scope and aggressive multi-vehicle / multi-year pull-through. "
     "Confidence is LOW for play 1 (PMTEC procurement model unclear), MEDIUM for plays 2 + 3 "
     "(established DIU OTA + multi-vendor partnership patterns), MEDIUM for sub plays (clear "
     "incumbent-driven sub-tier sizing). Refine in 30-day TAM-sizing exercise authorized in §9."),
    ("Heading 2", "10.1 Prime-led plays"),
    ("List Paragraph", "Play 1 — Non-kinetic effects simulation (Spectral-Lite / Trojan-Lite): $20M–$100M TAM. Floor = single DIU OTA at PMTEC scale (~$10M, multi-year). Ceiling = multi-service expansion (Army DEVCOM EW training + Navy NIWC training + USAF 53rd EWG)."),
    ("List Paragraph", "Play 2 — Space integration via ARKA: $50M–$300M TAM. Floor = DIU Hybrid Space Architecture demo conversion (~$15M). Ceiling = SDA PWSA pull-through (commercial-data ingestion contracts running $50M–$100M each across PWSA Tranche 2+3) + ARKA legacy contract continuity."),
    ("List Paragraph", "Play 3 — C-UAS Blue-Cell Validation: $10M–$50M TAM. Floor = single-exercise validation contract (~$5M). Ceiling = recurring multi-exercise validation services across PMTEC + JPMRC + INDOPACOM allies (Australia, Japan, Philippines C-UAS training pulls)."),
    ("Heading 2", "10.2 Sub plays"),
    ("List Paragraph", "Play 4 — LVC content sub on HII NCTE: $5M–$30M TAM (3-year sub-tier billings at typical 15–25% of HII content awards, contingent on HII winning expansion task orders)."),
    ("List Paragraph", "Play 5 — IMN sub on SOSi: $5M–$20M (ICAM/Zero Trust module sub-tier sizing on the MPE-NES IDIQ task-order ramp)."),
    ("List Paragraph", "Play 6 — Lockheed JFN sub: $5M–$25M (EW/SIGINT plug-in scope; subject to JFN architecture decisions in FY27)."),
    ("Heading 2", "10.3 Aggregate"),
    ("BodyText",
     "Total addressable for FY27–FY29: $95M–$525M (prime + sub). Realistic capture probability "
     "weighting (10–30% per play, varies by maturity): $20M–$140M FY27–FY29 expected revenue. "
     "Compare to B&P cost: plays 1–3 capture team activation estimated $1.5M–$3M over 12 months "
     "(2–4 full-time analyst equivalents + travel + white-paper development). ROI floor ~7× at "
     "lower-bound capture probability."),
    ("BodyText",
     "Caveat: these are first-pass estimates from publicly-available program scales. The 30-day "
     "TAM-sizing exercise authorized in §9 will refine with internal CACI revenue-attribution data, "
     "competitive win-loss benchmarks, and live procurement-vehicle ceilings."),
]


# ── Helpers ─────────────────────────────────────────────────────────────

def _apply_replacements(doc):
    """Apply text replacements per-paragraph. First match per paragraph wins."""
    n_changed = 0
    matched_keys = set()
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for find, replace in REPLACEMENTS:
            if find in p.text and (id(p), find) not in matched_keys:
                # Replace entire text — discards run-level formatting in matched paragraph
                # (acceptable for body paragraphs; we don't replace headings)
                _replace_paragraph_text(p, p.text.replace(find, replace))
                matched_keys.add((id(p), find))
                n_changed += 1
                break
    # Also apply to table cells
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.text.strip():
                        continue
                    for find, replace in REPLACEMENTS:
                        if find in p.text:
                            _replace_paragraph_text(p, p.text.replace(find, replace))
                            n_changed += 1
                            break
    return n_changed


def _replace_paragraph_text(p, new_text):
    """Replace paragraph text while preserving the paragraph's style."""
    # Clear existing runs
    for r in list(p.runs):
        r.text = ""
    # Put new text in first run (create one if needed)
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


def _insert_section_10_before_appendices(doc):
    """Insert §10 TAM block right before 'Appendix A — POC directory'."""
    body = doc.element.body
    target = None
    for p in doc.paragraphs:
        if p.text.startswith("Appendix A"):
            target = p._p
            break
    if target is None:
        # Fallback: append at end
        target = None

    # Build the new paragraphs (direct formatting — avoids missing-style errors)
    new_paragraphs = []

    def _p(text, *, bold=False, italic=False, size=None, indent_left=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        if bold: r.bold = True
        if italic: r.italic = True
        if size: r.font.size = Pt(size)
        if indent_left:
            from docx.shared import Inches
            p.paragraph_format.left_indent = Inches(0.25)
        return p

    # H1 — bold, large
    new_paragraphs.append(_p(NEW_SECTION_10_HEADING, bold=True, size=18))

    for style, text in NEW_SECTION_10_PARAGRAPHS:
        if style == "Heading 2":
            new_paragraphs.append(_p(text, bold=True, size=14))
        elif style == "List Paragraph":
            new_paragraphs.append(_p("• " + text, indent_left=True))
        else:
            new_paragraphs.append(_p(text))

    if target is not None:
        # Move all new paragraphs from the end (where add_paragraph put them) to just before target
        for p in new_paragraphs:
            body.remove(p._p)
            target.addprevious(p._p)


def _bump_version_table(doc):
    """Find the version table at the end and bump v0.1 → v0.2."""
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "v0.1 — draft" in p.text:
                        _replace_paragraph_text(
                            p,
                            p.text.replace(
                                "v0.1 — draft",
                                "v0.2 — draft (post Gemini 3.1 Pro red-team, 2026-05-14)",
                            ),
                        )


def main():
    print(f"Reading {SRC.name}...")
    doc = Document(SRC)

    print("Applying inline corrections from red-team findings...")
    n = _apply_replacements(doc)
    print(f"  applied {n} text replacements")

    print("Inserting new §10 (TAM per play)...")
    _insert_section_10_before_appendices(doc)

    print("Bumping version metadata v0.1 → v0.2...")
    _bump_version_table(doc)

    doc.save(OUT)
    print(f"\n✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    main()
