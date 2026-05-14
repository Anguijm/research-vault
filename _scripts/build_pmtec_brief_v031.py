#!/usr/bin/env python3
"""
build_pmtec_brief_v031.py — Roll v0.3 → v0.3.1 (drafting-hygiene patches only,
NOT strategic changes — v0.4 stays reserved for what comes back from
operator calls).

Four surgical fixes:
1. §10.1 Play 1 — drop "earlier draft cited $100M…" framing
2. §7.1 Play 1 — drop meta-narration about "frames them as feasibility questions"
3. §6.1 ARKA — strip "SIGINT/GEOINT" → "GEOINT" (Decision F scope-correction)
4. §7.2 Lockheed JFN — separate ARKA-feeds from legacy-SIGINT plug-ins
"""

from pathlib import Path
from docx import Document

VAULT_ROOT = Path(__file__).parent.parent
SRC = VAULT_ROOT / "opportunities/PMTEC-USINDOPACOM/04_artifacts/capture-brief-v0.3-draft.docx"
OUT = VAULT_ROOT / "opportunities/PMTEC-USINDOPACOM/04_artifacts/capture-brief-v0.3.1-draft.docx"


REPLACEMENTS = [
    # ── #1 §10.1 — strip "earlier draft cited $100M" framing ─────────────
    (
        "Play 1 — Non-kinetic effects simulation (Spectral-Lite / Trojan-Lite): "
        "$10M–$40M TAM PMTEC-DIRECT (single OTA + 2-3 yr expansion). The earlier "
        "draft cited $100M including Army DEVCOM EW training and Navy NIWC training "
        "— those buy through separate procurement chains, NOT PMTEC, so they are "
        "properly classified as ADJACENT market opportunities (estimated additional "
        "$30M-$60M over 5 years if Play 1 PMTEC win establishes a reusable platform). "
        "The two numbers should not be conflated when sizing the PMTEC capture decision.",

        "Play 1 — Non-kinetic effects simulation (Spectral-Lite / Trojan-Lite): "
        "$10M–$40M TAM (PMTEC-direct). Floor scenario: single DIU OTA at PMTEC "
        "scale (~$10M, multi-year). Ceiling: PMTEC-scoped 2-3 year expansion across "
        "exercises (Cobra Gold, Valiant Shield, Balikatan). Adjacent (non-PMTEC) "
        "market opportunity: $30M–$60M over 5 years if the PMTEC win establishes a "
        "reusable training-grade EW/SIGINT injection platform — addressable through "
        "separate procurement chains (Army DEVCOM EW training, Navy NIWC training, "
        "USAF 53rd EWG), which are distinct from PMTEC's vehicle. The PMTEC-direct "
        "TAM is the relevant number for the PMTEC capture decision; the adjacent "
        "number is upside contingent on a follow-on platform-reuse decision.",
    ),

    # ── #2 §7.1 Play 1 — drop meta-narration ─────────────────────────────
    (
        "Plan acknowledges the hard constraints (and frames them as feasibility "
        "questions to be answered, not assumed): (1) classification —",

        "Plan acknowledges three hard constraints, each of which is a government "
        "feasibility question rather than a CACI deliverable: (1) classification —",
    ),

    # ── #3 §6.1 ARKA — strip SIGINT scope creep (Decision F) ─────────────
    (
        "ARKA brings satellite EO/IR and hyperspectral imaging plus autonomous "
        "threat-classification software for SIGINT/GEOINT — relevant to SDA PWSA "
        "pull-through.",

        "ARKA brings satellite EO/IR and hyperspectral imaging plus autonomous "
        "threat-classification software (CACI press release describes the software "
        "as 'Agentic AI-based') for geospatial intelligence — relevant to SDA PWSA "
        "pull-through. (Scope note: ARKA does not market SIGINT; CACI's legacy "
        "SIGINT capabilities — Spectral, Trojan — are a separate stack.)",
    ),

    # ── #4 §7.2 Lockheed JFN — separate ARKA-feed from legacy-SIGINT ─────
    (
        "CACI's realistic JFN play is NOT to displace L3Harris on the EW thread "
        "but to offer complementary SIGINT-to-targeting plug-ins on parts of the "
        "JFN architecture L3Harris isn't directly covering (e.g., post-ARKA "
        "space-based detection feeding JFN's targeting fusion).",

        "CACI's realistic JFN play is NOT to displace L3Harris on the EW thread "
        "but to offer two distinct complementary feeds: (a) CACI legacy SIGINT "
        "(Spectral, Trojan) feeding JFN's targeting fusion on parts of the "
        "architecture L3Harris isn't directly covering, and (b) ARKA EO/IR + "
        "hyperspectral space-based detection feeding JFN's targeting fusion as a "
        "separate sensor track. These are two distinct stacks; do not conflate "
        "them with each other or with ARKA's product scope.",
    ),
]


def _apply_replacements(doc):
    n = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for find, replace in REPLACEMENTS:
            if find in p.text:
                _set_text(p, p.text.replace(find, replace))
                n += 1
                break
    return n


def _set_text(p, new_text):
    for r in list(p.runs):
        r.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


def _bump_version(doc):
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "v0.3 — draft" in p.text:
                        _set_text(p, p.text.replace(
                            "v0.3 — draft (second red-team pass, 2026-05-14)",
                            "v0.3.1 — draft (drafting-hygiene patch, 2026-05-14)",
                        ))


def main():
    print(f"Reading {SRC.name}...")
    doc = Document(SRC)
    n = _apply_replacements(doc)
    print(f"Applied {n} text replacements (drafting-hygiene only — no strategic changes)")
    _bump_version(doc)
    doc.save(OUT)
    print(f"\n✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    main()
