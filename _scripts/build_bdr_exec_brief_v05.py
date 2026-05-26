#!/usr/bin/env python3
"""
build_bdr_exec_brief_v05.py — Build BDR-FLEET-READINESS executive brief v0.5.

v0.5 fixes per operator feedback 2026-05-26 (fourth round):
  1. Back off the self-congratulatory tone. The brief presents a hypothesis backed by adjacent
     evidence, not a verified finding. Phrases like "concrete," "solid," "real BD work" got cut
     or softened.
  2. Expand every acronym on first appearance, format: "Spelled-out form (ACRONYM)."
  3. Drop the "operator" and "analyst" role framing. Single author voice. The OCI concern is
     rewritten so it stands without invoking either label or the vault-internal §9.3 reference.
"""

from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

VAULT_ROOT = Path(__file__).parent.parent
OUT = VAULT_ROOT / "opportunities/BDR-FLEET-READINESS/04_artifacts/executive-brief-v0.5-draft.docx"


def _set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _set_tight_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)


def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)


def _body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def _bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"• {text}")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_default_font(doc)
    _set_tight_margins(doc)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run("BDR-FLEET-READINESS — Executive Brief v0.5 (Draft)")
    r.font.bold = True
    r.font.size = Pt(13)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run(f"U.S. Navy Battle Damage Assessment and Repair (BDAR) decision-rehearsal training. Five-level progression. {date.today().isoformat()}. For CACI executive review (open-source-only).")
    r.font.size = Pt(9)
    r.font.italic = True

    # BLUF
    _heading(doc, "Bottom Line Up Front")

    _body(doc,
        "The U.S. Navy is publicly signaling demand for better wartime ship-repair readiness. The rest of the "
        "Department of Defense (DoD) buys related contractor work at multi-hundred-million-dollar scale. The Navy "
        "could plausibly buy this kind of work too. As of this draft date, public sources do not show a Navy "
        "procurement line for the specific sub-product proposed here. That gap is the opportunity this brief asks "
        "CACI leadership to consider.")

    _body(doc,
        "Proposed product. A five-level progression of gamified decision-rehearsal scenarios for Navy Battle Damage "
        "Assessment and Repair (BDAR) teams and the fleet commanders they support. The progression starts analog "
        "(printed scenario cards), builds through 1-hour wardroom sessions, and ends with software-driven scenarios "
        "injected into live fleet exercises. The same scenario content serves two audiences: the wardroom at a "
        "Navy repair activity preparing to receive a damaged ship, and the fleet commander deciding where to send "
        "her. Port selection is the canonical scenario serving both audiences.")

    _body(doc,
        "Proposed procurement path. CACI National Security Solutions, LLC (CACI NSS, LLC) already holds three "
        "General Services Administration (GSA) governmentwide contracting vehicles: GSA One Acquisition Solution "
        "for Integrated Services (OASIS), the National Institutes of Health Information Technology Acquisition "
        "and Assessment Center Chief Information Officer–Solutions and Partners 3 contract (NITAAC CIO-SP3), and "
        "GSA Multiple Award Schedule (MAS). A Navy fleet command could issue a task order against any of these "
        "vehicles without standing up a new Navy-specific contract. The $194 million task order CACI NSS, LLC "
        "holds today for support to U.S. Africa Command (AFRICOM) is the precedent for the kind of task order the "
        "Navy could issue.")

    # Why It Matters
    _heading(doc, "Why It Matters")

    _bullet(doc,
        "Senior leadership demand signal. Admiral Daryl L. Caudle, Chief of Naval Operations (CNO), named "
        "maintenance a warfighting requirement at the House Armed Services Committee (HASC) on 14 May 2026. The "
        "Department of the Navy Fiscal Year 2027 (FY27) budget funds Ship Maintenance at $17 billion toward an "
        "80-percent Combat Surge Ready (CSR) posture. The Ship Repair Facility–Japan Regional Maintenance Center "
        "(SRF-JRMC) ran a wartime-repair exercise with the Philippine Navy at Cebu in April 2026.")

    _bullet(doc,
        "DoD buys related contractor work-types at scale. AFRICOM via CACI NSS, LLC holds $194 million for broad "
        "professional services with exercise scope. The Missile Defense Agency (MDA) via Axient holds $234 million "
        "for missile-defense exercise and wargames support. The Office of the Secretary of Defense (OSD) via "
        "Parsons Government Services holds $557 million for real-time decision support to the Combatant Commands. "
        "These are three different work-types, not three examples of the same one. The Navy could plausibly procure "
        "any of them but has not yet done so under a public contract.")

    _bullet(doc,
        "Navy professional community amplifying the gap. The U.S. Naval Institute (USNI) Proceedings published "
        "\"Fix the Navy's Expeditionary Repair.\" The Center for International Maritime Security (CIMSEC) published "
        "\"If the U.S. Navy Can't Repair Ships in Peacetime, How Will It Do So in War?\" Both publications reach "
        "the customer organizations.")

    _bullet(doc,
        "Source-grounding caveat. Public sources support the hypothesis that the Navy has the problem and that the "
        "rest of DoD buys related work. Public sources do not yet demonstrate a Navy procurement line for the "
        "specific sub-product. The case for entry rests on the adjacent-precedent argument, not on a confirmed "
        "Navy demand signal.")

    # The Five-Level Progression
    _heading(doc, "The Five-Level Progression")

    _body(doc,
        "Start analog. Move to software as scale justifies it. Each level can be procured separately. Each level "
        "builds reference material for the next.")

    _bullet(doc,
        "Level 1. Flash drills. Five to fifteen minutes. One decision on a printed scenario card. Audience builds "
        "pattern recognition.")
    _bullet(doc,
        "Level 2. Linked-decision sequences. Thirty to forty-five minutes. Three to five decisions in a row, each "
        "shaping the next. Audience learns that early choices constrain later options.")
    _bullet(doc,
        "Level 3. Full 1-hour gamified wardroom session. Multi-discipline audience. All six operational-decision "
        "moments integrated into one scenario.")
    _bullet(doc,
        "Level 4. Multi-session campaign. Two to four hours across a deployment cycle. Artificial intelligence (AI) "
        "scenario generation is justified at this level because each participating wardroom needs a different "
        "variant.")
    _bullet(doc,
        "Level 5. Live exercise injection into Composite Training Unit Exercise (COMPTUEX), the Rim of the Pacific "
        "Exercise (RIMPAC), the Navy Large-Scale Exercise (LSE) series, and the Ship Wartime Repair and Maintenance "
        "Exercise (SWARMEX) program. Software-driven, federated into Navy synthetic training architecture.")

    # Procurement Path
    _heading(doc, "Procurement Path — Task Order Against an Existing CACI Vehicle")

    _body(doc,
        "CACI NSS, LLC holds three governmentwide acquisition contracts open to any federal customer, including "
        "any Navy command. A task order against any of these vehicles would not require a new Navy-specific "
        "contract.")

    _bullet(doc, "GSA OASIS — Federal-wide Professional Services contract.")
    _bullet(doc, "NITAAC CIO-SP3 — Information Technology services contract.")
    _bullet(doc, "GSA Multiple Award Schedule (contract number GS-35F-349CA).")

    _body(doc,
        "The pitch CACI executives could make at a Navy customer meeting has a defined starting point: here is "
        "the work-type, here is the AFRICOM precedent, here are the vehicles CACI already holds, the customer "
        "could issue a Level 1 pilot task order without committing to new procurement infrastructure.")

    # Recommendations
    _heading(doc, "Recommendations for Executive Action")

    _bullet(doc,
        "Pursue the five-level progression at one initial customer. Operational sponsor: SRF-JRMC, where access "
        "and operational match are strongest. Contracting authority: Commander, U.S. Pacific Fleet (PACFLT) "
        "Training and Readiness directorate (N7) or Logistics directorate (N4), or Commander, U.S. Fleet Forces "
        "Command (FLTFORCOM) N7 or N4. Contracting vehicle: CACI NSS, LLC's existing OASIS or CIO-SP3 holdings.")

    _bullet(doc,
        "Initiate CACI executive engagement with the named Navy commands. The Navy is not running a Request for "
        "Proposal (RFP) for this work today. Procurement is likely to require proactive executive-level engagement "
        "at PACFLT N7, FLTFORCOM N7, the U.S. Indo-Pacific Command Joint Operations Directorate (INDOPACOM J7), "
        "the Naval Sea Systems Command Logistics, Maintenance, and Industrial Operations Directorate (NAVSEA 04, "
        "which is the headquarters above PAE Industrial Operations or PAE-IO), or SRF-JRMC leadership.")

    _bullet(doc,
        "Decide the Hull, Mechanical, and Electrical (HM&E) data bridge approach early. Repair-side scenario "
        "realism needs Navy HM&E engineering and logistics data. ARKA (the CACI subsidiary acquired in March 2026) "
        "covers threat-environment realism only. Three options: partner with a prime that already holds HM&E data; "
        "leverage CACI's existing naval information-technology footprint; or negotiate Government Furnished "
        "Information (GFI). The most likely GFI source is Naval Surface Warfare Center Carderock Division (NSWC "
        "Carderock), with possible additional sources at another Naval Surface Warfare Center or somewhere in PAE "
        "Industrial Operations. The GFI path carries 12-to-18-month schedule risk for the Authority to Operate "
        "(ATO) on a gamified non-system-of-record environment. The decision is needed before pitch meetings.")

    _bullet(doc,
        "Prepare an alternative entry path via the Defense Innovation Unit (DIU), Small Business Innovation "
        "Research (SBIR), or Small Business Technology Transfer (STTR) for Level 1 or Level 2 scope. This is "
        "ready to run if direct executive engagement does not produce a task order inside the planning horizon.")

    # Risks and Flagged Concerns
    _heading(doc, "Risks and Flagged Concerns")

    _bullet(doc,
        "No existing Navy task order. The OASIS and CIO-SP3 vehicle holdings reduce this risk because no new "
        "vehicle is required, but the Navy still has to decide to issue the task order. If executive engagement "
        "does not produce one, entry collapses to the DIU, SBIR, or STTR pilot path or a subcontract on someone "
        "else's vehicle.")

    _bullet(doc,
        "Incumbent lockout at the fleet-command exercise-planning layer. Candidate incumbents (Booz Allen Hamilton, "
        "Science Applications International Corporation (SAIC), Huntington Ingalls Industries Mission Technologies "
        "(HII Mission Technologies), and specialized sub-tier firms) are not yet source-grounded in this research. "
        "A source-finding pass on this question was initiated 2026-05-26; results will inform a later brief "
        "version.")

    _bullet(doc,
        "Federal Acquisition Regulation (FAR) Subpart 9.5 Organizational Conflict of Interest (OCI) consideration, "
        "flagged for executive review. The analytical direction of this research is informed in part by non-public "
        "knowledge of contractor personnel currently working at SRF-JRMC's wartime readiness group. This non-public "
        "knowledge has not entered any vault content as a citable claim, but its existence may trigger an OCI "
        "analysis under FAR Subpart 9.5. CACI legal and contracts should review the OCI question before any "
        "pursuit decision or executive engagement with the named commands.")

    _bullet(doc,
        "HM&E-data bridge schedule risk. If the chosen bridge approach is the GFI path, securing the ATO for a "
        "gamified non-system-of-record environment can take 12 to 18 months. Engagement sequencing should account "
        "for this lead time.")

    # Footer
    f = doc.add_paragraph()
    f.paragraph_format.space_before = Pt(5)
    r = f.add_run(
        f"v0.5 built {date.today().isoformat()}. Distilled from the long-form capture brief (v0.4 draft). v0.5 "
        f"changes from v0.4: tone tightened away from confirmatory framing; every acronym expanded on first "
        f"appearance; single-author voice (the \"operator\" and \"analyst\" labels removed)."
    )
    r.font.size = Pt(8)
    r.font.italic = True

    doc.save(str(OUT))
    print(f"✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    build()
