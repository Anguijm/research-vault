#!/usr/bin/env python3
"""
build_pmtec_exec_brief_v031.py — Roll exec brief v0.3 → v0.3.1.

Two drafting-hygiene patches:
1. §Recommendation/sub-plays — strip operator-process detail from JHU bullet
2. §Asks — strip meta-commentary about why SAM.gov scan was removed
"""

from pathlib import Path
from docx import Document

VAULT_ROOT = Path(__file__).parent.parent
SRC = VAULT_ROOT / "opportunities/PMTEC-USINDOPACOM/04_artifacts/executive-brief-v0.3-draft.docx"
OUT = VAULT_ROOT / "opportunities/PMTEC-USINDOPACOM/04_artifacts/executive-brief-v0.3.1-draft.docx"


REPLACEMENTS = [
    # ── #1 §Recommendation/sub-plays — drop JHU operator-process detail ──
    (
        "Johns Hopkins University (specific JHU unit and partnership scope to be "
        "confirmed via direct outreach to Stridiron's office in first 30 days; "
        "not a UARC teaming MOU).",

        "Johns Hopkins University.",
    ),

    # ── #2 §Asks — strip meta-commentary about removed SAM.gov ask ──
    (
        "Authorize 30-day TAM-sizing exercise (per-play numeric revenue ranges + "
        "customer procurement-model verification). Note: SAM.gov / DIU / DVIDS "
        "scanning is already in motion via automated vault tooling — not an exec "
        "decision.",

        "Authorize 30-day TAM-sizing exercise — per-play numeric revenue ranges + "
        "customer procurement-model verification, delivered before Play 1 capture "
        "activation.",
    ),
]


def _apply_replacements(doc):
    n = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for find, replace in REPLACEMENTS:
            if find in p.text:
                _set_text(p, p.text.replace(find, replace))
                n += 1
                break
    return n


def _set_text(p, new_text):
    for r in list(p.runs):
        r.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


def main():
    print(f"Reading {SRC.name}...")
    doc = Document(SRC)
    n = _apply_replacements(doc)
    print(f"Applied {n} drafting-hygiene replacements")
    doc.save(OUT)
    print(f"\n✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    main()
