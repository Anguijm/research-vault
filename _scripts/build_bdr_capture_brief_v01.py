#!/usr/bin/env python3
"""
build_bdr_capture_brief_v01.py — Build BDR-FLEET-READINESS capture brief v0.1
from scratch, derived from the sealed §§1, 3, 4, 5, 6, 7 of 00_research-file.md
under the 2026-05-26 corrected scope.

This is a one-shot from-scratch build — there is no prior v0.1 to edit.
Subsequent versions (v0.2, v0.3 ...) should use the YAML-driven _scripts/build_brief.py
surgical-edit pattern documented at _scripts/briefs/README.md.

Structure: 10-section BD-team capture brief (target 10-12 pages):
  1. Cover & metadata
  2. Bottom Line Up Front (BLUF)
  3. Demand signal (from §3)
  4. Customer landscape (from §4)
  5. Competitive landscape (from §5)
  6. Our fit (from §6)
  7. Working hypothesis (from §7)
  8. Recommendation (from §7 recommendation draft)
  9. Risks
 10. Asks / next actions
Plus: Sources / appendix
"""

from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VAULT_ROOT = Path(__file__).parent.parent
OUT = VAULT_ROOT / "opportunities/BDR-FLEET-READINESS/04_artifacts/capture-brief-v0.1-draft.docx"


def _set_styles(doc):
    """Set default font + add heading sizes."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(18)
    return p


def _add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(14)
    return p


def _add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(12)
    return p


def _add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)
    return p


def _add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(f"• {text}")
    return p


def _add_bold(doc, text):
    """A bold-leading paragraph (like a callout)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.bold = True
    return p


def _add_italic(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.italic = True
    return p


def _add_page_break(doc):
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_styles(doc)

    # ── Cover ──────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(36)
    r = cover.add_run("CAPTURE BRIEF — v0.1 DRAFT\n")
    r.font.bold = True
    r.font.size = Pt(22)

    r = cover.add_run("\nBDR-FLEET-READINESS\n")
    r.font.bold = True
    r.font.size = Pt(18)

    r = cover.add_run("\nGamified operational-decision scenarios for Navy BDAR / BDAT teams at repair activities\n")
    r.font.size = Pt(14)
    r.font.italic = True

    r = cover.add_run(f"\nDraft date: {date.today().isoformat()}  |  Classification: For BD-team review (OSI-only).\n")
    r.font.size = Pt(10)

    r = cover.add_run("\nDerived from 00_research-file.md §§1, 3, 4, 5, 6, 7 (sealed 2026-05-26 under the small-ships-workflow pilot).  "
                       "Cross-AI red-team: Gemini Pro (3 rounds §3, 2 rounds §§4-7).  "
                       "Verifier: claude-sonnet-4-6.\n")
    r.font.size = Pt(9)
    r.font.italic = True

    _add_page_break(doc)

    # ── 1. BLUF ────────────────────────────────────────────────────────
    _add_heading_1(doc, "1. Bottom Line Up Front")

    _add_body(doc,
        "The U.S. Navy has a real and publicly-signaled demand for contractor-supplied operational-decision-scenario content "
        "for its Battle Damage Assessment and Repair (BDAR) and Battle Damage Assessment Team (BDAT) personnel at Navy repair "
        "activities — the four public naval shipyards (Norfolk, Puget Sound, Pearl Harbor, Portsmouth), the Regional Maintenance "
        "Centers (RMCs), and the forward-deployed Ship Repair Facility–Japan Regional Maintenance Center (SRF-JRMC). The "
        "differentiated sub-product CACI should pursue is a 1-hour turn-based gamified decision-rehearsal session for staff-cell "
        "and wardroom audiences, executed inside the rhythm of fleet exercise events rather than as a standalone training pipeline.")

    _add_body(doc,
        "CACI's right to win rests on three pillars: the existing CACI NSS LLC contract at AFRICOM ($194M for Plans, Operations, "
        "Logistics, Engagement, Training, Exercise, and Assessment Support — the exact work-type at multi-hundred-million-dollar "
        "scale at a sister Combatant Command); the applied capability lineage from CACI's INDOPACOM Pacific Multi-Domain Training "
        "and Experimentation Capability (PMTEC) work that translates down-vertical to the naval-repair domain; and the March 2026 "
        "ARKA acquisition signature libraries that drive threat-environment realism for the scenario content. The most-likely "
        "Navy procurement-pathway runs through the fleet-command exercise-authority layer (PACFLT N7/N4 or FLTFORCOM N7/N4) with "
        "SRF-JRMC as the operational end-user and demand sponsor, with a DIU / SBIR / STTR pilot as the secondary entry path.")

    _add_body(doc,
        "Recommendation (draft, contingent on further disconfirming-evidence resolution): position CACI as a candidate prime for "
        "the 1-hour sub-product targeting an initial entry through SRF-JRMC's wartime readiness cell as end-user with PACFLT or "
        "FLTFORCOM N7/N4 as the contracting authority. Initial entry most likely as a subcontractor on an existing fleet-command "
        "exercise-planning vehicle; path to prime contingent on a new fleet-command procurement vehicle that names the sub-product "
        "or a SBIR / STTR / DIU pilot that establishes scale and references. ARKA-side IP / release-authority work, the HM&E-data "
        "bridge approach for repair-side scenario realism, and the FAR 9.5 OCI analysis are operator-owned scoping items running "
        "parallel to the engagement and disconfirming-evidence work in section 11 of the research file.")

    _add_page_break(doc)

    # ── 2. Demand signal ───────────────────────────────────────────────
    _add_heading_1(doc, "2. Demand signal — what the Navy is publicly saying and buying")

    _add_body(doc,
        "Four independent signals converge on the corrected-scope demand. Senior Navy leadership is publicly committing dollar "
        "volume to wartime repair capability and combat surge readiness. The Navy is publicly executing the most directly "
        "relevant exercise type (Ship Wartime Repair and Maintenance Exercise, SWARMEX) in the most directly relevant venue "
        "(forward Pacific port, allied participation, host-nation contractor coordination). The defense policy community is "
        "publicly arguing the gap is real. And the Department of Defense — including the Joint Staff and the Marine Corps — "
        "procures contractor-delivered exercise design and scenario generation at funded line-item scale.")

    _add_heading_2(doc, "2.1 Senior-leadership signal")

    _add_body(doc,
        "Admiral Daryl L. Caudle, the Chief of Naval Operations, delivered a Statement on the Posture of the United States Navy "
        "before the House Armed Services Committee on 14 May 2026, naming maintenance as a warfighting requirement and "
        "announcing a planned deliberate study of Navy yard capacity. Caudle's HASC testimony frames the FY27 posture around "
        "four priorities: Lethal & Effective Force, Total Force Readiness (covering Infrastructure, Maintenance, Operations, "
        "and Spares), Capable & Resilient Warfighter, and Industrial & Logistics Capacity.")

    _add_body(doc,
        "The Department of the Navy FY 2027 President's Budget Press Brief dated 28 April 2026 funds Ship Maintenance at "
        "$17.0 billion to drive the fleet toward an \"80% Combat Surge Ready (CSR) posture by reducing maintenance delays "
        "and applying a disciplined focus across manning, training, modernization, and sustainment.\" The press brief also "
        "includes a separate $0.6 billion line item for Contested Logistics, justified explicitly because \"adversaries will "
        "target supply lines, ports, and communications.\" The DON FY27 Press Brief lists 37 Forward Deployed Naval Forces.")

    _add_heading_2(doc, "2.2 Navy public exercise events — SWARMEX-Cebu")

    _add_body(doc,
        "The single most directly scope-relevant artifact is SWARMEX — a Navy-named, Navy-executed exercise event that "
        "rehearses exactly the operational-decision moments the corrected scope's product addresses. U.S. Seventh Fleet, "
        "SRF-JRMC, U.S. Pacific Command, and U.S. Pacific Fleet all published parallel releases announcing that USS Ashland "
        "(LSD 48) completed Ship Wartime Repair and Maintenance in Cebu, Philippines, in coordination with Philippine Navy "
        "partners and local Philippine contractors. The exercise involved an amphibious warfare ship operating at a foreign "
        "port, host-nation contractor coordination for repair work, and Philippine Naval Sea Systems Command participation.")

    _add_body(doc,
        "SWARMEX-Cebu's structure maps directly onto three of the operational-decision moments the corrected-scope product "
        "addresses: forward team mobilization to where the damaged ship actually is, emergency contracting in a foreign port, "
        "and host-nation legal and operational framework navigation.")

    _add_heading_2(doc, "2.3 Policy-community pressure")

    _add_body(doc,
        "The U.S. Naval Institute Proceedings published \"Fix the Navy's Expeditionary Repair,\" arguing that the current "
        "expeditionary repair capability is inadequate to support fleet operations under wartime contestation. The Center for "
        "International Maritime Security (CIMSEC) published \"If the U.S. Navy Can't Repair Ships in Peacetime, How Will It "
        "Do So in War?\" — arguing that current peacetime repair throughput is the floor and wartime repair throughput needs "
        "to be higher. Both publications have established readership inside the customer organizations the corrected scope "
        "names as candidates.")

    _add_heading_2(doc, "2.4 Contractor-procurement signal — Joint and USMC levels")

    _add_body(doc,
        "The FY27 Department of Defense budget justification book funds contractor-delivered exercise design and scenario "
        "generation as explicit capability lines. The Marine Air-Ground Task Force (MAGTF) Tactical Warfare Simulation includes "
        "\"developing an exercise planning, design, implementation, execution, and control tool ... enabling exercise designers "
        "the ability to rapidly build new scenarios.\" The Joint Staff initiates \"design and development of a joint exercise "
        "design and control tool ... providing exercise planning, design and control within various joint simulation constructs.\" "
        "And the FY27 budget funds AI-powered scenario generation as a named capability line: \"This will be a massive education "
        "and training effort to get JExD [Joint Exercise Design] and exercise planners up to speed with new capability.\" "
        "These are direct primary-source evidence that the Department of Defense procures contractor-delivered exercise design "
        "and scenario generation at the Joint and Marine Corps levels.")

    _add_heading_2(doc, "2.5 Procurement-pathway prototype — what Navy will likely buy under")

    _add_body(doc,
        "Three non-Navy federal contracts establish the work-type at multi-hundred-million-dollar scale and tell the BD team "
        "what kind of Navy-equivalent named vehicle to expect. CACI NSS LLC holds a $194 million contract for \"Plans, "
        "Operations, Logistics, Engagement, Training, Exercise, and Assessment Support to AFRICOM\" (Army funder via GSA). "
        "Parsons Government Services holds a $556.8 million Command, Control, Communications, Computers, Combat Systems, "
        "Intelligence, Surveillance, Reconnaissance, and Targeting / Exercises / Operations / Information Services (CEOIS) "
        "task order providing decision support to Combatant Commands (OSD funder via GSA). Axient LLC holds a $233.8 million "
        "Test, Exercise and Wargames Support contract (MDA funder). When the Department of Defense decides to buy this work-"
        "type, it buys under explicitly-named exercise / wargames / decision-support vehicles. The §5 work has to identify "
        "the Navy equivalent.")

    _add_page_break(doc)

    # ── 3. Customer landscape ──────────────────────────────────────────
    _add_heading_1(doc, "3. Customer landscape")

    _add_body(doc,
        "The corrected-scope customer landscape is layered. At the strategic-direction layer sit the Secretary of the Navy "
        "and the Chief of Naval Operations — they set priorities and authorize budget direction but do not procure the "
        "corrected-scope product directly. At the repair-activity layer sit the four public naval shipyards (Norfolk, Puget "
        "Sound, Pearl Harbor, Portsmouth), the Regional Maintenance Centers including the forward-deployed SRF-JRMC, and "
        "Commander, Navy Regional Maintenance Center (CNRMC) as parent command. At the doctrinal-authority layer sits the "
        "Naval Warfare Development Center (NWDC). At the fleet-command training-authority layer sit Commander U.S. Fleet "
        "Forces Command (FLTFORCOM), Commander U.S. Pacific Fleet (COMPACFLT), Commander Naval Surface Forces (COMNAVSURFOR), "
        "and U.S. Indo-Pacific Command Joint Operations Directorate (INDOPACOM J7).")

    _add_heading_2(doc, "3.1 Procurement model — separate the end-user from the checkbook")

    _add_body(doc,
        "A critical procurement-model correction surfaced during the 2026-05-26 multi-section red-team. The operationally-most-"
        "exposed end-user (SRF-JRMC) is not the same as the contracting authority that holds the right budget color for a net-"
        "new gamified software platform. SRF-JRMC has Operations and Maintenance (O&M) dollars to fix forward-deployed ships; "
        "it does not typically possess the Research, Development, Test, and Evaluation (RDT&E) or Other Procurement, Navy (OPN) "
        "budget authority to act as the primary acquisition buyer for a new technology platform. The capture model has to "
        "decouple the two layers.")

    _add_heading_3(doc, "End-user / sponsor ranking (operational pull, requirement authorization)")
    _add_bullet(doc, "SRF-JRMC wartime readiness cell — forward-deployed, audience and venue match the corrected-scope sub-product.")
    _add_bullet(doc, "CNRMC + subordinate RMCs + Surface Team One — Navy-wide repair-activity layer, broader audience.")
    _add_bullet(doc, "Public naval shipyards (NNSY, PSNS, PNSY, PHNS) — heaviest-investment repair activities, longer cycles.")

    _add_heading_3(doc, "Contracting-authority / pathway ranking (can actually buy the sub-product)")
    _add_bullet(doc, "PACFLT N7 / N4 or FLTFORCOM N7 / N4 — owns COMPTUEX, RIMPAC, Large-Scale Exercise rhythm; has fleet-exercise contracting authority and the right budget color (OPN, RDT&E, or O&M).")
    _add_bullet(doc, "DIU, SBIR, or STTR pilot vehicle — standard new-technology-acquisition first-step path; lower per-engagement scale but unblocks the no-existing-vehicle problem.")
    _add_bullet(doc, "NAVSEA 04 or PEO IWS / equivalent SYSCOM authority — Navy training-systems acquisition apparatus; longer cycle but applicable if the sub-product enters as part of a broader training-systems program of record.")
    _add_bullet(doc, "INDOPACOM J7 joint-exercise vehicle — joint multi-domain procurement; natural path if the sub-product enters as part of a joint exercise that includes naval-repair scenarios.")
    _add_bullet(doc, "CNRMC parent-command via a non-PSS vehicle — the PSS / SeaPort-NxG IDC is wrong color of money but a separate IDC at CNRMC is feasible.")

    _add_heading_2(doc, "3.2 Doctrinal authority — NWDC as collaborator")

    _add_body(doc,
        "The Naval Warfare Development Center (NWDC) owns the doctrinal authority for Navy Tactics, Techniques, and Procedures "
        "(NTTPs) and Tactical Memorandums (TACMEMOs). NWDC is not a procurement customer for the corrected-scope product "
        "directly — it is the doctrinal-content collaborator any scenario portfolio must be written against to be operationally "
        "credible. The engagement strategy in research-file section 11 routes NWDC collaboration through Navy professional "
        "society events where NWDC publishes (American Society of Naval Engineers, Naval Engineers Journal).")

    _add_page_break(doc)

    # ── 4. Competitive landscape ───────────────────────────────────────
    _add_heading_1(doc, "4. Competitive landscape")

    _add_body(doc,
        "The competitive picture under the corrected scope splits into four layers — cross-customer Exercise / Wargames / "
        "Decision-Support incumbents that establish the work-type at federal scale, Navy training-systems incumbents adjacent "
        "but not directly competitive, fleet-command exercise-planning incumbents under-mapped in public sources but the most-"
        "likely direct-competition layer, and Navy professional-services incumbents at the corrected-scope customer.")

    _add_heading_2(doc, "4.1 Cross-customer Exercise / Wargames / Decision-Support incumbents")

    _add_body(doc,
        "CACI itself is the dominant incumbent in this layer via NSS LLC at AFRICOM ($194M). Parsons Government Services holds "
        "the OSD CEOIS task order at $556.8M. Axient LLC holds the MDA Test, Exercise and Wargames Support contract at "
        "$233.8M. None of the three are direct competitors for the corrected-scope sub-product at the Navy repair-activity "
        "level, but they are the procurement-pathway prototype that establishes what kind of Navy vehicle the BD team has to "
        "hunt for.")

    _add_heading_2(doc, "4.2 Navy training-systems incumbents — adjacent scope, not direct competition")

    _add_body(doc,
        "On 31 March 2026 the Navy modified the NAWCTSD-managed training-systems support Indefinite Delivery / Indefinite "
        "Quantity (IDIQ) vehicle to add $1.2 billion in ceiling, raising total contract value to $2.51 billion. Nine prime "
        "contractors compete for task orders — BGI-Aero Simulation JV, CAE USA, Delaware Resource Group of Oklahoma, "
        "Engineering Support Personnel, Fidelity Technologies, FlightSafety Defense, LB&B Associates, LTSS JV, and Valiant "
        "Global Defense Services. The work covers sustainment and training support for fielded simulators and instructional "
        "services — NOT the corrected-scope sub-product. The nine are adjacent incumbents (Navy training relationships, "
        "contracting vehicles, installation access) but not direct competitors. Three (CAE USA, FlightSafety Defense, BGI-"
        "Aero Simulation JV) have known capability in scenario design from adjacent aviation training contracts and could "
        "plausibly pivot if a procurement vehicle existed.")

    _add_heading_2(doc, "4.3 Fleet-command exercise-planning incumbents — under-mapped, most-likely direct-competition layer")

    _add_body(doc,
        "The largest research gap. Fleet-command exercise authorities (FLTFORCOM N7, COMPACFLT N7, INDOPACOM J7, SRF-JRMC's "
        "wartime-readiness cell) procure exercise-planning support for COMPTUEX, RIMPAC, Large-Scale Exercise, SWARMEX, and "
        "unit-level workups — but the specific contractor incumbencies are not yet source-grounded in public material this "
        "research has ingested. The typical Master Scenario Events List (MSEL) prime base under similar procurement vehicles "
        "at adjacent Combatant Commands includes Booz Allen Hamilton, SAIC, and HII Mission Technologies plus specialized "
        "exercise-design sub-tier players; the actual subset active at the named Navy commands has to be source-grounded "
        "before any partnering or capture decision. The next find_sources pass against fleet-command exercise-planning "
        "queries is the load-bearing research task to close this gap.")

    _add_heading_2(doc, "4.4 Navy PSS incumbents at the corrected-scope customer — right customer, wrong vehicle")

    _add_body(doc,
        "Invictus Associates LLC holds the Navy delivery order (PIID N0016424F3006) against parent IDC N0017819D7883 for "
        "Professional Support Services — Fleet Readiness Support for CNRMC, subordinate RMCs, and Surface Team One; total "
        "obligation $19,384,385; Norfolk VA place of performance; Department of the Navy funder. The parent IDC is the Navy's "
        "generic-services SeaPort-NxG vehicle. PSS through SeaPort-NxG buys staff augmentation, program management, "
        "administrative support, and metrics tracking; it does not by category buy wargaming or scenario design. The Invictus "
        "contract is therefore evidence the right customer (CNRMC) procures contractor support generally, and evidence the "
        "wrong vehicle (PSS on SeaPort-NxG) is the wrong vehicle for the corrected-scope product.")

    _add_page_break(doc)

    # ── 5. Our fit ─────────────────────────────────────────────────────
    _add_heading_1(doc, "5. Our fit — CACI right to win")

    _add_body(doc,
        "CACI's right to win on the corrected-scope sub-product rests on three pillars: an existing federal contract footprint "
        "in the exact work-type at multi-hundred-million-dollar scale; an applied capability lineage that translates down-"
        "vertical from the INDOPACOM J7 multi-domain exercise-design ecosystem to the naval-repair domain; and the March 2026 "
        "ARKA acquisition signature libraries that drive threat-environment realism for the scenario content. The right-to-win "
        "confidence under the corrected scope is materially higher than under the prior schoolhouse / NAWCTSD-pipeline framing "
        "the corrected scope superseded — the prior framing positioned CACI as a new entrant fighting an established training-"
        "systems pipeline; the corrected framing positions CACI as a capability-lineage holder translating existing capability "
        "down-vertical to a tighter Navy customer set.")

    _add_heading_2(doc, "5.1 CACI NSS LLC at AFRICOM — work-type at scale")

    _add_body(doc,
        "CACI NSS LLC holds a delivery order (PIID 47QFCA20F0002) against parent IDC GS00Q14OADU121 for \"Plans, Operations, "
        "Logistics, Engagement, Training, Exercise, and Assessment Support to AFRICOM\" — total obligation $194,034,792, "
        "Department of the Army funder via GSA Federal Acquisition Service awarder. The scope covers the full work-type stack "
        "named in the corrected-scope product description. The AFRICOM contract demonstrates CACI has been executing the exact "
        "work-type at multi-hundred-million-dollar scale to a Combatant Command customer since 2020. AFRICOM's combatant-"
        "command exercise support is structurally similar to fleet-command exercise support — which the §3 customer landscape "
        "names as the most-likely procurement layer for the corrected scope. The Navy fleet-command equivalents (FLTFORCOM, "
        "COMPACFLT, INDOPACOM J7) buy a similar shape of work; the gap is identifying the specific named-vehicle that maps the "
        "AFRICOM-style contracting pattern onto the Navy fleet-command structure.")

    _add_heading_2(doc, "5.2 PMTEC INDOPACOM J7 capability lineage")

    _add_body(doc,
        "CACI is an active engaged participant in the USINDOPACOM J7 Pacific Multi-Domain Training and Experimentation "
        "Capability (PMTEC) ecosystem, tracked separately in the parallel PMTEC-USINDOPACOM opportunity in this vault. The "
        "capabilities CACI demonstrates under PMTEC — Live-Virtual-Constructive (LVC) integration, scenario design for multi-"
        "domain decision-making, AI and digital-twin technologies, Regional Joint Training Infrastructure — are directly "
        "transferable down-vertical to the naval-repair domain. The corrected-scope sub-product is a focused sub-application "
        "of the broader multi-domain exercise-design capability CACI demonstrates under PMTEC. The right-to-win argument is "
        "that CACI does not have to build a new capability for the corrected-scope product — it has to translate an existing "
        "capability to a tighter vertical and a smaller delivery format.")

    _add_heading_2(doc, "5.3 ARKA — threat-environment realism, not repair-side realism")

    _add_body(doc,
        "The March 2026 ARKA acquisition gives CACI electro-optical, infrared, and hyperspectral sensor signature libraries "
        "derived from intelligence, surveillance, and reconnaissance (ISR) and space-domain work. For the corrected-scope "
        "product, signature-library content drives scenario realism on the THREAT / attack-vector side of an operational-"
        "decision rehearsal — what a damaged ship's sensors would actually report under specific damage profiles, what threat "
        "indicators the team would see in the moments after a hit, what the operational picture looks like when the ship is "
        "steaming with degraded sensors.")

    _add_body(doc,
        "ARKA does NOT directly drive REPAIR-side scenario realism — pump-out status, bulkhead compromise, drydock blocking "
        "availability, supply-chain routing for replacement valves, Material Financial Operations Management (MFOM) or Navy "
        "Maintenance Database (NMD) data. Repair-side realism requires a bridge to NAVSEA / SYSCOM technical-authority data, "
        "ship-class engineering documents, or Hull, Mechanical, and Electrical (HM&E) databases. The bridge has three candidate "
        "closures: partnering with an HM&E-data-holding prime, leveraging CACI's own naval-IT footprint, or negotiating "
        "Government Furnished Information (GFI) as part of any contract that procures the sub-product. The GFI path carries "
        "schedule risk — securing Authority to Operate (ATO) and data-sharing agreements for a gamified non-system-of-record "
        "environment can take 12-18 months.")

    _add_heading_2(doc, "5.4 OCI and contact-protection — operator-owned")

    _add_body(doc,
        "The corrected scope and the operator's research-origin context together raise specific Organizational Conflict of "
        "Interest (OCI) considerations under FAR 9.5 and contact-protection considerations under the §9.3 research-discipline "
        "rule. The OCI analysis is operator-owned per the gray-box model and is documented separately from this brief. The "
        "§9.3 discipline keeps any operator-side knowledge from contaminating analytical content; the recommendation in this "
        "brief names only OSI-sourced claims as load-bearing, and any operator-side direction informs which OSI sources to "
        "pursue without entering the recommendation chain directly.")

    _add_page_break(doc)

    # ── 6. Working hypothesis ──────────────────────────────────────────
    _add_heading_1(doc, "6. Working hypothesis and falsifying legs")

    _add_body(doc,
        "The working hypothesis is that the Navy has a real and publicly-signaled demand for contractor-supplied operational-"
        "decision-scenario content for its BDAR / BDAT teams at Navy repair activities, that the specific sub-product is a "
        "1-hour turn-based gamified decision-rehearsal session for staff-cell / wardroom audiences executed inside fleet "
        "exercise rhythm, that CACI's right to win rests on the three pillars in section 5, and that the most-likely "
        "procurement-pathway is at the fleet-command exercise-authority layer or direct SRF-JRMC wartime-readiness procurement "
        "rather than NAWCTSD or CNRMC PSS.")

    _add_body(doc,
        "The hypothesis breaks into six falsifying legs. If any one fails, the corresponding part of the hypothesis is dead "
        "and the recommendation shape has to change. The full research plan in section 10 of the research file is designed "
        "to try to break each leg.")

    _add_heading_2(doc, "6.1 Six falsifying legs")

    _add_bullet(doc, "Leg 1 — Navy operational-decision-scenario demand gap. Killed by senior signals turning ceremonial AND existing fleet exercises already injecting equivalent content organically. Counter-evidence: a fleet-command training-readiness instruction or After-Action Report (AAR) documenting existing scenario-design coverage.")
    _add_bullet(doc, "Leg 2 — Sub-product moat (gamified-software engine vs. analog tabletop). Killed by incumbents replicating the 1-hour TTX form factor with analog tabletops OR the AI scenario-generation capability maturing rapidly enough among incumbents (via the FY27 comptroller-funded JExD scenario-generation work) that CACI's PMTEC head start does not translate to a moat by procurement time.")
    _add_bullet(doc, "Leg 3 — CACI capability transferability (PMTEC to naval-repair). Killed by PMTEC's multi-domain Combatant Command staff focus being structurally different from naval-repair small-unit decision-rehearsal in ways that prevent direct capability transfer.")
    _add_bullet(doc, "Leg 4 — ARKA signature-library differentiator viability. Killed by (a) IP posture blocking release, OR (b) comparable data available to incumbents, OR (c) HM&E-data bridge cannot be closed, OR (d) GFI / ATO friction introduces 12-18 month schedule delay neutralizing the head start.")
    _add_bullet(doc, "Leg 5 — Customer-access pathway feasibility. Killed by fleet-command incumbents locking out CACI entry timeline, or contractual right-of-first-refusal preventing new-entrant access.")
    _add_bullet(doc, "Leg 6 — Sub-product viability as fundable line. Killed by (a) sub-product too small for standalone procurement (only exists embedded in multi-day contracts — partner not prime), OR (b) sub-product too specialized to attract a procurement vehicle (need DIU / SBIR / STTR pilot first).")

    _add_page_break(doc)

    # ── 7. Recommendation ──────────────────────────────────────────────
    _add_heading_1(doc, "7. Recommendation (draft)")

    _add_body(doc,
        "Position CACI as a candidate prime for the corrected-scope sub-product — 1-hour turn-based gamified decision-rehearsal "
        "sessions for BDAR / BDAT staff-cell and wardroom audiences — on a two-layer customer-procurement model.")

    _add_body(doc,
        "End-user / sponsor: SRF-JRMC wartime readiness cell, with CNRMC and the subordinate RMCs as secondary expansion targets. "
        "Contracting authority: PACFLT N7 / N4 or FLTFORCOM N7 / N4 as primary path, with DIU / SBIR / STTR pilot as the "
        "alternative entry. The CACI NSS LLC AFRICOM contract is the procurement-pathway prototype; the PMTEC INDOPACOM J7 "
        "work is the applied capability lineage; the ARKA signature-library asset is the threat-environment differentiator with "
        "the HM&E-data bridge as a separate-effort closure (partnering, internal naval-IT, or GFI / ATO negotiated as part of "
        "the sub-product procurement).")

    _add_body(doc,
        "Initial entry is most likely as a subcontractor on an existing fleet-command exercise-planning vehicle. The path to "
        "prime is contingent on either (a) a new fleet-command procurement vehicle that names the sub-product specifically, "
        "or (b) a SBIR / STTR / DIU pilot that establishes scale and references sufficient to anchor a follow-on prime "
        "engagement. The engagement strategy in research-file section 11 sequences the relationship-development work in "
        "parallel with the disconfirming-evidence checks in section 10.")

    _add_body(doc,
        "The recommendation shape will be reassessed after (a) the next find_sources pass against fleet-command-exercise-"
        "planning queries completes and produces source-grounded incumbent identification at the §4.3 / research-file §5.3 "
        "layer, (b) the engagement-surface inventory maps the specific contracting offices and active solicitations at "
        "FLTFORCOM N7 / N4, PACFLT N7 / N4, SRF-JRMC, and the relevant SYSCOM authorities (NAVSEA 04, PEO IWS / equivalent), "
        "and (c) the operator's OCI analysis closes per the FAR 9.5 procedure documented in `_meta/oci-primer.md`.")

    _add_page_break(doc)

    # ── 8. Risks ───────────────────────────────────────────────────────
    _add_heading_1(doc, "8. Risks")

    _add_heading_2(doc, "8.1 Procurement-vehicle risk")
    _add_body(doc,
        "The specific Navy fleet-command vehicle for the corrected-scope sub-product is not yet identified. If no such vehicle "
        "exists or comes online inside the planning horizon, the entry path collapses to DIU / SBIR / STTR pilot or subcontract "
        "to an existing prime — both reduce per-engagement scale and lengthen the path to material revenue.")

    _add_heading_2(doc, "8.2 Incumbent-lockout risk")
    _add_body(doc,
        "The fleet-command exercise-planning incumbent base is real but under-mapped in OSI sources. The candidate set "
        "(Booz Allen Hamilton, SAIC, HII Mission Technologies, plus specialized sub-tier players) is Assessment-tagged pending "
        "source-grounding; the actual subset active at FLTFORCOM N7, COMPACFLT N7, INDOPACOM J7, and SRF-JRMC could be either "
        "more crowded (forcing partnering) or sparser (enabling first-mover prime entry) than current evidence supports. The "
        "next find_sources pass against fleet-command exercise-planning queries is the load-bearing research task.")

    _add_heading_2(doc, "8.3 HM&E-data bridge schedule risk")
    _add_body(doc,
        "The repair-side scenario realism requires a bridge to NAVSEA / SYSCOM technical-authority data, ship-class engineering "
        "documents, or HM&E databases (MFOM, NMD) — not directly addressed by the ARKA acquisition. If the GFI / ATO path is "
        "chosen, securing Authority to Operate for a gamified non-system-of-record environment can take 12-18 months, "
        "effectively neutralizing the PMTEC-leveraged head start before procurement. Mitigations: (a) partner with an HM&E-"
        "data-holding prime; (b) leverage CACI's existing naval-IT footprint; (c) start the GFI / ATO process now, in parallel "
        "with the engagement work, rather than waiting for procurement to drive it.")

    _add_heading_2(doc, "8.4 OCI risk")
    _add_body(doc,
        "The operator's research-origin context triggers FAR 9.5 OCI analysis. The OCI analysis is operator-owned and runs "
        "separately from this brief; the brief itself uses only OSI-sourced claims as load-bearing. If the OCI analysis "
        "surfaces a finding that would limit CACI's competitive position or require a mitigation plan, the recommendation "
        "shape may need to change before pursuit.")

    _add_heading_2(doc, "8.5 Verifier and source-corpus risks")
    _add_body(doc,
        "Three small-ships-workflow pilot findings were surfaced during the §3-§7 rebuild and remain unresolved at the time of "
        "this v0.1 draft: the BDR-local FACT-label format in research-file §§8-11 (the verifier reads paragraph-leading "
        "`**FACT:**` but BDR §§8-11 still use trailing `**FACT.**`); the cross-AI temporal-mismatch pattern (Gemini's training-"
        "data state pre-2026 caused a false NO-GO on the Caudle-as-CNO claim in §3 Round 1); and the verifier truncation issue "
        "on large source files (verify_facts.py false-negatives FACT claims whose underlying quotes are late in a 3000+-line "
        "source file). None of these affect this brief's load-bearing claims; all three are recommendation items for the next "
        "research cycle.")

    _add_page_break(doc)

    # ── 9. Asks / next actions ─────────────────────────────────────────
    _add_heading_1(doc, "9. Asks / next actions")

    _add_heading_2(doc, "9.1 Asks of the BD team")
    _add_bullet(doc, "Confirm the corrected-scope frame (operational-decision-scenario, NOT schoolhouse) at the BD-team gate review. The 2026-05-26 scope correction reshapes which Navy organizations and contractors are competition vs. partner candidates.")
    _add_bullet(doc, "Authorize the next find_sources pass against fleet-command-exercise-planning queries to close the §4.3 research gap on incumbent identification at FLTFORCOM N7, COMPACFLT N7, INDOPACOM J7, and SRF-JRMC.")
    _add_bullet(doc, "Approve the engagement-surface inventory work (research-file §11.1) targeting the specific contracting offices at the named commands plus the SYSCOM authorities (NAVSEA 04, PEO IWS / equivalent).")
    _add_bullet(doc, "Decision on the HM&E-data bridge approach: partner with an HM&E-data-holding prime, leverage CACI's existing naval-IT footprint, or start the GFI / ATO process now in parallel with engagement work.")

    _add_heading_2(doc, "9.2 Asks of the operator")
    _add_bullet(doc, "Close the FAR 9.5 OCI analysis given the operator-side knowledge of contractor exercise planners at SRF-JRMC's wartime readiness group. The OCI analysis is operator-owned and runs separately from this research file.")
    _add_bullet(doc, "Confirm ARKA intellectual-property and release-authority posture for training-environment use of EO/IR/hyperspectral signature libraries.")
    _add_bullet(doc, "Decision on whether to pursue prime, sub, or pilot path as the initial entry strategy. The brief recommends sub-on-existing-vehicle as the most likely initial entry with prime contingent on a new procurement vehicle or pilot maturation.")

    _add_heading_2(doc, "9.3 Asks of the vault and tooling")
    _add_bullet(doc, "Decide on bulk-convert vs. per-section-loop for the BDR §§8-11 FACT-label format conversion (third small-ships pilot finding).")
    _add_bullet(doc, "Address the verifier-truncation issue on large source files (small-ships pilot finding #3) — chunked-context evaluation, page-range support, or per-source-section verification.")
    _add_bullet(doc, "Save the cross-AI temporal-mismatch memory note (pending operator approval).")

    _add_page_break(doc)

    # ── 10. Source ledger reference ────────────────────────────────────
    _add_heading_1(doc, "10. Source ledger reference")

    _add_body(doc,
        "All load-bearing FACT claims in this brief are sourced from primary or quasi-primary documents ingested in the "
        "BDR-FLEET-READINESS opportunity's `01_sources/` directory. The full source ledger lives in research-file §8 with "
        "citation tags in the format [s.YYYY-MM-DD-slug]. Key sources supporting this brief's load-bearing claims:")

    _add_bullet(doc, "DON FY 2027 President's Budget Press Brief, 28 Apr 2026 — Ship Maintenance $17.0B / Combat Surge Ready posture / Contested Logistics $0.6B / SIOP $1.8B / MIB $7.0B. [s.2026-05-26-don-fy27-press-brief]")
    _add_bullet(doc, "CNO Caudle HASC Statement, 14 May 2026 — maintenance as warfighting requirement / four FY27 priorities / planned deliberate study of Navy yard capacity. [s.2026-05-26-2026-05-14-caudle-testimony]")
    _add_bullet(doc, "SWARMEX-Cebu primary press releases — Navy.mil, C7F, USINDOPACOM, COMPACFLT, NavalNews, NavalToday — USS Ashland wartime repair in Cebu with Philippine Navy + host-nation contractor coordination.")
    _add_bullet(doc, "USNI Proceedings — \"Fix the Navy's Expeditionary Repair\". [s.2026-05-26-fix-the-navys-expeditionary-re]")
    _add_bullet(doc, "CIMSEC — \"If the U.S. Navy Can't Repair Ships in Peacetime, How Will It Do So in War?\". [s.2026-05-26-if-the-u-s-navy-cant-repair-sh]")
    _add_bullet(doc, "FY27 DoD Comptroller Justification Book — MAGTF Tactical Warfare Simulation / JExD / AI scenario generation funded as contractor-procurable lines. [s.2026-05-26-justification-book]")
    _add_bullet(doc, "USAspending awards: CACI NSS LLC $194M AFRICOM, Parsons Government Services $556.8M CEOIS, Axient LLC $233.8M MDA Wargames, Invictus Associates $19.4M CNRMC PSS. (Four sources in 01_sources/_quarantine/ replaced by API-fetched records on 2026-05-26.)")
    _add_bullet(doc, "GovConWire — Navy nine-contractors $1.2B NAWCTSD training-systems IDIQ modification. [s.2026-05-26-navy-selects-nine-contractors-1-2b]")
    _add_bullet(doc, "Norfolk Naval Shipyard $442M drydock modernization (workboat.com). [s.2026-05-26-norfolk-naval-shipyard-begins-]")
    _add_bullet(doc, "Pacific NAVFAC SIOP Brief — PSNS-SBS infrastructure. [s.2026-05-26-siop-brief-psns-sbs-2025-1]")

    _add_body(doc,
        "Verifier run 2026-05-26: 20 SUPPORTS, 14 PARTIAL, 4 DOES_NOT_SUPPORT (three of which are verifier-truncation false "
        "negatives against the 3822-line comptroller-book source file; one is a citation-resolution issue on the SWARMEX "
        "duplicate-slug cluster). Cross-AI red-team Gemini Pro: 3 rounds on §3 (GO), 2 rounds on §§4-7 (SEAL after micro-fix). "
        "Both red-team dialogues archived in opportunity's `_red-teams/` folder.")

    _add_italic(doc,
        f"End of capture brief v0.1 draft. Build date: {date.today().isoformat()}. Built from sealed research-file sections; "
        "all claims trace to citations in 01_sources/. Next version expected after BD-team gate review.")

    doc.save(str(OUT))
    print(f"✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    build()
