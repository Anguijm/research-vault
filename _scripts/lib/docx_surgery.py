"""
docx_surgery.py — Reusable helpers for surgical edits to a Word document.

Extracted from the five build_pmtec_*.py scripts that share the same
edit-the-previous-version-and-write-the-next-one pattern. The point of pulling
these out is so the per-version drivers shrink to a YAML config plus a single
call to apply_operations(), rather than each one re-implementing the same
helper functions.

See _scripts/build_brief.py for the driver that consumes these helpers, and
_scripts/briefs/ for the YAML config schema.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Inches, Pt


# ── Paragraph text manipulation ────────────────────────────────────────────


def set_paragraph_text(p, new_text: str) -> None:
    """Replace a paragraph's text while preserving the paragraph's style.

    Clears all existing runs and puts the new text into the first run, which
    drops run-level formatting inside the matched paragraph. That tradeoff
    matches the behavior of every build_pmtec_*.py script that came before
    this module; do not change it without checking those scripts first.
    """
    for r in list(p.runs):
        r.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


# ── Operation: apply text replacements ─────────────────────────────────────


def apply_text_replacements(
    doc,
    replacements: Iterable[tuple[str, str]],
) -> int:
    """Apply (find, replace) substitutions per-paragraph. First match wins.

    Walks both body paragraphs and table-cell paragraphs. Returns the count
    of paragraphs modified. A single paragraph is modified at most once per
    call (first matching replacement wins), preserving the behavior of the
    pre-consolidation scripts.
    """
    n = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for find, replace in replacements:
            if find in p.text:
                set_paragraph_text(p, p.text.replace(find, replace))
                n += 1
                break
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.text.strip():
                        continue
                    for find, replace in replacements:
                        if find in p.text:
                            set_paragraph_text(p, p.text.replace(find, replace))
                            n += 1
                            break
    return n


# ── Operation: strip paragraphs by prefix ──────────────────────────────────


def strip_paragraphs_starting_with(
    doc,
    prefixes: Iterable[str],
) -> list[str]:
    """Remove every body paragraph whose stripped text starts with any prefix.

    Used to strip template artifacts like "AI PROMPT —" lines that the
    drafting model failed to remove. Returns the truncated first 80 chars
    of each removed paragraph for caller logging.

    Does NOT walk table cells; the original scripts only stripped body
    paragraphs and we preserve that behavior.
    """
    body = doc.element.body
    removed: list[str] = []
    prefixes = tuple(prefixes)
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if text.startswith(prefixes):
            body.remove(p._p)
            removed.append(text[:80])
    return removed


# ── Operation: insert a new section before an anchor paragraph ─────────────

# Paragraph style names allowed in the insert_sections config. Each maps to a
# concrete formatting recipe applied via direct formatting (NOT named styles),
# which is what the v02 build script does to dodge missing-style errors when
# the source docx was produced from a template that lacks Word's default
# Heading 1 / Heading 2 styles.

_PARAGRAPH_STYLES = {
    "heading-1": {"bold": True, "size_pt": 18},
    "heading-2": {"bold": True, "size_pt": 14},
    "body": {},
    "bullet": {"bullet": True, "indent_inches": 0.25},
    "italic": {"italic": True},
}


def insert_section(
    doc,
    *,
    before_paragraph_starting_with: str | None,
    paragraphs: list[dict],
) -> bool:
    """Insert a sequence of paragraphs into the document.

    If `before_paragraph_starting_with` is given and an anchor is found,
    the new paragraphs are moved immediately before that anchor. Otherwise
    they are appended at the end of the document.

    `paragraphs` is a list of dicts with keys:
        style: one of _PARAGRAPH_STYLES (default "body")
        text:  the paragraph text

    Returns True if an anchor was found and used, False if appended at end.
    """
    body = doc.element.body
    anchor = None
    if before_paragraph_starting_with:
        for p in doc.paragraphs:
            if p.text.startswith(before_paragraph_starting_with):
                anchor = p._p
                break

    created = []
    for item in paragraphs:
        style_name = item.get("style", "body")
        if style_name not in _PARAGRAPH_STYLES:
            raise ValueError(
                f"Unknown paragraph style {style_name!r}; "
                f"valid: {sorted(_PARAGRAPH_STYLES)}"
            )
        recipe = _PARAGRAPH_STYLES[style_name]
        text = item.get("text", "")
        if recipe.get("bullet"):
            text = "• " + text
        p = doc.add_paragraph()
        r = p.add_run(text)
        if recipe.get("bold"):
            r.bold = True
        if recipe.get("italic"):
            r.italic = True
        if "size_pt" in recipe:
            r.font.size = Pt(recipe["size_pt"])
        if "indent_inches" in recipe:
            p.paragraph_format.left_indent = Inches(recipe["indent_inches"])
        created.append(p)

    if anchor is not None:
        for p in created:
            body.remove(p._p)
            anchor.addprevious(p._p)
        return True
    return False


# ── Operation: bump a version-stamp string in a table cell ─────────────────


def bump_version_in_tables(
    doc,
    *,
    find: str,
    replace: str,
) -> int:
    """Replace a version stamp anywhere it appears inside a table cell.

    The version table at the end of each brief carries strings like
    "v0.3 — draft (second red-team pass, 2026-05-14)" — this helper updates
    just that table content without touching body paragraphs.

    Returns the count of table-cell paragraphs modified.
    """
    n = 0
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if find in p.text:
                        set_paragraph_text(p, p.text.replace(find, replace))
                        n += 1
    return n


# ── Top-level: apply a sequence of operations described in a config dict ───


def apply_operations(doc, config: dict) -> dict:
    """Apply the standard sequence of operations to a docx document.

    The order is fixed: strip → replace → insert → bump. This matches the
    order used by every pre-consolidation script. Returns a summary dict of
    counts and any anchor-not-found warnings, so the calling driver can log
    a clean per-operation report.

    Expected config keys (all optional):
      - remove_paragraphs_starting_with: list[str]
      - text_replacements: list[{find: str, replace: str}]
      - insert_sections: list[{before_paragraph_starting_with: str, paragraphs: list}]
      - version_bump: {find: str, replace: str}
    """
    summary: dict = {
        "stripped": [],
        "replaced": 0,
        "inserted": [],
        "version_bumped": 0,
        "warnings": [],
    }

    prefixes = config.get("remove_paragraphs_starting_with") or []
    if prefixes:
        summary["stripped"] = strip_paragraphs_starting_with(doc, prefixes)

    replacements_raw = config.get("text_replacements") or []
    if replacements_raw:
        pairs = [(r["find"], r["replace"]) for r in replacements_raw]
        summary["replaced"] = apply_text_replacements(doc, pairs)

    for section in config.get("insert_sections") or []:
        anchor = section.get("before_paragraph_starting_with")
        paragraphs = section.get("paragraphs") or []
        found = insert_section(
            doc,
            before_paragraph_starting_with=anchor,
            paragraphs=paragraphs,
        )
        if anchor and not found:
            summary["warnings"].append(
                f"insert_sections: anchor not found, appended at end: {anchor!r}"
            )
        summary["inserted"].append(
            {"anchor": anchor, "paragraph_count": len(paragraphs), "found_anchor": found}
        )

    bump = config.get("version_bump")
    if bump:
        summary["version_bumped"] = bump_version_in_tables(
            doc, find=bump["find"], replace=bump["replace"]
        )

    return summary
