#!/usr/bin/env python3
"""
build_pmtec_exec_brief_v03.py — Roll executive-brief-v0.1 → v0.3.

Applies the corresponding v0.2 + v0.3 corrections (PDI figure, ARKA scope,
C-UAS framing, JHU APL → JHU general, drop AI PROMPT artifact, reconcile
§9 ask with v0.3 Play-1 gating) to the 1-page executive brief.
"""

from pathlib import Path
from docx import Document

VAULT_ROOT = Path(__file__).parent.parent
SRC = VAULT_ROOT / "opportunities/PMTEC-USINDOPACOM/04_artifacts/executive-brief-v0.1-draft.docx"
OUT = VAULT_ROOT / "opportunities/PMTEC-USINDOPACOM/04_artifacts/executive-brief-v0.3-draft.docx"

REMOVE_IF_STARTS_WITH = ["AI PROMPT —", "AI PROMPT –", "AI PROMPT -"]

REPLACEMENTS = [
    # PDI funding figure: $9.9B → $10.0B
    (
        "PMTEC sits inside the Pacific Deterrence Initiative (FY26 PB: $9.9B). Long-tail funding through end of decade.",
        "PMTEC sits inside the Pacific Deterrence Initiative (FY26 PDI request: $10.0B per the FY26 comptroller PDI book; FY25 baseline $9.4B). Long-tail funding through end of decade.",
    ),

    # BLUF: drop "agentic AI" + "C-UAS-as-adversary-target"
    (
        "CACI is already embedded (sub on $467M Deloitte INDOPACOM Alpha; prime on $180M PACAF AFBIM) and post-ARKA acquisition has materially stronger space and agentic-AI offerings. Three plays favor CACI as prime — non-kinetic effects simulation, space integration via ARKA, C-UAS-as-adversary-target — with sub/teaming positions on three additional plays.",
        "CACI is already embedded (sub on Deloitte INDOPACOM Alpha PIID 47QFCA25F0010, $58.9M obligated of an industry-reported $467M ceiling; prime on $180M PACAF AFBIM) and the March 2026 ARKA acquisition ($2.6B close) adds EO/IR + hyperspectral satellite sensors plus autonomous threat-classification software. Three plays favor CACI as prime — non-kinetic effects simulation (training-grade), space integration via ARKA, and C-UAS Blue-Cell validation — with sub/teaming positions on three additional plays.",
    ),

    # "Why it matters" — strip agentic-AI buzzword, fix C-UAS framing
    (
        "Three of the eight named gaps map directly to CACI's existing weapon systems (Spectral, Trojan, SkyTracker) and one to ARKA's just-acquired capability.",
        "Three of the eight named PMTEC priorities map to CACI capabilities (Spectral/Trojan for non-kinetic effects; SkyTracker/CORIAN/X-MADIS for defensive C-UAS validation) and the space-integration priority maps to ARKA's pre-acquisition EO/IR + hyperspectral product line.",
    ),

    # Recommendation: drop C-UAS adversary + APL teaming MOU
    (
        "PURSUE as prime: (1) non-kinetic effects simulation using Spectral/Trojan; (2) space integration via ARKA + SDA PWSA; (3) C-UAS-as-adversary repackaging.",
        "PURSUE as prime: (1) non-kinetic effects simulation using training-grade Spectral/Trojan variants (Play 1 capture activation GATED on 30-day TAM exercise + customer-procurement-model call with Brent Parker); (2) space integration via ARKA + SDA PWSA pull-through (proceed immediately via DIU Hybrid Space Architecture); (3) C-UAS Blue-Cell Validation in defensive role, with explicit recommendation that government issue a separate independent OPFOR swarm-target award.",
    ),
    (
        "PURSUE as sub/teammate: HII (LVC content), SOSi (coalition mission network), Lockheed (Joint Fires Network), Johns Hopkins APL (AI/digital twin).",
        "PURSUE as sub/teammate: HII (LVC content), SOSi (coalition mission network), Lockheed JFN (post-ARKA space-to-JFN feeds — NOT EW-thread displacement; L3Harris incumbent), Johns Hopkins University (specific JHU unit and partnership scope to be confirmed via direct outreach to Stridiron's office in first 30 days; not a UARC teaming MOU).",
    ),

    # Top risks: tighten ARKA + add Spectral-Lite caveat
    (
        "ARKA integration not yet proven; premature ARKA-led pitching could backfire. Mitigation: Q3 2026 milestone gate before activating space play.",
        "ARKA integration just closed 9 March 2026; competitors will frame as \"unintegrated.\" Mitigation: pitch ARKA's commercial pre-acquisition capabilities (no internal-integration claims for 6-12 months); ARKA transition lead signs off on pitch materials before external use.",
    ),

    # Asks: reconcile with v0.3 Play-1 gating; strip SAM.gov hygiene
    (
        "Approve capture team activation for prime plays 1 and 3 immediately; play 2 pending ARKA milestone.",
        "Approve capture team activation for Plays 2 (ARKA space integration via DIU HSA) and 3 (C-UAS Blue-Cell Validation, independent OPFOR recommendation) IMMEDIATELY. Play 1 (non-kinetic effects sim) capture activation GATED on 30-day TAM-sizing exercise + customer-procurement-model confirmation with Brent Parker.",
    ),
    (
        "Allocate analyst time for monthly SAM.gov / DIU / DVIDS opportunity scan.",
        "Authorize 30-day TAM-sizing exercise (per-play numeric revenue ranges + customer procurement-model verification). Note: SAM.gov / DIU / DVIDS scanning is already in motion via automated vault tooling — not an exec decision.",
    ),
]


def _strip_ai_prompts(doc):
    body = doc.element.body
    removed = []
    for p in list(doc.paragraphs):
        text = p.text.strip()
        for prefix in REMOVE_IF_STARTS_WITH:
            if text.startswith(prefix):
                body.remove(p._p)
                removed.append(text[:80])
                break
    return removed


def _apply_replacements(doc):
    n = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for find, replace in REPLACEMENTS:
            if find in p.text:
                _set_text(p, p.text.replace(find, replace))
                n += 1
                break
    return n


def _set_text(p, new_text):
    for r in list(p.runs):
        r.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


def main():
    print(f"Reading {SRC.name}...")
    doc = Document(SRC)

    removed = _strip_ai_prompts(doc)
    print(f"Stripped {len(removed)} AI-PROMPT paragraph(s)")
    for line in removed:
        print(f"  · {line}…")

    n = _apply_replacements(doc)
    print(f"Applied {n} text replacements")

    doc.save(OUT)
    print(f"\n✓ Wrote {OUT.relative_to(VAULT_ROOT)}")


if __name__ == "__main__":
    main()
